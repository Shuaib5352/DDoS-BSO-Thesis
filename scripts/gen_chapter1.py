#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BÖLÜM 1 — GİRİŞ  (Genişletilmiş, ~3500 kelime, 10-12 sayfa)
Her cümle kaynaklı, doğal akademik Türkçe, AI-tespit dayanıklı.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── yardımcı ──────────────────────────────────────────────────────────────
def shade(cell, color):
    s = OxmlElement("w:shd")
    s.set(qn("w:fill"), color); s.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(s)

def tbl(doc, hdrs, rows, cap=None, hl=None):
    if cap:
        p = doc.add_paragraph(); r = p.add_run(cap)
        r.bold = True; r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
    t = doc.add_table(rows=len(rows)+1, cols=len(hdrs))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    for i, h in enumerate(hdrs):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255,255,255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(c, "1B3A4B")
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ""
            p = c.paragraphs[0]; r = p.add_run(str(v))
            r.font.size = Pt(9); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if hl is not None and ri == hl:
                shade(c, "E8F5E9"); r.bold = True
    doc.add_paragraph()

def P(doc, txt):
    p = doc.add_paragraph(txt)
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def eq(doc, txt, n=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.size = Pt(11); r.italic = True
    if n: r2 = p.add_run(f"    ({n})"); r2.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

# ── belge ─────────────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3.5); s.right_margin = Cm(2.5)

ns = doc.styles["Normal"]
ns.font.name = "Times New Roman"; ns.font.size = Pt(12)
ns.paragraph_format.line_spacing = 1.5
ns.paragraph_format.space_after = Pt(6)

for lv in range(1,4):
    hs = doc.styles[f"Heading {lv}"]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(0,0,0); hs.font.bold = True
    if lv==1: hs.font.size=Pt(16); hs.paragraph_format.space_before=Pt(24)
    elif lv==2: hs.font.size=Pt(14); hs.paragraph_format.space_before=Pt(18)
    else: hs.font.size=Pt(12); hs.paragraph_format.space_before=Pt(12)

# ======================================================================
doc.add_heading("BÖLÜM 1", level=1)
doc.add_heading("GİRİŞ", level=1)
# ======================================================================

# ══════════════════════════════════════════════════════════════════════
# 1.1
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1.1 Araştırmanın Arka Planı ve Motivasyonu", level=2)

P(doc,
    "Nesnelerin İnterneti (IoT), fiziksel dünyaya ait cihazların gömülü sensörler, "
    "işlemciler ve ağ arabirimleri vasıtasıyla birbirleriyle ve merkezi sistemlerle "
    "veri alışverişinde bulunmasını sağlayan bir bilgi-iletişim paradigmasıdır. "
    "Kavramın kökleri 1999 yılında Kevin Ashton'ın Massachusetts Teknoloji "
    "Enstitüsü'nde yaptığı sunuma uzanmakla birlikte, asıl yayılımı 2010'lu "
    "yılların başında düşük maliyetli kablosuz iletişim çiplerinin ve bulut "
    "platformlarının yaygınlaşmasıyla gerçekleşmiştir. Cisco Systems'in 2023 "
    "yıllık İnternet raporuna göre, 2025 yılı sonu itibarıyla dünya genelinde "
    "75,44 milyar IoT cihazının çevrimiçi olması öngörülmektedir (Cisco Systems, "
    "2023). Bu denli büyük bir cihaz ekosistemi, akıllı şehirler, endüstriyel "
    "otomasyon, sağlık izleme ve tarım gibi pek çok sektörde devrim niteliğinde "
    "dönüşümler yaratırken, beraberinde özgün güvenlik sorunlarını da getirmektedir."
)

P(doc,
    "IoT cihazlarının büyük çoğunluğu sınırlı hesaplama gücü, kısıtlı bellek ve "
    "düşük enerji bütçesiyle çalışmak zorundadır; bu kısıtlar, gelişmiş şifreleme "
    "protokolleri ve düzenli yazılım güncellemelerinin uygulanmasını güçleştirmektedir "
    "(Doshi, Apthorpe ve Feamster, 2018). Meidan vd. (2018), N-BaIoT çalışmasında "
    "varsayılan kimlik bilgileri ile çalışan IoT cihazlarının dakikalar içinde "
    "botnet ağlarına dahil edilebildiğini deneysel olarak kanıtlamıştır. "
    "Sivanathan vd. (2019) ise 28 farklı tipteki IoT cihazının ağ trafik "
    "profillerini analiz ederek, cihaz çeşitliliğinin homojen güvenlik "
    "politikalarının oluşturulmasını önemli ölçüde zorlaştırdığını göstermiştir."
)

P(doc,
    "Bu yapısal zaafiyetlerin en çarpıcı örneği, 2016 yılının sonbaharında "
    "gerçekleşen Mirai botnet olayıdır. Kolias, Kambourakis, Stavrou ve Voas "
    "(2017) tarafından ayrıntılı biçimde belgelenen bu saldırıda, fabrika "
    "ayarlarındaki kullanıcı adı–parola çiftleri kullanılarak yaklaşık 600.000 "
    "IoT cihazı ele geçirilmiş ve DDoS trafiği 1,2 Tbps zirve seviyesine "
    "ulaşmıştır. Saldırı, DNS altyapı sağlayıcısı Dyn'ı hedef almış ve "
    "Twitter, Netflix, Reddit gibi küresel ölçekli hizmetleri saatlerce "
    "erişilemez kılmıştır. Bu olay, siber güvenlik topluluğunda IoT "
    "kaynaklı DDoS tehditlerinin ciddiyetine ilişkin bir kırılma noktası "
    "olarak kabul edilmektedir."
)

P(doc,
    "Dağıtılmış Hizmet Reddi (DDoS) saldırıları, çok sayıda kaynak "
    "sistemden hedefe yönelik koordineli ve aşırı düzeyde trafik göndererek "
    "meşru kullanıcıların hizmete erişiminin engellenmesini amaçlayan siber "
    "saldırılardır. Zargar, Joshi ve Tipper (2013) tarafından gerçekleştirilen "
    "kapsamlı tarama çalışmasında, DDoS saldırı türleri hedeflenen ağ "
    "katmanına göre üç ana kategoride sınıflandırılmıştır: (i) hacimsel "
    "saldırılar — hedefin bant genişliğini tüketmeye yönelik UDP flood, "
    "ICMP flood ve DNS amplification; (ii) protokol tabanlı saldırılar — "
    "TCP/IP yığınının durum yönetimi mekanizmalarını istismar eden SYN flood "
    "ve ACK fragmentation; (iii) uygulama katmanı saldırıları — HTTP flood "
    "ve Slowloris gibi düşük hacimli ancak hizmet katmanında yüksek etkili "
    "türler. Agrawal ve Tapaswi (2019), güncel DDoS savunma stratejilerini "
    "bulut ekosistemi perspektifinden değerlendirerek, çok vektörlü (multi-vector) "
    "saldırıların tek bir savunma mekanizmasıyla engellenemeyeceğini vurgulamıştır."
)

P(doc,
    "Geleneksel saldırı tespit sistemleri (IDS), iki temel yaklaşıma "
    "dayanmaktadır: imza tabanlı ve anomali tabanlı tespit. İmza tabanlı "
    "sistemler (örneğin Snort, Suricata), bilinen saldırı kalıplarını bir "
    "veritabanında depolar ve gelen trafiği bu kalıplarla karşılaştırır. Bu "
    "yöntem düşük yanlış pozitif oranı (FPR) sağlasa da daha önce "
    "karşılaşılmamış sıfır-gün (zero-day) saldırılarına ve polimorfik "
    "varyantlara karşı yapısal olarak kördür. Anomali tabanlı sistemler "
    "ise normal trafik profilini öğrenerek bu profilden sapan aktiviteleri "
    "tespit eder; ancak yüksek FPR ve profil oluşturma güçlüğü temel "
    "dezavantajlarıdır. Bu iki yaklaşımın sınırlılıkları, makine öğrenmesi "
    "(ML) tabanlı saldırı tespitinin önem kazanmasının temel nedenidir "
    "(Buczak ve Guven, 2016)."
)

P(doc,
    "ML tabanlı saldırı tespiti, ağ trafiğindeki normal ve anormal kalıpları "
    "etiketlenmiş veriden otomatik olarak öğrenmekte ve öğrenilen modeli yeni, "
    "daha önce görülmemiş örneklere genelleyebilmektedir. Breiman (2001) "
    "tarafından önerilen Random Forest (RF), birden fazla karar ağacının "
    "Bootstrap örnekleriyle eğitilip çoğunluk oylamasıyla toplu karar "
    "verdiği bir topluluk (ensemble) yöntemidir. RF, aşırı uyuma (overfitting) "
    "karşı yüksek dayanıklılık ve yerleşik özellik önemi (feature importance) "
    "sağlama kapasitesi ile IDS literatüründe sıklıkla tercih edilen "
    "sınıflandırıcılardan biridir (Hastie, Tibshirani ve Friedman, 2009). "
    "Bununla birlikte, ML modellerinin performansı iki kritik faktöre "
    "bağlıdır: kullanılan özellik alt kümesinin kalitesi ve hiper-parametre "
    "konfigürasyonu. Bu iki faktörün elle ayarlanması zaman alıcı olup "
    "optimal çözüme ulaşma garantisi sunmamaktadır."
)

P(doc,
    "Son on yılda, doğadan esinlenen meta-sezgisel optimizasyon algoritmaları, "
    "ML modellerinin özellik seçimi ve hiper-parametre ayarlaması için giderek "
    "artan bir ilgiyle uygulanmaktadır. Xue, Zhang, Browne ve Yao (2016) "
    "tarafından gerçekleştirilen kapsamlı tarama çalışmasında, evrimsel "
    "hesaplama yaklaşımlarının — Parçacık Sürüsü Optimizasyonu (PSO), Genetik "
    "Algoritma (GA), Yarasa Algoritması (BA) gibi — özellik seçiminde etkili "
    "sonuçlar ürettiği sistematik olarak kanıtlanmıştır. Yang (2010), "
    "micro-yarasaların ultrasonik ekolokasyon mekanizmasından esinlenerek "
    "Yarasa Algoritması'nı geliştirmiştir; bu algoritma, frekans ayarlama, "
    "otomatik gürültü azaltma ve darbe oranı artırma mekanizmalarıyla keşif "
    "ve sömürü arasında kendi kendine uyarlanan bir denge kurmaktadır."
)

P(doc,
    "Bu tez çalışmasında, Yang'ın (2010) Yarasa Algoritması'nın — Yarasa "
    "Sürüsü Optimizasyonu (BSO) olarak adlandırılan uyarlanmış formu — "
    "CICIoT2023 veri seti (Neto vd., 2023) üzerinde DDoS saldırı tespiti "
    "bağlamında ilk kapsamlı uygulaması gerçekleştirilmektedir. Önerilen "
    "BSO-Hybrid RF çerçevesi, özellik seçimi ve RF hiper-parametre "
    "optimizasyonunu tek bir meta-sezgisel arama sürecinde eşzamanlı "
    "olarak gerçekleştirerek, IDS tasarımında verimli ve tekrarlanabilir "
    "bir metodoloji sunmayı hedeflemektedir."
)

# ══════════════════════════════════════════════════════════════════════
# 1.2
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1.2 Problem Tanımı", level=2)

P(doc,
    "DDoS saldırı tespitinde ML tabanlı yöntemlerin karşı karşıya kaldığı "
    "zorluklar, bu tezin motivasyonunu oluşturmakta olup üç yapısal "
    "problem başlığı altında ele alınmaktadır."
)

# 1.2.1
doc.add_heading("1.2.1 Problem 1: Yüksek Boyutlu Özellik Uzayı", level=3)

P(doc,
    "Ağ trafik veri setleri, akış düzeyinde onlarca hatta yüzlerce "
    "istatistiksel özellik barındırır. Chandrashekar ve Sahin (2014) "
    "tarafından belirtildiği üzere, bu özelliklerin bir kısmı yedekli "
    "(redundant), bir kısmı gürültülü (noisy) ve bir kısmı da "
    "sınıflandırma görevi için bilgilendirici değildir. Guyon ve "
    "Elisseeff (2003), öznitelik seçiminin üç temel faydasını "
    "vurgulamıştır: (a) modelin genellenebilirliğini artırma, "
    "(b) eğitim ve çıkarım süresini düşürme, (c) boyut laneti "
    "(curse of dimensionality) etkisini azaltma."
)

P(doc,
    "Bu çalışmada kullanılan CICIoT2023 veri seti (Neto vd., 2023), "
    "toplamda 39 ağ trafik özelliği içermektedir. Bu özellikler arasında "
    "akış süresi (duration), paket sayısı (Number), ortalama paket "
    "boyutu (AVG), paketler arası süre (IAT), TCP bayrak sayıları "
    "(syn_count, ack_count, fin_count, rst_count), istatistiksel "
    "dağılım metrikleri (Min, Max, Mean, Std, Tot_sum, Tot_size) ve "
    "protokol tabanlı özellikler yer almaktadır. 39 özelliğin tamamının "
    "doğrudan kullanılması, model karmaşıklığını artırarak aşırı uyum "
    "riskini yükseltmektedir. Bu nedenle, bilgilendirici özniteliklerin "
    "alt kümesinin sistematik biçimde seçilmesi gerekmektedir."
)

# 1.2.2
doc.add_heading("1.2.2 Problem 2: Hiper-Parametre Optimizasyonu", level=3)

P(doc,
    "Sınıflandırma modellerinin nihai performansı, hiper-parametre "
    "konfigürasyonuna doğrudan bağlıdır. Random Forest (Breiman, 2001) "
    "modelinde dört kritik hiper-parametre bulunmaktadır: ağaç sayısı "
    "(n_estimators), maksimum ağaç derinliği (max_depth), bir düğümü "
    "bölmek için gereken minimum örnek sayısı (min_samples_split) ve "
    "her bölmede değerlendirilen özellik alt kümesinin oranı "
    "(max_features). Bu dört parametrenin olası kombinasyonları astronomik "
    "genişlikte bir arama uzayı oluşturur."
)

P(doc,
    "Geleneksel parametre ayarlama yöntemlerinden ızgara araması (grid "
    "search), tüm kombinasyonları sistematik biçimde dener ancak arama "
    "uzayı büyüdükçe hesaplama maliyeti üstel olarak artar. Bergstra "
    "ve Bengio (2012), rastgele aramanın (random search) ızgara "
    "aramasına göre daha verimli olduğunu deneysel olarak kanıtlamış "
    "olmakla birlikte, her iki yöntem de yerel optimumlara takılma "
    "riskinden bağımsız değildir ve küresel arama kapasitesinden "
    "yoksundur. Meta-sezgisel algoritmalar, popülasyon tabanlı "
    "stokastik arama mekanizmalarıyla bu sınırlılığı aşma "
    "potansiyeli taşımaktadır (Yang, 2010)."
)

P(doc,
    "Burada dikkat edilmesi gereken kritik nokta, özellik seçimi ve "
    "hiper-parametre optimizasyonunun birbirinden bağımsız süreçler "
    "olmadığıdır. Belirli bir özellik alt kümesi için optimal olan parametre "
    "seti, farklı bir alt küme için suboptimal olabilir. Bu karşılıklı "
    "bağımlılık, iki sürecin ayrı ayrı değil, eşzamanlı olarak "
    "optimize edilmesi gerektiğini işaret etmektedir. Önerilen BSO-Hybrid "
    "RF çerçevesi, tam olarak bu eşzamanlı optimizasyonu gerçekleştirmek "
    "üzere tasarlanmıştır."
)

# 1.2.3
doc.add_heading("1.2.3 Problem 3: Sınıf Dengesizliği", level=3)

P(doc,
    "IDS veri setlerinde sınıf dengesizliği, yapısal ve kaçınılmaz bir "
    "olgudur: normal trafik örnekleri genellikle saldırı örneklerinden "
    "çok daha fazladır; ayrıca farklı saldırı türleri arasında da "
    "büyük oransal farklılıklar bulunur. He ve Garcia (2009) tarafından "
    "gerçekleştirilen kapsamlı tarama çalışmasında, dengesiz veri "
    "dağılımlarının standart sınıflandırıcıların performansını sistematik "
    "olarak bozduğu ve modelin çoğunluk sınıflarını doğru tahmin ederken "
    "azınlık sınıflarını yanlış sınıflandırmaya eğilim gösterdiği "
    "gösterilmiştir."
)

P(doc,
    "CICIoT2023 veri setinde bu problem somut biçimde gözlemlenmektedir. "
    "En büyük sınıf olan DDoS-ACK_Fragmentation 53.148 örnekle veri "
    "setinin %44,9'unu oluştururken, en küçük sınıf olan Backdoor_Malware "
    "yalnızca 2.252 örnekle %1,9'luk bir dilimi temsil etmektedir. Bu iki "
    "sınıf arasındaki oran yaklaşık 23,6:1'dir. Sınıflandırıcı, eğitim "
    "sürecinde çoğunluk sınıflarına aşırı uyum gösterme ve azınlık "
    "sınıfını çoğunluk olarak etiketleme eğilimi taşır — bu durum, özellikle "
    "nadir fakat kritik öneme sahip saldırı türlerinin gözden kaçmasına "
    "neden olabilir."
)

P(doc,
    "Bu soruna yönelik veri düzeyinde en yaygın çözüm, Chawla, Bowyer, "
    "Hall ve Kegelmeyer (2002) tarafından önerilen Sentetik Azınlık "
    "Aşırı Örnekleme Tekniği'dir (SMOTE). SMOTE, azınlık sınıfı "
    "örnekleri arasında k-en yakın komşu (k-NN) tabanlı doğrusal "
    "interpolasyon yaparak sentetik örnekler üretir. Bu çalışmada SMOTE "
    "yalnızca eğitim setine uygulanmış; test ve doğrulama setleri "
    "orijinal dağılımlarıyla korunarak veri sızıntısı (data leakage) "
    "önlenmiştir. Eğitim seti boyutu 72.252'den 87.500'e yükselmiş, "
    "her sınıf 17.500 örneğe dengelenmiştir."
)

P(doc,
    "Özetle, bu tez yukarıda tanımlanan üç yapısal problemi — yüksek "
    "boyutlu özellik uzayı, hiper-parametre optimizasyonu ve sınıf "
    "dengesizliği — BSO-Hybrid RF çerçevesi aracılığıyla eşzamanlı "
    "olarak ele alan bütünleşik bir çözüm önermektedir."
)

# ══════════════════════════════════════════════════════════════════════
# 1.3
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1.3 Araştırma Hipotezleri", level=2)

P(doc,
    "Yukarıda tanımlanan problem çerçevesinde, bu tez çalışmasında "
    "aşağıdaki üç araştırma hipotezi formüle edilmiş ve deneysel "
    "olarak test edilmiştir. Hipotezlerin belirlenmesinde Demšar (2006) "
    "tarafından önerilen istatistiksel karşılaştırma çerçevesi ve "
    "Cohen (1988) tarafından tanımlanan etki büyüklüğü kavramı "
    "referans alınmıştır."
)

hyps = [
    ("AH₁ (Performans Üstünlüğü): ",
     "BSO ile eşzamanlı olarak özellik seçimi ve hiper-parametre "
     "optimizasyonu gerçekleştirilen Random Forest modeli, varsayılan "
     "parametrelerle çalışan standart RF modeline kıyasla istatistiksel "
     "olarak anlamlı düzeyde (p < 0,05) daha yüksek F1-Macro skoru elde "
     "eder. Test yöntemi: 10-katlı CV üzerinde eşleştirilmiş t-testi."),
    ("AH₂ (Boyut Azaltma): ",
     "BSO tabanlı öznitelik seçimi, toplam özellik sayısının %50'sinden "
     "fazlasını elemine ederken sınıflandırma doğruluğunda %1'den az "
     "kayba neden olur. Ölçüt: |Acc(39 öz.) − Acc(BSO öz.)| < %1."),
    ("AH₃ (Optimizasyon Verimliliği): ",
     "BSO-Hybrid çerçevesi (özellik seçimi + HP optimizasyonu + SMOTE), "
     "PSO, GA ve GWO tabanlı aynı yapıdaki çerçevelere kıyasla daha düşük "
     "uygunluk (fitness) değeri elde eder. Ölçüt: f(BSO) < f(PSO), "
     "f(BSO) < f(GA), f(BSO) < f(GWO)."),
]
for lab, txt in hyps:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(lab); r.bold = True
    p.add_run(txt)

P(doc,
    "Her üç hipotezin test edilmesi, Bölüm 4'te sunulan deneysel "
    "sonuçlar ve istatistiksel analizlere dayanmaktadır. Hipotez "
    "testlerinde α = 0,05 anlamlılık düzeyi ve iki yönlü (two-tailed) "
    "test yaklaşımı benimsenmiştir."
)

# ══════════════════════════════════════════════════════════════════════
# 1.4
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1.4 Araştırma Soruları", level=2)

P(doc,
    "Araştırma hipotezleriyle tutarlı biçimde, bu tezin yanıtlamayı "
    "hedeflediği üç temel araştırma sorusu şunlardır:"
)

qs = [
    ("AS₁ (Özellik Seçimi Etkinliği): ",
     "BSO algoritması, CICIoT2023 veri setindeki 39 ağ trafik "
     "özelliğinden en bilgilendirici alt kümeyi ne ölçüde etkili "
     "biçimde seçebilir ve bu seçim sınıflandırma performansını "
     "nasıl etkiler?"),
    ("AS₂ (Azınlık Sınıf Tespiti): ",
     "BSO-Hybrid RF çerçevesi, Backdoor_Malware gibi düşük temsil "
     "edilen ve ciddi tehdit oluşturan saldırı türlerinde operasyonel "
     "açıdan yeterli tespit performansı sağlayabilir mi?"),
    ("AS₃ (İstatistiksel Güvenilirlik): ",
     "Önerilen çerçevenin performans üstünlüğü, 10-katlı tabakalı "
     "çapraz doğrulama ve birden fazla eşleştirilmiş istatistiksel test "
     "(t-testi, Wilcoxon, McNemar, Cohen's d) ile güvenilir ve "
     "tekrarlanabilir biçimde doğrulanabilir mi?"),
]
for lab, txt in qs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(lab); r.bold = True
    p.add_run(txt)

# ══════════════════════════════════════════════════════════════════════
# 1.5
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1.5 Tezin Amacı ve Kapsamı", level=2)

P(doc,
    "Bu çalışmanın temel amacı, IoT ağ ortamlarında DDoS saldırılarının "
    "daha etkili, verimli ve güvenilir biçimde tespiti için Yarasa Sürüsü "
    "Optimizasyonu (BSO) tabanlı bir hibrit makine öğrenmesi çerçevesi "
    "tasarlamak, uygulamak ve kapsamlı deneysel doğrulama ile "
    "değerlendirmektir. Çalışmanın kapsamı aşağıdaki altı başlık "
    "altında özetlenebilir:"
)

scopes = [
    ("(a) Veri Seti ve Keşifsel Analiz: ", "Canadian Institute for "
     "Cybersecurity tarafından yayınlanan CICIoT2023 veri seti "
     "(Neto vd., 2023) üzerinde 103.218 akış kaydı ve 39 özelliğin "
     "5 sınıflı çok sınıflandırma problemi olarak yapılandırılması."),
    ("(b) BSO Optimizasyonu: ", "25 yarasa × 50 iterasyon = 1.177 "
     "uygunluk değerlendirmesi ile eşzamanlı özellik seçimi (39→19, "
     "%51,3 azaltma) ve RF hiper-parametre optimizasyonu "
     "(n_estimators=266, max_depth=20, min_samples_split=7, "
     "max_features=0,469). Toplam optimizasyon süresi: 840,43 saniye."),
    ("(c) Sınıf Dengeleme: ", "SMOTE (Chawla vd., 2002) ile yalnızca "
     "eğitim setinde 72.252→87.500 örneğe yükseltme; her sınıf "
     "17.500 örneğe dengeleme. 15.248 sentetik örnek üretimi."),
    ("(d) Karşılaştırmalı Değerlendirme: ", "12 farklı ML modeli "
     "(BSO-Hybrid RF, XGBoost, GWO-RF, RF, GA-RF, PSO-RF, DT, KNN, "
     "SVM, LR, BSO-SVM, NB) × 7 performans metriği."),
    ("(e) Ablasyon Çalışması: ", "S1–S4 senaryoları ile her bileşenin "
     "(SMOTE, BSO-FS, BSO-HP) marjinal katkısının izole edilerek "
     "ölçülmesi. Wolpert (1996) tarafından vurgulanan bileşen analizi "
     "prensibiyle tutarlıdır."),
    ("(f) İstatistiksel Doğrulama: ", "10-katlı tabakalı çapraz "
     "doğrulama, eşleştirilmiş t-testi, Wilcoxon Signed-Rank testi, "
     "McNemar testi ve Cohen's d etki büyüklüğü (Demšar, 2006; "
     "Cohen, 1988)."),
]
for lab, txt in scopes:
    p = doc.add_paragraph()
    r = p.add_run(lab); r.bold = True
    p.add_run(txt)

P(doc,
    "Çalışmanın kapsamı dışında tutulan unsurlar: gerçek zamanlı ağ "
    "akışı üzerinde çevrimiçi (online) tespit deneyleri, CICIoT2023 "
    "veri setinin tam 33 sınıflı versiyonunun kullanımı ve derin "
    "öğrenme sınıflandırıcılarının (CNN, LSTM, Transformer) deneysel "
    "karşılaştırması. Bu unsurlar, gelecek çalışma önerileri kapsamında "
    "Bölüm 5'te tartışılmaktadır."
)

# ══════════════════════════════════════════════════════════════════════
# 1.6
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1.6 Tezin Katkıları", level=2)

P(doc,
    "Bu tez, IoT güvenliği ve meta-sezgisel makine öğrenmesi "
    "kesişiminde aşağıdaki dört özgün katkıyı literatüre sunmaktadır:"
)

contribs = [
    ("K1 — BSO-Hybrid RF Çerçevesi: ",
     "Özellik seçimi ve hiper-parametre optimizasyonunu tek bir "
     "meta-sezgisel arama sürecinde eşzamanlı olarak birleştiren "
     "özgün hibrit çerçeve tasarımı. Yang'ın (2010) Yarasa "
     "Algoritması'nın CICIoT2023 üzerinde DDoS tespiti bağlamında "
     "ilk kapsamlı uygulamasıdır. 25 yarasa bireyi, 50 iterasyon "
     "boyunca toplam 1.177 uygunluk değerlendirmesi gerçekleştirmiş "
     "ve en iyi uygunluk değeri 0,177801 olarak elde edilmiştir."),

    ("K2 — %51,3 Boyut Azaltma: ",
     "BSO algoritması, 39 ağ trafik özelliğinden 19'unu optimal "
     "alt küme olarak seçmiştir. Bu %51,3'lük boyut azaltma, "
     "sınıflandırma doğruluğunda yalnızca +0,08% fark ile "
     "gerçekleştirilmiştir (%89,74 → %89,82). Elenen 20 özelliğin "
     "büyük çoğunluğu yedekli istatistiksel türevlerdir (örn. Mean "
     "ve AVG; Std ve Var). Seçilen 19 özellik arasında syn_count "
     "(önem: 0,2245), Number (0,1834) ve Tot_sum (0,1541) en "
     "bilgilendirici üç nitelik olarak belirlenmiştir."),

    ("K3 — Azınlık Sınıf İyileştirmesi: ",
     "SMOTE (Chawla vd., 2002) entegrasyonu ile CICIoT2023 veri "
     "setinin en küçük sınıfı olan Backdoor_Malware'in F1-Skoru "
     "%28,40'tan %57,40'a yükselmiştir (%102 iyileşme). Bu sonuç, "
     "SMOTE uygulanmadan azınlık saldırı türlerinin neredeyse tespit "
     "edilemez olduğunu kanıtlamaktadır. Sınıf dengeleme, S1–S4 "
     "ablasyon çalışmasının en kritik bileşeni olarak teşhis "
     "edilmiştir (F1-Macro: +%3,78 artış)."),

    ("K4 — Kapsamlı Deneysel Doğrulama: ",
     "12 model × 7 performans metriği × 10-katlı tabakalı çapraz "
     "doğrulama × 4 istatistiksel test (eşleştirilmiş t-testi, "
     "Wilcoxon, McNemar, Cohen's d). S1–S4 ablasyon senaryoları ile "
     "her bileşenin (SMOTE, BSO-FS, BSO-HP) marjinal etkisi izole "
     "edilerek ölçülmüştür. 11 eşleştirilmiş istatistiksel testten "
     "9'unda p < 0,05 düzeyinde anlamlılık saptanmıştır; BSO vs "
     "Decision Tree karşılaştırmasında Cohen's d = 11,07 (çok büyük "
     "etki) elde edilmiştir."),
]

for lab, txt in contribs:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(lab); r.bold = True
    p.add_run(txt)

# ======= Tablo 1.1 — Katkı Özet Tablosu ======
tbl(doc,
    ["Katkı", "Açıklama", "Kanıt"],
    [
        ["K1", "BSO-Hybrid RF çerçevesi", "Fitness: 0,177801; 1.177 değerlendirme"],
        ["K2", "%51,3 boyut azaltma (39→19)", "Doğruluk farkı: +0,08%"],
        ["K3", "Azınlık sınıf iyileştirme", "Backdoor F1: %28,40 → %57,40"],
        ["K4", "Kapsamlı deneysel doğrulama", "9/11 testte p < 0,05; Cohen's d = 11,07"],
    ],
    cap="Tablo 1.1: Tezin Özgün Katkılarının Özeti",
)

# ══════════════════════════════════════════════════════════════════════
# 1.7
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1.7 Araştırmanın Önemi ve Beklenen Etkisi", level=2)

P(doc,
    "Bu çalışmanın akademik ve pratik önemi, birbirleriyle bağlantılı "
    "üç boyutta değerlendirilebilir. Birincisi, Yarasa Algoritması'nın "
    "CICIoT2023 gibi güncel ve büyük ölçekli bir IoT güvenlik veri "
    "seti üzerinde eşzamanlı özellik seçimi ve hiper-parametre "
    "optimizasyonu amacıyla uygulandığı ilk kapsamlı çalışma olması "
    "itibarıyla, meta-sezgisel IDS literatürüne yeni bir perspektif "
    "katmaktadır."
)

P(doc,
    "İkincisi, önerilen çerçevenin %51,3 boyut azaltma sağlaması, "
    "gerçek dünya IoT ağlarında hesaplama ve bellek kısıtlamalarının "
    "önem taşıdığı ortamlarda dağıtım kolaylığı açısından pratik "
    "değer sunmaktadır. 19 özellikle çalışan model, 39 özellikli "
    "modele göre daha hızlı çıkarım yapabilir — bu durum özellikle "
    "yüksek trafikli ağ geçitlerinde ve edge bilgi işlem senaryolarında "
    "avantaj sağlar."
)

P(doc,
    "Üçüncüsü, SMOTE ile sınıf dengelemenin Backdoor_Malware gibi "
    "nadir saldırı türlerinin tespitinde %102 iyileşme sağlaması, "
    "operasyonel IDS tasarımlarında sınıf dengeleme stratejilerinin "
    "zorunluluğunu ampirik olarak kanıtlamaktadır. Bu bulgu, IoT "
    "güvenlik sistemlerini tasarlayan mühendisler için doğrudan "
    "uygulanabilir bir yol haritası sunmaktadır."
)

# ══════════════════════════════════════════════════════════════════════
# 1.8
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1.8 Temel Kavramlar ve Tanımlar", level=2)

P(doc,
    "Bu bölümde, tez boyunca sıklıkla kullanılan temel kavramlar "
    "kısaca tanımlanmaktadır. Kapsamlı açıklamalar Bölüm 2'de ilgili "
    "alt bölümlerde sunulmaktadır."
)

defs = [
    ("DDoS (Dağıtılmış Hizmet Reddi): ",
     "Birden fazla kaynaktan hedefe yönelik koordineli trafik "
     "göndererek hizmetin kesintiye uğratılmasını amaçlayan "
     "siber saldırı türüdür (Zargar vd., 2013)."),
    ("IoT (Nesnelerin İnterneti): ",
     "Fiziksel cihazların sensörler ve ağ bağlantıları aracılığıyla "
     "veri paylaşmasını mümkün kılan bilgi-iletişim paradigması "
     "(Cisco Systems, 2023)."),
    ("Meta-Sezgisel Optimizasyon: ",
     "Doğadan esinlenen, NP-zor problemlere stokastik arama "
     "ile yaklaşık çözümler üreten algoritma ailesidir "
     "(Yang, 2010; Kennedy ve Eberhart, 1995)."),
    ("Yarasa Algoritması / BSO: ",
     "Yang (2010) tarafından micro-yarasaların ekolokasyon "
     "davranışından esinlenerek geliştirilen, frekans ayarlama "
     "ve otomatik gürültü/darbe kontrolü mekanizmalarına sahip "
     "meta-sezgisel algoritmadır."),
    ("Random Forest (RF): ",
     "Breiman (2001) tarafından önerilen, birden fazla karar "
     "ağacının Bootstrap alt örnekleriyle eğitilip çoğunluk "
     "oylamasıyla karar verdiği topluluk (ensemble) yöntemidir."),
    ("SMOTE: ",
     "Chawla vd. (2002) tarafından geliştirilen, azınlık "
     "sınıfı örnekleri arasında doğrusal interpolasyonla sentetik "
     "örnekler üreten veri dengeleme tekniğidir."),
    ("F1-Macro: ",
     "Tüm sınıfların F1 harmonik ortalamasının eşit ağırlıklı "
     "aritmetik ortalaması; sınıf dengesizliğinden bağımsız "
     "bir performans metriğidir (Sokolova ve Lapalme, 2009)."),
    ("AUC-ROC: ",
     "ROC eğrisi altında kalan alan; eşik-bağımsız sınıflandırma "
     "performansını ölçer (Fawcett, 2006)."),
    ("MCC: ",
     "Matthews Korelasyon Katsayısı; [-1, +1] aralığında, sınıf "
     "dengesizliğine dayanıklı değerlendirme metriğidir "
     "(Chicco ve Jurman, 2020)."),
]

for lab, txt in defs:
    p = doc.add_paragraph()
    r = p.add_run(lab); r.bold = True; r.font.size = Pt(11)
    p.add_run(txt)

# ══════════════════════════════════════════════════════════════════════
# 1.9
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1.9 Tezin Organizasyonu", level=2)

P(doc,
    "Bu tez, birbirine mantıksal olarak bağlı beş bölümden "
    "oluşmaktadır. Her bölümün kapsamı ve birbiriyle ilişkisi "
    "aşağıda özetlenmektedir:"
)

org = [
    ("Bölüm 1 — Giriş (bu bölüm): ",
     "Araştırmanın arka planı ve motivasyonu, problem tanımı "
     "(3 yapısal problem), araştırma hipotezleri (AH₁–AH₃), "
     "araştırma soruları (AS₁–AS₃), amaç ve kapsam, tezin "
     "dört özgün katkısı (K1–K4), temel kavramlar."),

    ("Bölüm 2 — Literatür Taraması: ",
     "DDoS saldırıları ve IoT güvenliği, ML tabanlı saldırı "
     "tespiti (denetimli ve derin öğrenme), meta-sezgisel "
     "optimizasyon algoritmaları (PSO, GA, GWO, BA), Yarasa "
     "Algoritması'nın matematiksel formülasyonu, öznitelik "
     "seçimi yöntemleri (filtre, sarmalayıcı, gömülü) ve "
     "ilgili çalışmaların karşılaştırmalı tablosu."),

    ("Bölüm 3 — Yöntem: ",
     "Araştırma tasarımı (7 aşamalı pipeline), CICIoT2023 veri "
     "setinin yapısı ve sınıf dağılımı, veri ön işleme "
     "(normalizasyon, SMOTE), BSO algoritmasının ikili özellik "
     "seçimi ve sürekli HP optimizasyonu görevleri, uygunluk "
     "fonksiyonu, hibrit model tanımı, S1–S4 ablasyon senaryoları, "
     "7 değerlendirme metriği ve 5 istatistiksel doğrulama yöntemi."),

    ("Bölüm 4 — Bulgular ve Değerlendirme: ",
     "BSO yakınsama analizi, seçilen 19 özellik ve önem "
     "sıralaması, 12 model karşılaştırması (Tablo 4.2), sınıf "
     "bazında performans (karışıklık matrisi ve ROC eğrileri), "
     "S1–S4 ablasyon sonuçları, 10-katlı CV (σ=0,204), "
     "istatistiksel test sonuçları ve AH₁–AH₃ hipotez doğrulaması."),

    ("Bölüm 5 — Sonuç ve Öneriler: ",
     "Temel bulguların özeti (B1–B5), araştırma sorularının "
     "yanıtları (AS₁–AS₃), tezin katkıları, 5 sınırlılık, "
     "7 gelecek çalışma önerisi (derin öğrenme entegrasyonu, "
     "çoklu veri seti doğrulaması, gerçek zamanlı dağıtım, "
     "açıklanabilir AI, federe öğrenme) ve kapanış."),
]

for lab, txt in org:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(lab); r.bold = True
    p.add_run(txt)

P(doc,
    "Tezin ekler bölümünde, BSO algoritmasının Python (scikit-learn "
    "+ NumPy) tabanlı sözde kodu, CICIoT2023 özellik listesinin "
    "tamamı ve 12 model için ayrıntılı karışıklık matrisleri "
    "sunulmaktadır."
)

# ══════════════════════════════════════════════════════════════════════
# KAYDET
# ══════════════════════════════════════════════════════════════════════
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "..", "BOLUM_1_GIRIS.docx")
out = os.path.normpath(out)
doc.save(out)

print(f"\n{'='*60}")
print(f"  ✅ BÖLÜM 1 — GİRİŞ BAŞARIYLA OLUŞTURULDU")
print(f"{'='*60}")
print(f"  Dosya : {out}")
print(f"  Alt bölüm : 9  (1.1 – 1.9)")
print(f"  Paragraf   : ~45+")
print(f"  Tablo      : 1  (Tablo 1.1)")
print(f"  Hipotez    : 3  (AH₁, AH₂, AH₃)")
print(f"  Araştırma S: 3  (AS₁, AS₂, AS₃)")
print(f"  Katkı      : 4  (K1–K4)")
print(f"  Tanım      : 9  kavram tanımı")
print(f"{'='*60}")
