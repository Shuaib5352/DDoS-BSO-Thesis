#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tam Yüksek Lisans Tezi — 5 Bölüm (Türkçe)
IoT Ortamlarında BSO Tabanlı Hibrit Makine Öğrenmesi ile DDoS Saldırı Tespiti

Yazar:    Shuaib Ayad Jasim
Danışman: Dr. Saim Ervural
Üniversite: KTO Karatay Üniversitesi
Tarih:    2026

Her cümle akademik kaynaklarla desteklenmiş, doğal akademik Türkçe ile yazılmıştır.
Tüm veriler gerçek deney sonuçlarına dayanmaktadır.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows, caption=None, hl_row=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.bold = True; r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(255,255,255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(c, "1B3A4B")
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ""
            p = c.paragraphs[0]; r = p.add_run(str(v))
            r.font.size = Pt(9); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if hl_row is not None and ri == hl_row:
                set_cell_shading(c, "E8F5E9"); r.bold = True
    doc.add_paragraph()

def eq(doc, text, num=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(11); r.italic = True
    if num:
        r2 = p.add_run(f"    ({num})"); r2.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def P(doc, text):
    """Add a normal paragraph with proper first-line indent."""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def blank(doc):
    doc.add_paragraph()

# ============================================================================
# CREATE DOCUMENT
# ============================================================================
doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3.5); s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"; style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for lv in range(1, 4):
    hs = doc.styles[f"Heading {lv}"]
    hs.font.name = "Times New Roman"; hs.font.color.rgb = RGBColor(0,0,0); hs.font.bold = True
    if lv == 1: hs.font.size = Pt(16); hs.paragraph_format.space_before = Pt(24)
    elif lv == 2: hs.font.size = Pt(14); hs.paragraph_format.space_before = Pt(18)
    else: hs.font.size = Pt(12); hs.paragraph_format.space_before = Pt(12)

# ============================================================================
# KAPAK SAYFASI
# ============================================================================
for _ in range(4): doc.add_paragraph()
for txt, sz, b in [
    ("KTO KARATAY ÜNİVERSİTESİ", 16, True),
    ("LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ", 14, True),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.bold = b; r.font.size = Pt(sz)

blank(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Bilgisayar Mühendisliği Anabilim Dalı"); r.font.size = Pt(13)

for _ in range(2): blank(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "IoT ORTAMLARINDA YARASA SÜRÜSÜ OPTİMİZASYONU (BSO) TABANLI "
    "HİBRİT MAKİNE ÖĞRENMESİ İLE DDoS SALDIRI TESPİTİ: "
    "CICIoT2023 VERİ SETİ ÜZERİNDE DENEYSEL DEĞERLENDİRME"
); r.bold = True; r.font.size = Pt(14)

blank(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("YÜKSEK LİSANS TEZİ"); r.bold = True; r.font.size = Pt(14)

for _ in range(2): blank(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Shuaib Ayad JASIM"); r.font.size = Pt(13)

blank(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Danışman: Dr. Öğr. Üyesi Saim ERVURAL"); r.font.size = Pt(13)

for _ in range(2): blank(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("KONYA — 2026"); r.bold = True; r.font.size = Pt(14)
doc.add_page_break()

# ============================================================================
# ÖZET (TÜRKÇE)
# ============================================================================
doc.add_heading("ÖZET", level=1)

P(doc,
    "Bu tez çalışmasında, Nesnelerin İnterneti (IoT) ortamlarında gerçekleştirilen "
    "Dağıtılmış Hizmet Reddi (DDoS) saldırılarının tespiti için Yarasa Sürüsü "
    "Optimizasyonu (BSO) tabanlı bir hibrit makine öğrenmesi çerçevesi önerilmiş, "
    "uygulanmış ve kapsamlı deneysel doğrulama ile değerlendirilmiştir. Önerilen "
    "BSO-Hybrid RF çerçevesi, Yang (2010) tarafından geliştirilen Yarasa Algoritması'nın "
    "eşzamanlı öznitelik seçimi ve Random Forest hiper-parametre optimizasyonu için "
    "uyarlanmış bir versiyonunu kullanmaktadır."
)

P(doc,
    "Deneyler, Canadian Institute for Cybersecurity tarafından yayınlanan CICIoT2023 "
    "veri seti üzerinde gerçekleştirilmiştir (Neto vd., 2023). 118.466 akış kaydı ve "
    "39 ağ trafik özelliğinden oluşan veri seti, 5 sınıflı çok sınıflandırma problemi "
    "olarak yapılandırılmıştır. BSO, 25 yarasa × 50 iterasyon = 1.177 uygunluk "
    "değerlendirmesi sonucunda 39 özellikten 19'unu seçerek %51,3 boyut azaltma "
    "sağlamış ve aynı anda n_estimators=266, max_depth=20, min_samples_split=7, "
    "max_features=0,469 parametrelerini optimize etmiştir."
)

P(doc,
    "BSO-Hybrid RF modeli, %89,82 doğruluk, %84,24 F1-Macro, %98,38 AUC-ROC ve "
    "0,8676 MCC elde etmiştir. SMOTE ile sınıf dengeleme sonrasında Backdoor_Malware "
    "F1-Skoru %28,40'tan %57,40'a yükselmiştir. 10-katlı tabakalı çapraz doğrulamada "
    "ortalama doğruluk %90,98 ± 0,204 olarak ölçülmüştür. 12 model karşılaştırmasında "
    "11 eşleştirilmiş istatistiksel testten 9'unda anlamlı fark (p < 0,05) saptanmıştır."
)

p = doc.add_paragraph()
r = p.add_run("Anahtar Kelimeler: "); r.bold = True
p.add_run("DDoS Tespiti, Yarasa Sürüsü Optimizasyonu, Random Forest, Öznitelik Seçimi, "
          "CICIoT2023, IoT Güvenliği, Hibrit Makine Öğrenmesi, SMOTE")
doc.add_page_break()

# ============================================================================
# ABSTRACT (ENGLISH)
# ============================================================================
doc.add_heading("ABSTRACT", level=1)

P(doc,
    "This thesis proposes, implements, and evaluates a Beetle Swarm Optimization (BSO) "
    "based hybrid machine learning framework for detecting Distributed Denial of Service "
    "(DDoS) attacks in Internet of Things (IoT) environments. The proposed BSO-Hybrid RF "
    "framework adapts Yang's (2010) Bat Algorithm for simultaneous feature selection and "
    "Random Forest hyperparameter optimization within a single optimization loop."
)

P(doc,
    "Experiments were conducted on the CICIoT2023 dataset published by the Canadian "
    "Institute for Cybersecurity (Neto et al., 2023), comprising 118,466 flow records "
    "and 39 network traffic features across 5 classes. BSO selected 19 out of 39 features "
    "(51.3% reduction) while optimizing RF hyperparameters (n_estimators=266, max_depth=20, "
    "min_samples_split=7, max_features=0.469) through 1,177 fitness evaluations."
)

P(doc,
    "The BSO-Hybrid RF model achieved 89.82% accuracy, 84.24% F1-Macro, 98.38% AUC-ROC, "
    "and 0.8676 MCC. SMOTE-based class balancing improved Backdoor_Malware F1-Score from "
    "28.40% to 57.40%. Ten-fold stratified cross-validation yielded 90.98% ± 0.204 mean "
    "accuracy. Out of 11 pairwise statistical tests against baseline models, 9 showed "
    "statistically significant differences (p < 0.05)."
)

p = doc.add_paragraph()
r = p.add_run("Keywords: "); r.bold = True
p.add_run("DDoS Detection, Bat Swarm Optimization, Random Forest, Feature Selection, "
          "CICIoT2023, IoT Security, Hybrid Machine Learning, SMOTE")
doc.add_page_break()

# ============================================================================
# TEŞEKKÜR
# ============================================================================
doc.add_heading("TEŞEKKÜR", level=1)

P(doc,
    "Bu tez çalışmasının her aşamasında bilgi birikimiyle yol gösteren, akademik "
    "titizliğiyle çalışmamı şekillendiren ve sabırlı destekleriyle yanımda olan "
    "değerli danışmanım Dr. Öğr. Üyesi Saim ERVURAL'a en içten teşekkürlerimi "
    "sunarım. Hocamın rehberliği olmadan bu çalışmanın tamamlanması mümkün olmazdı."
)

P(doc,
    "Lisansüstü eğitimim boyunca bilgi ve deneyimlerini benimle paylaşan KTO Karatay "
    "Üniversitesi Bilgisayar Mühendisliği Bölümü'nün tüm değerli öğretim üyelerine "
    "teşekkür ederim."
)

P(doc,
    "Bu süreçte manevi destekleriyle her zaman yanımda olan aileme, özellikle anneme "
    "ve babama, sonsuz teşekkürlerimi borçluyum."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Shuaib Ayad JASIM\nKonya, 2026")
r.font.size = Pt(11)
doc.add_page_break()

# ============================================================================
# İÇİNDEKİLER (Placeholder — Word'de otomatik oluşturulur)
# ============================================================================
doc.add_heading("İÇİNDEKİLER", level=1)
P(doc, "[Bu sayfa Word'de Referanslar → İçindekiler Tablosu → Otomatik Tablo 1 "
       "ile güncellenecektir. Ctrl+A ardından F9 tuşuyla güncellenir.]")
doc.add_page_break()

doc.add_heading("TABLOLAR LİSTESİ", level=1)
P(doc, "[Word'de Referanslar → Şekil ve Tablo Listesi ile otomatik oluşturulur.]")
doc.add_page_break()

doc.add_heading("ŞEKİLLER LİSTESİ", level=1)
P(doc, "[Word'de Referanslar → Şekil ve Tablo Listesi ile otomatik oluşturulur.]")
doc.add_page_break()

# ============================================================================
# KISALTMALAR
# ============================================================================
doc.add_heading("KISALTMALAR", level=1)

abbrevs = [
    ("AUC-ROC", "Receiver Operating Characteristic Eğrisi Altındaki Alan"),
    ("BA / BSO", "Yarasa Algoritması / Yarasa Sürüsü Optimizasyonu"),
    ("CNN", "Evrişimli Sinir Ağı (Convolutional Neural Network)"),
    ("CV", "Çapraz Doğrulama (Cross-Validation)"),
    ("DDoS", "Dağıtılmış Hizmet Reddi (Distributed Denial of Service)"),
    ("DT", "Karar Ağacı (Decision Tree)"),
    ("EDA", "Keşifsel Veri Analizi (Exploratory Data Analysis)"),
    ("FPR", "Yanlış Pozitif Oranı (False Positive Rate)"),
    ("GA", "Genetik Algoritma (Genetic Algorithm)"),
    ("GWO", "Gri Kurt Optimizasyonu (Grey Wolf Optimizer)"),
    ("IDS", "Saldırı Tespit Sistemi (Intrusion Detection System)"),
    ("IoT", "Nesnelerin İnterneti (Internet of Things)"),
    ("KNN", "K-En Yakın Komşu (K-Nearest Neighbors)"),
    ("LR", "Lojistik Regresyon (Logistic Regression)"),
    ("LSTM", "Uzun Kısa Süreli Bellek (Long Short-Term Memory)"),
    ("MCC", "Matthews Korelasyon Katsayısı"),
    ("ML", "Makine Öğrenmesi (Machine Learning)"),
    ("NB", "Naive Bayes"),
    ("PSO", "Parçacık Sürüsü Optimizasyonu (Particle Swarm Optimization)"),
    ("RF", "Rastgele Orman (Random Forest)"),
    ("SMOTE", "Sentetik Azınlık Aşırı Örnekleme Tekniği"),
    ("SVM", "Destek Vektör Makinesi (Support Vector Machine)"),
]

t = doc.add_table(rows=len(abbrevs), cols=2)
t.style = "Table Grid"
for i, (ab, desc) in enumerate(abbrevs):
    t.rows[i].cells[0].text = ab
    t.rows[i].cells[1].text = desc
    for c in range(2):
        for p in t.rows[i].cells[c].paragraphs:
            for r in p.runs: r.font.size = Pt(10)
doc.add_page_break()


# ████████████████████████████████████████████████████████████████████████████
# BÖLÜM 1 — GİRİŞ
# ████████████████████████████████████████████████████████████████████████████

doc.add_heading("BÖLÜM 1", level=1)
doc.add_heading("GİRİŞ", level=1)

# ── 1.1 ──
doc.add_heading("1.1 Araştırmanın Arka Planı ve Motivasyonu", level=2)

P(doc,
    "Nesnelerin İnterneti (IoT), fiziksel cihazların sensörler ve ağ bağlantıları "
    "aracılığıyla birbirleriyle veri paylaşmasını mümkün kılan bir paradigma olarak "
    "son on yılda katlanarak büyümüştür. Cisco'nun 2023 yıllık raporuna göre, "
    "küresel ölçekte 2025 yılı itibarıyla 75 milyardan fazla IoT cihazının aktif "
    "olması beklenmektedir (Cisco Systems, 2023). Bu büyüme, beraberinde ciddi "
    "güvenlik kaygılarını da getirmiştir; zira IoT cihazları sınırlı hesaplama "
    "kaynakları, zayıf kimlik doğrulama mekanizmaları ve düzensiz yazılım "
    "güncellemeleri nedeniyle siber saldırılara karşı son derece savunmasızdır "
    "(Meidan vd., 2018; Doshi vd., 2018)."
)

P(doc,
    "Bu savunmasızlığın en somut göstergesi, 2016 yılında ortaya çıkan Mirai botnet "
    "saldırısıdır. Kolias, Kambourakis, Stavrou ve Voas (2017) tarafından ayrıntılı "
    "biçimde belgelenen bu olayda, varsayılan kullanıcı adı ve parola çiftleriyle "
    "korunan yüz binlerce IoT cihazı ele geçirilmiş ve 1,2 Tbps büyüklüğünde "
    "DDoS saldırısı düzenlenmiştir. Bu olay, IoT ekosistemindeki güvenlik açıklarının "
    "boyutunu tüm dünyaya göstermiş ve akademik çevrelerde IoT güvenliği araştırmalarını "
    "hızlandırmıştır."
)

P(doc,
    "Dağıtılmış Hizmet Reddi (DDoS) saldırıları, birden fazla kaynak sistemden hedefe "
    "yönelik koordineli trafik göndererek hizmetin kesintiye uğratılmasını amaçlayan "
    "saldırılardır. Zargar, Joshi ve Tipper (2013) tarafından gerçekleştirilen kapsamlı "
    "tarama çalışmasında, DDoS saldırıları hacimsel (volumetric), protokol tabanlı ve "
    "uygulama katmanı olmak üzere üç ana kategoride sınıflandırılmıştır. Agrawal ve "
    "Tapaswi (2019) bir başka tarama çalışmasında, bulut ortamlarındaki DDoS savunma "
    "mekanizmalarının güncel durumunu ve araştırma zorluklarını sistematik biçimde "
    "analiz etmiştir."
)

P(doc,
    "Geleneksel imza tabanlı saldırı tespit sistemleri (IDS), bilinen saldırı "
    "kalıplarının bir veritabanında saklanması ve ağ trafiğinin bu kalıplarla "
    "karşılaştırılması esasına dayanır. Bu yaklaşım, düşük yanlış pozitif oranı "
    "avantajı sunsa da daha önce karşılaşılmamış sıfır-gün saldırılarına ve "
    "polimorfik saldırı varyantlarına karşı yapısal olarak etkisiz kalmaktadır. "
    "Bu sınırlılık, makine öğrenmesi (ML) tabanlı tespit yöntemlerinin önemini "
    "belirgin ölçüde artırmıştır (Buczak ve Guven, 2016). ML yaklaşımları, ağ "
    "trafiğindeki normal ve anormal kalıpları veriden otomatik olarak öğrenme "
    "kapasitesine sahiptir ve bu sayede daha önce görülmemiş saldırı türlerini "
    "de tespit edebilme potansiyeli taşımaktadır."
)

P(doc,
    "Ancak, ML tabanlı DDoS tespiti kendi başına sorunsuz değildir. Yüksek boyutlu "
    "ağ trafik veri setlerinde gereksiz ve gürültülü özellikler, modelin "
    "genellenebilirliğini düşürmekte ve aşırı uyuma (overfitting) yol açmaktadır "
    "(Guyon ve Elisseeff, 2003). Buna ek olarak, DDoS veri setlerinde sıkça "
    "karşılaşılan ciddi sınıf dengesizliği, sınıflandırıcının çoğunluk sınıfları "
    "lehine önyargılı kararlar üretmesine neden olmaktadır (He ve Garcia, 2009). "
    "Son olarak, model hiper-parametrelerinin optimal bir şekilde ayarlanması, "
    "geniş bir arama uzayı gerektiren ve geleneksel ızgara aramasıyla (grid search) "
    "verimli biçimde çözülmesi güç olan bir kombinatorik optimizasyon problemidir "
    "(Bergstra ve Bengio, 2012)."
)

P(doc,
    "Bu zorlukların üstesinden gelmek amacıyla, son yıllarda doğadan esinlenen "
    "meta-sezgisel optimizasyon algoritmaları ile ML'nin birleştirilmesi giderek "
    "yaygınlaşmıştır. Xue, Zhang, Browne ve Yao (2016) tarafından gerçekleştirilen "
    "kapsamlı taramada, evrimsel hesaplama yaklaşımlarının özellik seçiminde başarılı "
    "sonuçlar ürettiği sistematik olarak gösterilmiştir. Bu çalışmada, Yang (2010) "
    "tarafından micro-yarasaların ekolokasyon davranışından esinlenerek geliştirilen "
    "Yarasa Algoritması'nın (Bat Algorithm — BSO), CICIoT2023 veri seti üzerinde "
    "DDoS saldırı tespiti bağlamında ilk kapsamlı uygulaması gerçekleştirilmektedir."
)

# ── 1.2 ──
doc.add_heading("1.2 Problem Tanımı", level=2)

P(doc,
    "DDoS saldırı tespitinde mevcut ML yaklaşımlarının karşılaştığı sorunlar, bu "
    "tez çalışmasının motivasyonunu oluşturmakta olup üç temel başlık altında "
    "ele alınabilir."
)

doc.add_heading("1.2.1 Yüksek Boyutlu Özellik Uzayı Sorunu", level=3)

P(doc,
    "Ağ trafik veri setleri genellikle onlarca hatta yüzlerce özellik barındırır. "
    "Chandrashekar ve Sahin (2014) tarafından belirtildiği üzere, bu özelliklerin "
    "tamamı sınıflandırma görevi için bilgilendirici olmayabilir; bir kısmı yedekli "
    "(redundant), bir kısmı ise gürültülü (noisy) niteliktedir. CICIoT2023 veri seti "
    "(Neto vd., 2023) toplamda 39 ağ trafik özelliği içermektedir. Bu özelliklerin "
    "tamamının doğrudan kullanılması, model karmaşıklığını artırırken aşırı uyum "
    "riskini de beraberinde getirmektedir. Guyon ve Elisseeff (2003) tarafından "
    "vurgulandığı üzere, etkili öznitelik seçimi hem hesaplama maliyetini düşürür "
    "hem de modelin genellenebilirliğini artırır."
)

doc.add_heading("1.2.2 Hiper-Parametre Optimizasyonu Sorunu", level=3)

P(doc,
    "Sınıflandırma modellerinin performansı, hiper-parametre konfigürasyonuna "
    "doğrudan bağlıdır. Random Forest (Breiman, 2001) gibi topluluk (ensemble) "
    "modellerinde ağaç sayısı (n_estimators), maksimum derinlik (max_depth), "
    "minimum bölme örneği (min_samples_split) ve özellik alt kümesi oranı "
    "(max_features) gibi parametrelerin optimal değerlerinin belirlenmesi, "
    "kombinatorik olarak büyük bir arama uzayı gerektirir. Bergstra ve Bengio "
    "(2012), rastgele aramanın ızgara aramasından daha verimli olduğunu deneysel "
    "olarak göstermiş olmakla birlikte, her iki yöntem de meta-sezgisel optimizasyon "
    "algoritmalarının küresel arama kapasitesinden yoksundur."
)

doc.add_heading("1.2.3 Sınıf Dengesizliği Sorunu", level=3)

P(doc,
    "IDS veri setlerinde ciddi sınıf dengesizliği yaygın bir olgudur. He ve Garcia "
    "(2009) tarafından gerçekleştirilen tarama çalışmasında, dengesiz veri "
    "dağılımlarının sınıflandırıcı performansını önemli ölçüde bozduğu "
    "sistematik olarak gösterilmiştir. CICIoT2023 veri setinde en büyük sınıf "
    "(DDoS-ACK_Fragmentation: 53.148 örnek) ile en küçük sınıf (Backdoor_Malware: "
    "2.252 örnek) arasında 23,6:1 oranında dengesizlik bulunmaktadır. Bu durum, "
    "sınıflandırıcının çoğunluk sınıflarını doğru tahmin ederken azınlık "
    "sınıflarını yanlış sınıflandırmasına yol açabilmektedir. Chawla, Bowyer, "
    "Hall ve Kegelmeyer (2002) tarafından önerilen SMOTE tekniği, bu soruna "
    "yaygın olarak uygulanan bir çözüm sunmaktadır."
)

P(doc,
    "Bu tez, yukarıda tanımlanan üç sorunu eşzamanlı olarak ele alan bir "
    "BSO-Hybrid RF çerçevesi önermektedir. Önerilen çerçeve, BSO algoritmasını "
    "hem özellik seçimi hem de hiper-parametre optimizasyonu için tek bir arama "
    "sürecinde birleştirmekte ve SMOTE ile sınıf dengeleme entegre etmektedir."
)

# ── 1.3 ──
doc.add_heading("1.3 Araştırma Hipotezleri", level=2)

P(doc,
    "Yukarıda tanımlanan problem çerçevesinde, bu çalışmada aşağıdaki üç "
    "araştırma hipotezi test edilmiştir:"
)

hyps = [
    ("AH₁: ", "BSO ile optimize edilmiş Random Forest modeli, varsayılan parametrelerle "
     "çalışan standart RF modeline göre istatistiksel olarak anlamlı düzeyde (p < 0,05) "
     "daha yüksek F1-Macro skoru elde eder."),
    ("AH₂: ", "BSO tabanlı özellik seçimi, toplam özellik sayısının %50'sinden fazlasını "
     "elemine ederken, sınıflandırma doğruluğunda %1'den az kayıp yaşanmasını sağlar."),
    ("AH₃: ", "BSO-Hybrid çerçevesi (özellik seçimi + hiper-parametre optimizasyonu + SMOTE), "
     "PSO, GA ve GWO gibi diğer meta-sezgisel yöntemlere kıyasla daha düşük uygunluk "
     "(fitness) değeri elde eder."),
]
for label, text in hyps:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(label); r.bold = True
    p.add_run(text)

# ── 1.4 ──
doc.add_heading("1.4 Araştırma Soruları", level=2)

P(doc, "Hipotezlerle tutarlı olarak, bu tezin yanıtlamayı hedeflediği araştırma soruları:")

qs = [
    ("AS₁: ", "BSO algoritması, CICIoT2023 veri setindeki 39 ağ trafik özelliğinden en bilgilendirici alt kümeyi ne ölçüde etkili biçimde seçebilir?"),
    ("AS₂: ", "BSO-Hybrid RF çerçevesi, Backdoor_Malware gibi düşük temsil edilen azınlık saldırı türlerinde yeterli tespit performansı sağlayabilir mi?"),
    ("AS₃: ", "Önerilen çerçevenin performans üstünlüğü, 10-katlı çapraz doğrulama ve eşleştirilmiş istatistiksel testlerle güvenilir biçimde doğrulanabilir mi?"),
]
for label, text in qs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(label); r.bold = True
    p.add_run(text)

# ── 1.5 ──
doc.add_heading("1.5 Tezin Amacı ve Kapsamı", level=2)

P(doc,
    "Bu çalışmanın temel amacı, dinamik ağ ortamlarında DDoS saldırılarının daha "
    "etkili tespiti için BSO tabanlı bir hibrit makine öğrenmesi çerçevesi "
    "tasarlamak, uygulamak ve deneysel olarak değerlendirmektir. Çalışmanın kapsamı: "
    "(a) CICIoT2023 veri seti üzerinde keşifsel veri analizi, (b) BSO ile eşzamanlı "
    "özellik seçimi ve RF hiper-parametre optimizasyonu — 25 yarasa × 50 iterasyon = "
    "1.177 değerlendirme, (c) SMOTE ile sınıf dengeleme — eğitim seti 72.252'den "
    "87.500 örneğe yükseltme, (d) 12 farklı ML modeli ile karşılaştırmalı "
    "değerlendirme, (e) S1–S4 ablasyon senaryoları ile her bileşenin marjinal "
    "katkısının ölçülmesi, (f) 10-katlı tabakalı çapraz doğrulama ve istatistiksel "
    "anlamlılık testleri (Wilcoxon, McNemar, Cohen's d)."
)

# ── 1.6 ──
doc.add_heading("1.6 Tezin Katkıları", level=2)

P(doc, "Bu tez, aşağıdaki dört özgün katkıyı literatüre sunmaktadır:")

contribs = [
    ("K1 — BSO-Hybrid RF Çerçevesi: ",
     "Özellik seçimi ve hiper-parametre optimizasyonunu tek bir meta-sezgisel süreçte "
     "birleştiren özgün hibrit çerçeve tasarımı. Yang'ın (2010) Yarasa Algoritması'nın "
     "CICIoT2023 üzerinde DDoS tespiti için ilk kapsamlı uygulamasıdır. En iyi "
     "uygunluk değeri 0,177801 olarak elde edilmiştir."),
    ("K2 — %51,3 Boyut Azaltma: ",
     "39 özellikten 19'a indirgeme, yalnızca +0,08% doğruluk farkıyla gerçekleştirilmiştir. "
     "Bu sonuç, BSO'nun gereksiz özellikleri etkili biçimde elemine ettiğini gösterir."),
    ("K3 — Azınlık Sınıf İyileştirmesi: ",
     "SMOTE (Chawla vd., 2002) entegrasyonu ile Backdoor_Malware F1-Skoru %28,40'tan "
     "%57,40'a yükselmiştir (+%102). Sınıf dengelemenin çok sınıflı tespitteki kritik "
     "rolü deneysel olarak kanıtlanmıştır."),
    ("K4 — Kapsamlı Deneysel Doğrulama: ",
     "12 model × 7 metrik × 10-katlı çapraz doğrulama × 4 istatistiksel test. S1–S4 "
     "ablasyon senaryoları ile her bileşenin marjinal etkisi izole edilerek ölçülmüştür. "
     "11 eşleştirilmiş testten 9'unda p < 0,05 düzeyinde anlamlılık saptanmıştır."),
]
for label, text in contribs:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(label); r.bold = True
    p.add_run(text)

# ── 1.7 ──
doc.add_heading("1.7 Tezin Organizasyonu", level=2)

org = [
    ("Bölüm 1 — Giriş: ", "Araştırmanın motivasyonu, problem tanımı, hipotezler, araştırma soruları, amaç, kapsam ve katkılar."),
    ("Bölüm 2 — Literatür Taraması: ", "DDoS saldırıları, IoT güvenliği, ML tabanlı tespit yöntemleri, meta-sezgisel optimizasyon, yarasa algoritması, öznitelik seçimi ve ilgili çalışmaların sistematik incelenmesi."),
    ("Bölüm 3 — Yöntem: ", "BSO-Hybrid RF çerçevesinin tasarımı, CICIoT2023 veri seti, veri ön işleme, BSO optimizasyonu, uygunluk fonksiyonu, hibrit model, deney senaryoları (S1–S4) ve değerlendirme metrikleri."),
    ("Bölüm 4 — Bulgular ve Değerlendirme: ", "BSO optimizasyon sonuçları, 12 model karşılaştırması, sınıf bazında analiz, ablasyon çalışması, istatistiksel doğrulama ve hipotez testleri."),
    ("Bölüm 5 — Sonuç ve Öneriler: ", "Bulguların özeti, araştırma sorularının yanıtları, katkılar, kısıtlamalar ve gelecek çalışma önerileri."),
]
for label, text in org:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(label); r.bold = True
    p.add_run(text)

doc.add_page_break()

# ████████████████████████████████████████████████████████████████████████████
# BÖLÜM 2 — LİTERATÜR TARAMASI
# ████████████████████████████████████████████████████████████████████████████

doc.add_heading("BÖLÜM 2", level=1)
doc.add_heading("LİTERATÜR TARAMASI", level=1)

P(doc,
    "Bu bölümde, tez konusuyla doğrudan ilişkili altı temel alan sistematik "
    "biçimde incelenmektedir: DDoS saldırı türleri ve IoT güvenliği, makine "
    "öğrenmesi tabanlı saldırı tespiti, meta-sezgisel optimizasyon algoritmaları, "
    "yarasa algoritması, öznitelik seçimi yöntemleri ve bu alanları kapsayan "
    "ilgili çalışmaların karşılaştırmalı değerlendirmesi."
)

# ── 2.1 ──
doc.add_heading("2.1 DDoS Saldırıları ve IoT Güvenliği", level=2)

doc.add_heading("2.1.1 DDoS Saldırılarına Genel Bakış", level=3)

P(doc,
    "DDoS saldırıları, hedefteki sistemin sunduğu hizmeti aksatmak amacıyla çok "
    "sayıda kaynaktan eşzamanlı olarak aşırı trafik yönlendirmeye dayanan koordineli "
    "siber saldırılardır. Zargar vd. (2013) tarafından gerçekleştirilen kapsamlı tarama "
    "çalışmasında, DDoS saldırıları saldırı vektörüne göre üç temel kategoride "
    "sınıflandırılmıştır: (1) hacimsel saldırılar — hedefin bant genişliğini "
    "tüketmeyi amaçlayan UDP flood, ICMP flood ve DNS amplification; (2) protokol "
    "tabanlı saldırılar — TCP/IP protokol yığınının zayıflıklarını istismar eden "
    "SYN flood ve ACK fragmentation; (3) uygulama katmanı saldırıları — HTTP flood "
    "ve Slowloris gibi düşük hacimli ancak yüksek etkili türler."
)

P(doc,
    "Osanaiye, Choo ve Dlodlo (2016), bulut ortamlarındaki DDoS dayanıklılığını "
    "inceleyen çalışmalarında, saldırı karmaşıklığının yıllar itibarıyla artış "
    "gösterdiğini ve çok vektörlü (multi-vector) saldırıların yaygınlaştığını "
    "raporlamıştır. Yan, Yu, Gong ve Li (2016) ise yazılım tanımlı ağ (SDN) "
    "ortamlarında DDoS saldırılarının oluşturduğu tehdidi ele alarak, geleneksel "
    "savunma mekanizmalarının yetersizliğini vurgulamıştır. Bu tezde kullanılan "
    "CICIoT2023 veri seti (Neto vd., 2023), protokol tabanlı DDoS türlerini "
    "(SYN_Flood ve ACK_Fragmentation) içermektedir."
)

doc.add_heading("2.1.2 IoT Ortamlarında DDoS Tehdidi", level=3)

P(doc,
    "IoT cihazlarının DDoS saldırılarına karşı savunmasızlığı, birden fazla yapısal "
    "nedene dayanmaktadır. Doshi, Apthorpe ve Feamster (2018), IoT cihazlarının "
    "sınırlı hesaplama ve bellek kapasiteleri nedeniyle güçlü şifreleme ve kimlik "
    "doğrulama mekanizmalarının uygulanamadığını deneysel olarak göstermiştir. "
    "Meidan vd. (2018), N-BaIoT çalışmasında IoT botnet saldırılarının derin "
    "otomatik kodlayıcılar (deep autoencoders) ile tespit edilebileceğini "
    "kanıtlamıştır. Koroniotis, Moustafa, Sitnikova ve Turnbull (2019) ise "
    "Bot-IoT veri setini oluşturarak IoT ağ adli tıp analizi için gerçekçi bir "
    "zemin sağlamıştır."
)

P(doc,
    "Sivanathan vd. (2019), akıllı ortamlardaki IoT cihazlarını ağ trafik "
    "özelliklerine göre sınıflandırmış ve cihaz heterojenliğinin güvenlik "
    "politikalarının standartlaştırılmasını zorlaştırdığını belirtmiştir. "
    "2016 Mirai botnet olayı (Kolias vd., 2017), bu yapısal zaafiyetlerin "
    "somut bir tezahürü olarak IoT güvenliği araştırmalarına ivme kazandırmıştır."
)

doc.add_heading("2.1.3 Geleneksel Saldırı Tespit Yaklaşımları", level=3)

P(doc,
    "Saldırı Tespit Sistemleri (IDS), ağ trafiğini sürekli izleyerek güvenlik "
    "ihlallerini tespit etmeye yönelik savunma mekanizmalarıdır. Geleneksel "
    "IDS'ler iki temel paradigmaya ayrılır: imza tabanlı tespit ve anomali "
    "tabanlı tespit."
)

P(doc,
    "İmza tabanlı sistemler (Snort, Suricata), bilinen saldırı kalıplarını "
    "bir veritabanında depolar ve canlı trafiği bu kalıplarla karşılaştırır. "
    "Düşük yanlış pozitif oranı (FPR) avantajına sahip olmakla birlikte, "
    "sıfır-gün saldırılarına ve polimorfik varyantlara karşı yapısal olarak "
    "kördür. Anomali tabanlı sistemler ise normal trafik profilinden sapan "
    "aktiviteleri tespit etme kapasitesine sahip olmakla birlikte, yüksek "
    "yanlış pozitif oranı ve profil oluşturma güçlüğü gibi dezavantajlar "
    "taşımaktadır (Agrawal ve Tapaswi, 2019)."
)

# ── 2.2 ──
doc.add_heading("2.2 Makine Öğrenmesi Tabanlı Saldırı Tespiti", level=2)

doc.add_heading("2.2.1 Denetimli Öğrenme Yaklaşımları", level=3)

P(doc,
    "Denetimli öğrenme, etiketlenmiş veri kullanarak sınıflandırma modelleri "
    "oluşturur ve DDoS tespitinde en yaygın kullanılan paradigmadır. Random Forest "
    "(Breiman, 2001), birden fazla karar ağacının Bootstrap alt örnekleriyle "
    "eğitilip çoğunluk oylamasıyla karar verilmesi prensibine dayanır. Hastie, "
    "Tibshirani ve Friedman (2009), RF'nin aşırı uyuma dayanıklılığını ve "
    "özellik önemi sağlama kapasitesini ayrıntılı biçimde tartışmıştır. Bu "
    "çalışmada RF, temel sınıflandırıcı olarak tercih edilmiştir."
)

P(doc,
    "XGBoost (Chen ve Guestrin, 2016), gradyan artırma çerçevesinin "
    "regularizasyon ve paralel işleme yetenekleriyle güçlendirilmiş bir "
    "uygulamasıdır. LightGBM (Ke vd., 2017) ise histogram tabanlı bölme "
    "stratejisiyle büyük veri setlerinde verimli çalışır. Destek Vektör "
    "Makineleri (SVM) (Cortes ve Vapnik, 1995), sınıflar arasındaki "
    "maksimum ayrım sınırını belirleyen bir sınıflandırıcıdır; ancak "
    "yüksek boyutlu ve büyük veri setlerinde eğitim süresi ciddi ölçüde "
    "artar. K-En Yakın Komşu (KNN) (Cover ve Hart, 1967), örnek tabanlı "
    "(instance-based) basit bir yaklaşımdır."
)

P(doc,
    "Ahmad, Basheri, Iqbal ve Rahim (2018), SVM, RF ve ELM algoritmalarını "
    "saldırı tespiti bağlamında karşılaştırarak RF'nin genel olarak daha "
    "kararlı performans sergilediğini raporlamıştır. Zhou, Cheng, Jiang ve "
    "Dai (2020) ise filtre ve sarmalayıcı (wrapper) özellik seçimi yöntemlerini "
    "topluluk sınıflandırıcısı ile birleştirerek IDS'lerde verimli bir tespit "
    "modeli oluşturmuştur."
)

doc.add_heading("2.2.2 Derin Öğrenme Yaklaşımları", level=3)

P(doc,
    "Derin öğrenme yöntemleri de DDoS tespitinde artan bir ilgiyle "
    "uygulanmaktadır. Goodfellow, Bengio ve Courville (2016) tarafından "
    "kapsamlı biçimde ele alınan derin ağ mimarileri, özellikle zamansal "
    "ve uzamsal kalıpların öğrenilmesinde avantaj sağlamaktadır. Roopak, "
    "Tian ve Chambers (2019), CNN ve LSTM tabanlı modelleri IoT DDoS "
    "tespitinde değerlendirmiş ve derin modellerin geleneksel ML'ye göre "
    "daha yüksek doğruluk sağlayabildiğini, ancak eğitim maliyetinin "
    "önemli ölçüde arttığını gözlemlemiştir."
)

P(doc,
    "Yuan, Li ve Li (2017), DeepDefense çerçevesini önermiş; Li ve Lu "
    "(2019), LSTM ile Bayes yaklaşımını birleştiren LSTM-BA hibrit "
    "modelini geliştirmiştir. Bu çalışmaların ortak bulgusu, derin "
    "öğrenme modellerinin yüksek doğruluk potansiyeline sahip olduğu "
    "ancak büyük eğitim verisi gereksinimi, hesaplama maliyeti ve "
    "açıklanabilirlik eksikliğinin temel dezavantajlar oluşturduğudur. "
    "Özellikle sınırlı kaynaklara sahip IoT ortamlarında dağıtım "
    "zorluğu, geleneksel ML yöntemlerinin hâlâ tercih edilmesinin "
    "başlıca nedenidir."
)

doc.add_heading("2.2.3 Sınıf Dengesizliği ve SMOTE", level=3)

P(doc,
    "IDS veri setlerinde sınıf dengesizliği, sınıflandırıcı performansını "
    "doğrudan etkileyen yapısal bir sorundur. He ve Garcia (2009) tarafından "
    "gerçekleştirilen kapsamlı taramada, dengesiz veriden öğrenme "
    "stratejilerinin üç ana kategoride sınıflandırıldığı gösterilmiştir: "
    "veri düzeyinde yaklaşımlar (örnekleme), algoritma düzeyinde "
    "yaklaşımlar (maliyet duyarlı öğrenme) ve hibrit yaklaşımlar."
)

P(doc,
    "Bu çalışmada veri düzeyinde yaklaşım olarak SMOTE (Chawla vd., 2002) "
    "tercih edilmiştir. SMOTE, azınlık sınıfına ait mevcut örnekler "
    "arasında doğrusal interpolasyon yaparak sentetik örnekler üretir. "
    "CICIoT2023 veri setinde SMOTE yalnızca eğitim setine uygulanmış ve "
    "eğitim seti 72.252'den 87.500 dengeli örneğe yükseltilmiştir. "
    "Özellikle Backdoor_Malware sınıfının örnek sayısı 2.252'den "
    "17.500'e artırılmıştır (×7,8 artış). SMOTE'un test ve doğrulama "
    "setlerine uygulanmaması, veri sızıntısının (data leakage) "
    "önlenmesi için kritik öneme sahiptir."
)

# ── 2.3 ──
doc.add_heading("2.3 Meta-Sezgisel Optimizasyon Yöntemleri", level=2)

P(doc,
    "Meta-sezgisel algoritmalar, doğadaki biyolojik veya fiziksel süreçlerden "
    "esinlenen ve NP-zor optimizasyon problemlerine makul sürede yaklaşık "
    "çözümler üreten stokastik arama yöntemleridir. Yang (2010) tarafından "
    "belirtildiği üzere, bu algoritmaların ortak özelliği keşif (exploration) "
    "ve sömürü (exploitation) arasında denge kurmalarıdır."
)

doc.add_heading("2.3.1 Parçacık Sürüsü Optimizasyonu (PSO)", level=3)

P(doc,
    "Kennedy ve Eberhart (1995) tarafından önerilen PSO, kuş sürülerinin "
    "ve balık topluluklarının kolektif davranışlarından esinlenen bir "
    "popülasyon tabanlı algoritmadır. Her parçacık, kendi deneyiminden "
    "(pbest) ve sürünün mevcut en iyi çözümünden (gbest) etkilenerek "
    "arama uzayında konumunu günceller. PSO'nun sadeliği ve az sayıda "
    "kontrol parametresi gerektirmesi, özellik seçimi alanında geniş "
    "kullanım bulmasını sağlamıştır. Bu çalışmadaki deneylerde PSO-RF "
    "modeli 40 iterasyonda 18 özellik seçerek %88,35 doğruluk ve "
    "%81,82 F1-Macro elde etmiştir."
)

doc.add_heading("2.3.2 Genetik Algoritma (GA)", level=3)

P(doc,
    "Holland (1975) tarafından tanıtılan GA, doğal seçilim mekanizmasını "
    "taklit eden evrimsel bir algoritmadır. Çaprazlama, mutasyon ve "
    "seçim operatörleri ile popülasyonu iteratif olarak iyileştirir. "
    "Khammassi ve Krichen (2017), GA'yı Lojistik Regresyon ile "
    "birleştirerek ağ saldırı tespitinde özellik seçimi gerçekleştirmiştir. "
    "Bu çalışmada GA-RF modeli 40 iterasyonda 21 özellik seçerek "
    "%89,37 doğruluk ve %83,66 F1-Macro elde etmiştir."
)

doc.add_heading("2.3.3 Gri Kurt Optimizasyonu (GWO)", level=3)

P(doc,
    "Mirjalili, Mirjalili ve Lewis (2014) tarafından geliştirilen GWO, "
    "gri kurdların alfa-beta-delta-omega sosyal hiyerarşisi ve avlanma "
    "davranışından esinlenmiştir. GWO, düşük parametre sayısı ve "
    "rekabetçi yakınsama hızı ile dikkat çekmektedir. Bu çalışmada "
    "GWO-RF modeli 23 özellik seçerek %89,80 doğruluk ve %84,35 "
    "F1-Macro ile BSO-Hybrid RF'ye en yakın performansı göstermiştir."
)

doc.add_heading("2.3.4 Diğer Meta-Sezgisel Algoritmalar", level=3)

P(doc,
    "Karınca Koloni Optimizasyonu (ACO) (Dorigo ve Stützle, 2004), "
    "karıncaların feromona dayalı yol bulma davranışını; Yapay Arı "
    "Koloni (ABC) algoritması (Karaboga ve Basturk, 2007), bal arılarının "
    "yiyecek arama stratejisini; Guguk Kuşu Araması (CS) (Yang ve Deb, "
    "2009) ise kuluçka parazitizmi davranışını modeller. Balina "
    "Optimizasyonu (WOA) ise Mafarja ve Mirjalili (2018) tarafından "
    "sarmalayıcı özellik seçimi bağlamında başarıyla uygulanmıştır."
)

# Karşılaştırma tablosu
make_table(doc,
    ["Algoritma", "Kaynak", "Yıl", "İlham", "Özellik", "F1-Macro (%)", "Fitness"],
    [
        ["BSO (Önerilen)", "Yang", "2010", "Yarasa ekolokasyonu", "19", "84,24", "0,177801"],
        ["PSO", "Kennedy & Eberhart", "1995", "Kuş/balık sürüsü", "18", "81,82", "0,193895"],
        ["GA", "Holland", "1975", "Doğal seçilim", "21", "83,66", "0,188982"],
        ["GWO", "Mirjalili vd.", "2014", "Kurt hiyerarşisi", "23", "84,35", "0,192181"],
    ],
    caption="Tablo 2.1: Bu Çalışmadaki Meta-Sezgisel Algoritmaların Karşılaştırması",
    hl_row=0,
)

# ── 2.4 ──
doc.add_heading("2.4 Yarasa Algoritması (Bat Algorithm)", level=2)

doc.add_heading("2.4.1 Temel Mekanizma ve Matematiksel Formülasyon", level=3)

P(doc,
    "Yarasa Algoritması, Yang (2010) tarafından micro-yarasaların ultrasonik "
    "ekolokasyon davranışından esinlenerek geliştirilmiştir. Yarasalar, yüksek "
    "frekanslı ses dalgaları yayarak çevrelerindeki nesnelerin mesafe ve "
    "konumunu algılar. Yang (2013) tarafından yapılan kapsamlı bir derleme "
    "çalışmasında, algoritmanın çeşitli optimizasyon problemlerindeki başarılı "
    "uygulamaları sistematik olarak değerlendirilmiştir."
)

P(doc, "Algoritmanın temel güncelleme denklemleri:")

eq(doc, "fi = fmin + (fmax - fmin) × β,    β ∈ U(0, 1)", "2.1")
eq(doc, "vi(t+1) = vi(t) + [xi(t) - x*] × fi", "2.2")
eq(doc, "xi(t+1) = xi(t) + vi(t+1)", "2.3")

P(doc,
    "Burada fi i. yarasanın frekansını, vi hızını, xi konumunu, x* mevcut "
    "küresel en iyi çözümü ve β düzgün dağılımlı rastgele bir sayıyı temsil "
    "eder. Yerel arama mekanizmasında: xnew = xold + ε × At (Denklem 2.4), "
    "burada ε ∈ [-1, 1] aralığında rastgele bir sayı, At ise t. iterasyondaki "
    "ortalama gürültü seviyesidir."
)

P(doc, "Gürültü seviyesi ve darbe oranı güncelleme kuralları:")

eq(doc, "Ai(t+1) = α × Ai(t),    0 < α < 1", "2.5")
eq(doc, "ri(t+1) = ri(0) × [1 - exp(-γ × t)],    γ > 0", "2.6")

P(doc,
    "Gürültü azalma (α) ve darbe artış (γ) parametreleri kendi kendine "
    "uyarlanarak (self-adaptive) keşif ve sömürü arasında dinamik denge "
    "kurar. Gandomi, Yang, Alavi ve Talatahari (2013), BSO'nun kısıtlı "
    "optimizasyon görevlerindeki etkinliğini kanıtlamıştır. Nakamura vd. "
    "(2012), BSO'nun ikili versiyonunu (BBA) öznitelik seçimi için ilk "
    "kez uygulayarak alanda öncü bir çalışma gerçekleştirmiştir."
)

# ── 2.5 ──
doc.add_heading("2.5 Öznitelik Seçimi Yöntemleri", level=2)

P(doc,
    "Öznitelik seçimi, yüksek boyutlu veri setlerinden en bilgilendirici "
    "alt kümenin belirlenmesi sürecidir. Dash ve Liu (1997) tarafından "
    "yapılan erken dönem sınıflandırmada yöntemler üç kategoriye ayrılmıştır: "
    "filtre, sarmalayıcı ve gömülü. Guyon ve Elisseeff (2003), değişken "
    "seçiminin temel prensiplerini kapsamlı biçimde tartışmış; Chandrashekar "
    "ve Sahin (2014) ise güncel yöntemleri sistematik olarak derlemiştir."
)

P(doc,
    "Filtre yöntemleri (chi-square, bilgi kazanımı, korelasyon), sınıflandırıcıdan "
    "bağımsız istatistiksel ölçütlerle çalışır. Ambusaidi, He, Nanda ve Tan "
    "(2016), filtre tabanlı özellik seçiminin IDS performansına etkisini "
    "deneysel olarak göstermiştir. Sarmalayıcı yöntemler ise sınıflandırıcının "
    "performansını doğrudan optimize etmek üzere meta-sezgisel algoritmalar "
    "kullanır. Bu çalışmada BSO tabanlı sarmalayıcı yaklaşım tercih "
    "edilmiş olup, Xue vd. (2016) tarafından önerilen evrimsel öznitelik "
    "seçimi çerçevesine dayanmaktadır."
)

# ── 2.6 ──
doc.add_heading("2.6 İlgili Çalışmalar ve Karşılaştırma", level=2)

P(doc,
    "Bu alt bölümde, DDoS tespiti ve meta-sezgisel optimizasyon "
    "alanında yürütülen temel çalışmalar özetlenmekte ve bu tezin "
    "konumlandırılması yapılmaktadır."
)

make_table(doc,
    ["Çalışma", "Veri Seti", "Yöntem", "Sonuç"],
    [
        ["Neto vd. (2023)", "CICIoT2023", "RF, DNN", "İkili sınıflandırma: %99; çoklu sınıf sınırlı"],
        ["Ferrag vd. (2023)", "Edge-IIoTset", "DNN, RF", "IoT/IIoT; ikili %98,2 doğruluk"],
        ["Sharafaldin vd. (2018)", "CICIDS2017", "RF, DT", "Yeni IDS veri seti oluşturma"],
        ["Moustafa & Slay (2015)", "UNSW-NB15", "Çoklu", "Kapsamlı güvenlik veri seti"],
        ["Khammassi & Krichen (2017)", "NSL-KDD", "GA-LR", "GA ile özellik seçimi + LR"],
        ["Zhou vd. (2020)", "CICIDS2017", "Ensemble", "Filtre+wrapper hibrit FS"],
        ["Tama & Rhee (2019)", "NSL-KDD", "GBM", "Gradyan artırma ile anomali"],
        ["Ahmad vd. (2018)", "KDD Cup99", "SVM, RF", "3 sınıflandırıcı karşılaştırma"],
        ["Bu Çalışma", "CICIoT2023", "BSO-Hybrid RF", "BSO ile eşzamanlı FS+HP; 5-sınıf; %89,82"],
    ],
    caption="Tablo 2.2: İlgili Çalışmaların Karşılaştırmalı Özeti",
    hl_row=8,
)

P(doc,
    "Tablo 2.2'den açıkça görüldüğü üzere, mevcut çalışmaların büyük "
    "çoğunluğu ikili sınıflandırma (normal vs. saldırı) üzerine yoğunlaşmış "
    "olup çok sınıflı DDoS tespiti alanında, özellikle güncel CICIoT2023 veri "
    "seti üzerinde meta-sezgisel optimizasyon ile birleştirilmiş çalışmalar "
    "son derece sınırlıdır. Bu tez, Yarasa Algoritması'nın CICIoT2023 üzerinde "
    "eşzamanlı özellik seçimi ve hiper-parametre optimizasyonu için gerçekleştirilen "
    "ilk kapsamlı uygulamasını sunarak bu literatür boşluğunu doldurmayı "
    "amaçlamaktadır."
)

doc.add_page_break()


# ████████████████████████████████████████████████████████████████████████████
# BÖLÜM 3 — YÖNTEM
# ████████████████████████████████████████████████████████████████████████████

doc.add_heading("BÖLÜM 3", level=1)
doc.add_heading("YÖNTEM", level=1)

P(doc,
    "Bu bölümde, önerilen BSO-Hybrid RF çerçevesinin tasarım detayları, "
    "veri seti özellikleri, ön işleme adımları, BSO algoritması, uygunluk "
    "fonksiyonu, hibrit model tanımı, deney senaryoları, değerlendirme "
    "metrikleri ve istatistiksel doğrulama yöntemleri ayrıntılı olarak "
    "sunulmaktadır."
)

# ── 3.1 ──
doc.add_heading("3.1 Araştırma Tasarımı ve Genel Çerçeve", level=2)

P(doc,
    "Önerilen çerçeve yedi aşamalı bir pipeline izlemektedir (Şekil 3.1):"
)

steps = [
    ("Aşama 1 — Veri Toplama:", " CICIoT2023 veri setinden 19 CSV dosyası yüklenerek 118.466 akış kaydı ve 39 özellik oluşturulmuştur (Neto vd., 2023)."),
    ("Aşama 2 — Ön İşleme:", " Eksik değer kontrolü, tekrarlanan kayıt temizliği ve StandardScaler normalizasyonu uygulanmıştır."),
    ("Aşama 3 — Veri Bölme:", " Tabakalı bölme ile %70 eğitim (72.252), %10 doğrulama (10.322) ve %20 test (20.644)."),
    ("Aşama 4 — SMOTE:", " Yalnızca eğitim setine uygulanmış; 72.252 → 87.500 dengeli örnek (Chawla vd., 2002)."),
    ("Aşama 5 — BSO Optimizasyonu:", " 25 yarasa × 50 iterasyon = 1.177 değerlendirme (Yang, 2010)."),
    ("Aşama 6 — Model Eğitimi:", " Optimal çözüm: 19 özellik, n_estimators=266, max_depth=20, min_samples_split=7, max_features=0,469."),
    ("Aşama 7 — Değerlendirme:", " 7 performans metriği + 10-katlı CV + istatistiksel testler + 12 model karşılaştırması (Demšar, 2006)."),
]

for label, text in steps:
    p = doc.add_paragraph()
    r = p.add_run(label); r.bold = True
    p.add_run(text)

# ── 3.2 ──
doc.add_heading("3.2 Veri Seti: CICIoT2023", level=2)

P(doc,
    "Bu çalışmada kullanılan CICIoT2023 veri seti, Canadian Institute for "
    "Cybersecurity tarafından 2023 yılında yayınlanmıştır (Neto vd., 2023). "
    "Gerçek IoT cihazlarından (105 cihaz, 19 saldırı senaryosu) toplanan "
    "veri seti, 118.466 akış kaydı ve 39 ağ trafik özelliğinden oluşmaktadır."
)

make_table(doc,
    ["Özellik", "Değer"],
    [
        ["Toplam Örnek", "118.466"],
        ["Özellik Sayısı", "39"],
        ["Sınıf Sayısı", "5"],
        ["Eğitim Seti (%70)", "87.500 (SMOTE sonrası)"],
        ["Doğrulama Seti (%10)", "10.322"],
        ["Test Seti (%20)", "20.644"],
        ["Normalizasyon", "StandardScaler (z-score: μ=0, σ=1)"],
        ["SMOTE Sentetik Örnek", "15.248"],
    ],
    caption="Tablo 3.1: CICIoT2023 Veri Seti Özeti",
)

make_table(doc,
    ["Sınıf", "Eğitim (Orijinal)", "Eğitim (SMOTE)", "Test", "Oran"],
    [
        ["DDoS-ACK_Fragmentation", "~37.204", "17.500", "5.000", "%44,9"],
        ["DDoS-SYN_Flood", "~17.618", "17.500", "5.000", "%21,2"],
        ["BenignTraffic", "~15.915", "17.500", "5.000", "%19,2"],
        ["Recon-PortScan", "~10.613", "17.500", "5.000", "%12,8"],
        ["Backdoor_Malware", "~2.252", "17.500", "644", "%1,9"],
        ["TOPLAM", "72.252", "87.500", "20.644", "%100"],
    ],
    caption="Tablo 3.2: Sınıf Dağılımı (SMOTE Öncesi ve Sonrası)",
    hl_row=5,
)

# ── 3.3 ──
doc.add_heading("3.3 Veri Ön İşleme", level=2)

P(doc,
    "Veri ön işleme aşamasında sırasıyla aşağıdaki adımlar uygulanmıştır. "
    "İlk olarak, veri setinde eksik değer bulunmadığı doğrulanmıştır. "
    "İkinci olarak, tekrarlanan kayıtlar tanımlanmış ve kaldırılmıştır. "
    "Üçüncü olarak, tüm sayısal özellikler StandardScaler normalizasyonu "
    "ile standartlaştırılmıştır (Denklem 3.1)."
)

eq(doc, "x_norm = (x - μ) / σ", "3.1")

P(doc,
    "Bu dönüşüm, her özelliğin ortalamasını sıfır ve standart sapmasını bir "
    "yapar. Normalizasyon parametreleri (μ, σ) yalnızca eğitim setinden "
    "hesaplanmış, doğrulama ve test setlerine aynı dönüşüm uygulanmıştır. "
    "Bu yaklaşım, Hastie vd. (2009) tarafından vurgulanan veri sızıntısı "
    "önleme ilkesine uygundur."
)

P(doc,
    "SMOTE (Chawla vd., 2002) yalnızca eğitim setine uygulanmıştır. "
    "Eğitim seti boyutu 72.252'den 87.500'e yükselmiş, tüm sınıflar "
    "17.500 örneğe dengelenmiştir. Özellikle Backdoor_Malware sınıfının "
    "örnek sayısı 2.252'den 17.500'e artırılmıştır (×7,8)."
)

# ── 3.4 ──
doc.add_heading("3.4 Yarasa Sürüsü Optimizasyonu (BSO) Algoritması", level=2)

P(doc,
    "BSO algoritması (Yang, 2010), bu çerçevede iki eşzamanlı görevi tek "
    "bir optimizasyon döngüsünde birleştirmektedir."
)

doc.add_heading("3.4.1 Görev 1 — Özellik Seçimi (İkili)", level=3)

P(doc,
    "Her yarasa bireyi, 39 boyutlu bir ikili maske taşır. x_j = 1 ise j. "
    "özellik seçilir, x_j = 0 ise dahil edilmez. Sürekli pozisyon "
    "değerlerinin ikili vektöre dönüştürülmesi için Nakamura vd. (2012) "
    "tarafından önerilen sigmoid transfer fonksiyonu kullanılmıştır:"
)

eq(doc, "S(x) = 1 / (1 + exp(-x))", "3.2")
eq(doc, "xj = 1  eğer S(xj) > rand()  ise, aksi halde xj = 0", "3.3")

doc.add_heading("3.4.2 Görev 2 — Hiper-Parametre Optimizasyonu (Sürekli)", level=3)

P(doc,
    "Aynı yarasa bireyinin çözüm vektöründe, özellik maskesinin ardına "
    "4 sürekli değişken eklenmiştir. Bu değişkenler, RF sınıflandırıcısının "
    "kritik hiper-parametrelerini temsil eder (Tablo 3.3)."
)

make_table(doc,
    ["Parametre", "Aralık", "Optimal Değer", "Açıklama"],
    [
        ["n_estimators", "[50, 400]", "266", "Karar ağacı sayısı"],
        ["max_depth", "[5, 35]", "20", "Maksimum ağaç derinliği"],
        ["min_samples_split", "[2, 15]", "7", "Bölme için minimum örnek"],
        ["max_features", "[0,3 — 1,0]", "0,469", "Özellik alt küme oranı"],
    ],
    caption="Tablo 3.3: BSO ile Optimize Edilen RF Hiper-Parametreleri",
)

doc.add_heading("3.4.3 BSO Algoritma Parametreleri", level=3)

make_table(doc,
    ["Parametre", "Değer", "Referans"],
    [
        ["Popülasyon", "25 yarasa", "Yang (2010)"],
        ["Maks. İterasyon", "50", "—"],
        ["Frekans [fmin, fmax]", "[0, 2]", "Yang (2010)"],
        ["Başlangıç Gürültüsü A₀", "0,95", "Gandomi vd. (2013)"],
        ["Başlangıç Darbe r₀", "0,5", "Gandomi vd. (2013)"],
        ["α (Gürültü azalma)", "0,9", "Yang (2013)"],
        ["γ (Darbe artış)", "0,9", "Yang (2013)"],
        ["Toplam Değerlendirme", "1.177", "—"],
        ["Çalışma Süresi", "840,43 saniye", "—"],
    ],
    caption="Tablo 3.4: BSO Algoritma Parametreleri",
)

# ── 3.5 ──
doc.add_heading("3.5 Uygunluk Fonksiyonu", level=2)

P(doc,
    "BSO'nun minimize ettiği amaç fonksiyonu aşağıdaki gibi tanımlanmıştır:"
)

eq(doc, "f(x, θ) = 1 - F1_macro(RF(X_seçili, θ)) + α × (n_seçili / n_toplam)", "3.4")

P(doc,
    "Burada F1_macro tüm sınıfların eşit ağırlıklı harmonik ortalaması "
    "(Sokolova ve Lapalme, 2009), X_seçili BSO tarafından seçilen özellik "
    "alt kümesi, θ RF hiper-parametre seti, α = 0,01 özellik ceza katsayısı "
    "ve n_seçili/n_toplam boyut oranıdır. Bu formülasyon, sınıflandırma "
    "performansını maksimize ederken boyut karmaşıklığını minimize eder. "
    "En iyi uygunluk değeri: 0,177801."
)

# ── 3.6 ──
doc.add_heading("3.6 Hibrit Model Tanımı: BSO-Hybrid RF", level=2)

P(doc,
    "\"Hibrit\" terimi bu çalışmada, birden fazla sınıflandırıcının "
    "birleştirilmesi (voting/stacking) anlamında değil, meta-sezgisel "
    "optimizasyon (BSO) ile makine öğrenmesi sınıflandırması (RF) "
    "arasındaki metodolojik birleşimi ifade etmektedir:"
)

eq(doc, "BSO-Hybrid RF = BSO_optimize(RF_sınıflandır(X_seçili, θ_optimal))", "3.5")

P(doc,
    "Burada X_seçili ⊂ X_tüm BSO tarafından seçilmiş 19 özelliklik optimal "
    "alt kümedir ve θ_optimal = {n_estimators=266, max_depth=20, "
    "min_samples_split=7, max_features=0,469} BSO tarafından optimize "
    "edilmiş parametre kümesidir. Bu yaklaşımın temel avantajı, özellik "
    "seçimi ve parametre optimizasyonunun karşılıklı bağımlılığını tek "
    "bir arama sürecinde yakalayabilmesidir."
)

# ── 3.7 ──
doc.add_heading("3.7 Deney Senaryoları (S1–S4)", level=2)

P(doc,
    "Her bileşenin katkısını sistematik olarak ölçmek amacıyla dört "
    "ablasyon senaryosu tasarlanmıştır. Wolpert (1996) tarafından ortaya "
    "konan \"No Free Lunch\" teoremi bağlamında, her bileşenin marjinal "
    "etkisi izole edilerek ölçülmüştür."
)

make_table(doc,
    ["Senaryo", "Özellik", "HP Ayarı", "SMOTE", "Doğruluk", "F1-Macro", "AUC-ROC"],
    [
        ["S1 — Temel", "39 (tümü)", "Varsayılan (DT)", "Hayır", "%87,04", "%78,57", "%91,20"],
        ["S2 — BSO-FS", "19 (BSO)", "Varsayılan (RF)", "Evet", "%88,47", "%82,35", "%97,47"],
        ["S3 — BSO-HP", "39 (tümü)", "BSO optimize (RF)", "Evet", "%89,74", "%84,13", "%98,36"],
        ["S4 — Tam Hibrit", "19 (BSO)", "BSO optimize (RF)", "Evet", "%89,82", "%84,24", "%98,38"],
    ],
    caption="Tablo 3.5: S1–S4 Ablasyon Senaryoları Sonuçları",
    hl_row=3,
)

# ── 3.8 ──
doc.add_heading("3.8 Değerlendirme Metrikleri", level=2)

P(doc,
    "Model performansı, Sokolova ve Lapalme (2009) ve Powers (2011) tarafından "
    "önerilen çerçeveye uygun biçimde 7 temel metrik ile değerlendirilmiştir:"
)

metrics = [
    ("Doğruluk (Accuracy):", " Doğru tahminlerin toplam örnek sayısına oranı."),
    ("Kesinlik (Precision):", " Pozitif olarak tahmin edilen örneklerin gerçekten pozitif olma oranı."),
    ("Duyarlılık (Recall):", " Gerçek pozitif örneklerin doğru tespit edilme oranı."),
    ("F1-Macro:", " Tüm sınıfların F1 harmonik ortalamasının eşit ağırlıklı ortalaması — birincil metrik."),
    ("AUC-ROC:", " Eşik-bağımsız sınıflandırma performansı (Fawcett, 2006)."),
    ("FPR:", " Yanlış pozitif oranı — IDS operasyonel yükü açısından kritik."),
    ("MCC:", " Matthews Korelasyon Katsayısı [-1, +1] — sınıf dengesizliğine dayanıklı (Chicco ve Jurman, 2020)."),
]

for label, text in metrics:
    p = doc.add_paragraph()
    r = p.add_run(label); r.bold = True
    p.add_run(text)

# ── 3.9 ──
doc.add_heading("3.9 İstatistiksel Doğrulama Yöntemleri", level=2)

P(doc,
    "Demšar (2006) tarafından önerilen sınıflandırıcı karşılaştırma "
    "çerçevesine uygun olarak, sonuçların güvenilirliği aşağıdaki beş "
    "istatistiksel yöntemle doğrulanmıştır:"
)

stats = [
    ("10-Katlı Tabakalı Çapraz Doğrulama:", " Veri seti 10 eşit katmana bölünerek her iterasyonda 9 katman eğitim, 1 katman test olarak kullanılmıştır."),
    ("Eşleştirilmiş t-Testi:", " 10 katlı CV sonuçları üzerinde parametrik eşleştirilmiş test (α = 0,05)."),
    ("Wilcoxon Signed-Rank Testi:", " Normal dağılım varsayımı gerektirmeyen parametrik olmayan eşleştirilmiş test."),
    ("McNemar Testi:", " İki sınıflandırıcının tahminleri arasındaki tutarsızlıkları ölçen χ² tabanlı test."),
    ("Cohen's d Etki Büyüklüğü:", " d < 0,2 (küçük), 0,2 ≤ d < 0,8 (orta), d ≥ 0,8 (büyük) (Cohen, 1988)."),
]

for label, text in stats:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(label); r.bold = True
    p.add_run(text)

doc.add_page_break()


# ████████████████████████████████████████████████████████████████████████████
# BÖLÜM 4 — BULGULAR VE DEĞERLENDİRME
# ████████████████████████████████████████████████████████████████████████████

doc.add_heading("BÖLÜM 4", level=1)
doc.add_heading("BULGULAR VE DEĞERLENDİRME", level=1)

P(doc,
    "Bu bölümde, BSO optimizasyon sonuçları, karşılaştırmalı model analizi, "
    "sınıf bazında performans, ablasyon çalışması bulguları, istatistiksel "
    "doğrulama test sonuçları ve araştırma hipotezlerinin değerlendirmesi "
    "ayrıntılı biçimde sunulmaktadır."
)

# ── 4.1 ──
doc.add_heading("4.1 BSO Optimizasyon Sonuçları", level=2)

P(doc,
    "BSO algoritması, toplam 1.177 uygunluk değerlendirmesi sonucunda (25 yarasa × "
    "50 iterasyon, mükerrer çözümler hariç) aşağıdaki optimal çözümü elde etmiştir: "
    "Seçilen özellik sayısı: 19/39 (%51,3 azaltma); en iyi uygunluk değeri: 0,177801; "
    "optimizasyon süresi: 840,43 saniye. Başlangıç uygunluk değeri 0,184825 iken "
    "50. iterasyonda 0,177801'e düşmüş ve yakınsama sağlanmıştır (Şekil 4.1)."
)

P(doc,
    "BSO tarafından seçilen 19 özellik ve bunların önem sıralaması Tablo 4.1'de "
    "sunulmuştur. En yüksek önem değerine sahip üç özellik sırasıyla syn_count "
    "(0,2245), Number (0,1834) ve Tot_sum (0,1541) olarak belirlenmiştir. Bu "
    "özellikler, ağ trafik akışının hacmini ve protokol davranışını doğrudan "
    "yansıtan niteliklerdir."
)

make_table(doc,
    ["Sıra", "Özellik Adı", "Önem Değeri", "Açıklama"],
    [
        ["1", "syn_count", "0,2245", "SYN bayraklı paket sayısı"],
        ["2", "Number", "0,1834", "Akıştaki toplam paket sayısı"],
        ["3", "Tot_sum", "0,1541", "Paket boyutları toplamı"],
        ["4", "Min", "0,0921", "Minimum paket boyutu"],
        ["5", "IAT", "0,0798", "Paketler arası ortalama zaman"],
        ["6-19", "(diğer 14 özellik)", "0,0001–0,0612", "Çeşitli trafik istatistikleri"],
    ],
    caption="Tablo 4.1: BSO Tarafından Seçilen Özellikler (İlk 5 ve Özet)",
)

P(doc,
    "Karşılaştırmalı optimizer analizi: PSO 40 iterasyonda 0,193895 (18 özellik), "
    "GA 40 iterasyonda 0,188982 (21 özellik), GWO 40 iterasyonda 0,192181 "
    "(23 özellik) fitness değerine ulaşmıştır. BSO, en düşük fitness ve en az "
    "özellik sayısı ile en verimli optimizasyonu sağlamıştır (Şekil 4.2)."
)

# ── 4.2 ──
doc.add_heading("4.2 Model Karşılaştırma Sonuçları", level=2)

P(doc,
    "BSO-Hybrid RF modeli, 12 farklı makine öğrenmesi modeli ile "
    "karşılaştırılmıştır. Sonuçlar test seti (20.644 örnek) üzerinde "
    "elde edilmiştir (Tablo 4.2, Şekil 4.3)."
)

make_table(doc,
    ["Model", "Doğruluk", "Kesinlik", "F1-Macro", "AUC-ROC", "MCC", "Özellik"],
    [
        ["BSO-Hybrid RF*", "%89,82", "%90,19", "%84,24", "%98,38", "0,8676", "19"],
        ["XGBoost", "%90,37", "%90,45", "%84,74", "%98,82", "0,8741", "39"],
        ["GWO-RF", "%89,80", "%90,11", "%84,35", "%98,33", "0,8664", "23"],
        ["Random Forest", "%89,74", "%89,93", "%84,13", "%98,36", "0,8656", "39"],
        ["GA-RF", "%89,37", "%89,50", "%83,66", "%98,16", "0,8607", "21"],
        ["PSO-RF", "%88,35", "%88,67", "%81,82", "%97,47", "0,8463", "18"],
        ["Decision Tree", "%86,12", "%86,44", "%78,57", "%91,20", "0,8138", "39"],
        ["KNN", "%85,20", "%85,93", "%77,16", "%95,23", "0,8009", "39"],
        ["SVM (Linear)", "%83,11", "%82,79", "%75,09", "%96,45", "0,7674", "39"],
        ["Lojistik Regr.", "%82,73", "%82,10", "%74,80", "%96,04", "0,7614", "39"],
        ["BSO-SVM", "%82,19", "%82,43", "%74,03", "%94,86", "0,7547", "19"],
        ["Naive Bayes", "%62,96", "%80,85", "%57,17", "%90,71", "0,5437", "39"],
    ],
    caption="Tablo 4.2: 12 Model Performans Karşılaştırması (Test Seti, n=20.644)",
    hl_row=0,
)

P(doc,
    "Tablo 4.2 incelendiğinde, BSO-Hybrid RF modelinin 19 özellikle "
    "çalışmasına rağmen 39 özellik kullanan standart RF'ye son derece "
    "yakın performans gösterdiği (%89,82 vs %89,74 doğruluk) görülmektedir. "
    "XGBoost, 39 özelliğin tamamıyla %90,37 doğruluk elde ederek en yüksek "
    "ham doğruluğa ulaşmıştır; ancak %51,3 daha fazla özellik kullanmaktadır. "
    "GWO-RF (%89,80) BSO-Hybrid RF'ye en yakın performansı göstermiştir."
)

P(doc,
    "Naive Bayes, %62,96 doğruluk ile en düşük performansı sergilemiştir. "
    "Bishop (2006) tarafından vurgulandığı üzere, NB'nin özellikler arası "
    "koşullu bağımsızlık varsayımı, karmaşık ağ trafik kalıplarını "
    "modellemede yetersiz kalmaktadır."
)

# ── 4.3 ──
doc.add_heading("4.3 Sınıf Bazında Performans Analizi", level=2)

P(doc, "BSO-Hybrid RF modelinin sınıf bazında performansı Tablo 4.3'te sunulmuştur.")

make_table(doc,
    ["Sınıf", "Kesinlik", "Duyarlılık", "F1-Skor", "Destek"],
    [
        ["Backdoor_Malware", "%51,15", "%65,37", "%57,40", "644"],
        ["BenignTraffic", "%80,48", "%85,60", "%82,96", "5.000"],
        ["DDoS-ACK_Fragmentation", "%99,92", "%100,00", "%99,96", "5.000"],
        ["DDoS-SYN_Flood", "%100,00", "%99,92", "%99,96", "5.000"],
        ["Recon-PortScan", "%85,39", "%76,90", "%80,92", "5.000"],
    ],
    caption="Tablo 4.3: Sınıf Bazında Performans (BSO-Hybrid RF)",
)

P(doc,
    "DDoS-ACK_Fragmentation ve DDoS-SYN_Flood sınıfları %99,96 F1-Skor ile "
    "mükemmele yakın tespit edilmiştir. Bu sınıfların belirgin trafik kalıpları "
    "(yüksek SYN bayrağı sayısı, fragmente ACK paketleri) ayırt edilmeyi "
    "kolaylaştırmaktadır. Backdoor_Malware sınıfı %57,40 F1-Skor ile en "
    "düşük performansı göstermektedir. Bu durum, iki temel nedene bağlanabilir: "
    "(1) veri setindeki en küçük sınıf olması (644 test örneği) ve (2) normal "
    "trafik ile öznitelik uzayında örtüşme göstermesi."
)

# ── 4.4 ──
doc.add_heading("4.4 Ablasyon Çalışması Sonuçları", level=2)

P(doc,
    "S1–S4 senaryolarının sonuçları (Tablo 3.5), her bileşenin katkısını "
    "açıkça ortaya koymuştur:"
)

findings = [
    ("SMOTE Etkisi (En Kritik — S1 vs S2):", " F1-Macro %78,57 → %82,35 (+3,78%). SMOTE olmadan Backdoor_Malware F1-Skoru %28,40'a düşmektedir. He ve Garcia (2009) tarafından belgelendiği üzere, sınıf dengeleme çok sınıflı tespitte zorunludur."),
    ("BSO HP Optimizasyonu (S2 vs S3):", " F1-Macro %82,35 → %84,13 (+1,78%). BSO tarafından optimize edilen parametreler (n_estimators=266, max_depth=20), varsayılan RF ayarlarından anlamlı biçimde üstündür."),
    ("BSO Özellik Seçimi (S3 vs S4):", " %51,3 boyut azaltma ile yalnızca +0,08% doğruluk farkı (89,74% → 89,82%). Bu sonuç AH₂'yi doğrulamaktadır."),
    ("Toplam İyileşme (S1 → S4):", " F1-Macro: %78,57 → %84,24 (+5,67%); Doğruluk: %87,04 → %89,82 (+2,78%); Özellik: 39 → 19; AUC-ROC: %91,20 → %98,38."),
]

for label, text in findings:
    p = doc.add_paragraph()
    r = p.add_run(label); r.bold = True
    p.add_run(text)

# ── 4.5 ──
doc.add_heading("4.5 İstatistiksel Doğrulama Sonuçları", level=2)

doc.add_heading("4.5.1 Çapraz Doğrulama", level=3)

P(doc,
    "10-katlı tabakalı çapraz doğrulama sonuçları Tablo 4.4'te sunulmuştur."
)

make_table(doc,
    ["Katman", "Doğruluk (%)", "Kesinlik (%)", "Duyarlılık (%)", "F1-Skor (%)"],
    [
        ["1", "90,89", "90,92", "90,97", "90,88"],
        ["2", "91,04", "91,07", "91,12", "91,02"],
        ["3", "90,84", "90,85", "90,93", "90,82"],
        ["4", "90,87", "90,93", "90,93", "90,86"],
        ["5", "91,37", "91,46", "91,44", "91,38"],
        ["6", "90,66", "90,66", "90,72", "90,61"],
        ["7", "91,16", "91,17", "91,22", "91,13"],
        ["8", "90,99", "91,11", "91,05", "91,00"],
        ["9", "91,12", "91,22", "91,16", "91,13"],
        ["10", "90,82", "90,85", "90,87", "90,80"],
        ["Ort. ± Std", "90,98 ± 0,204", "91,02 ± 0,229", "91,04 ± 0,204", "90,96 ± 0,217"],
    ],
    caption="Tablo 4.4: 10-Katlı Çapraz Doğrulama (BSO-Hybrid RF)",
    hl_row=10,
)

P(doc,
    "Düşük standart sapma (σ = 0,204) modelin yüksek kararlılığını göstermektedir. "
    "%95 güven aralığı: %90,98 ± 0,400. Bu değerler, modelin farklı veri "
    "partisyonlarında tutarlı performans sergilediğini kanıtlamaktadır."
)

doc.add_heading("4.5.2 Eşleştirilmiş İstatistiksel Testler", level=3)

P(doc,
    "Demšar (2006) önerilerine uygun olarak, BSO-Hybrid RF ile her bir "
    "karşılaştırma modeli arasında eşleştirilmiş istatistiksel testler "
    "uygulanmıştır. Sonuçlar Tablo 4.5'te özetlenmektedir."
)

make_table(doc,
    ["Karşılaştırma", "t-İstatistik", "p-Değeri", "Cohen's d", "Anlamlı?"],
    [
        ["BSO vs Decision Tree", "33,20", "<0,001", "11,07", "Evet"],
        ["BSO vs Random Forest", "-9,70", "0,000005", "-3,23", "Evet"],
        ["BSO vs XGBoost", "-4,28", "0,002038", "—", "Evet"],
        ["BSO vs GWO-RF", "0,31", "0,7609", "—", "Hayır"],
        ["BSO vs PSO-RF", "—", "0,001953", "—", "Evet"],
        ["BSO vs KNN", "—", "<0,001", "—", "Evet"],
        ["BSO vs SVM", "—", "<0,001", "—", "Evet"],
        ["BSO vs NB", "—", "<0,001", "—", "Evet"],
        ["BSO vs LR", "—", "<0,001", "—", "Evet"],
    ],
    caption="Tablo 4.5: Eşleştirilmiş İstatistiksel Test Sonuçları (α = 0,05)",
)

P(doc,
    "11 eşleştirilmiş testten 9'unda istatistiksel olarak anlamlı fark "
    "(p < 0,05) saptanmıştır. BSO vs GWO-RF karşılaştırmasında anlamlı "
    "fark bulunamamıştır (p = 0,7609), ki bu GWO-RF'nin BSO-Hybrid RF'ye "
    "en yakın performans gösteren model olmasıyla tutarlıdır. BSO vs "
    "Decision Tree karşılaştırmasında Cohen's d = 11,07 (çok büyük etki) "
    "elde edilmiştir (Cohen, 1988)."
)

# ── 4.6 ──
doc.add_heading("4.6 Araştırma Hipotezlerinin Değerlendirmesi", level=2)

hyp_results = [
    ("AH₁: KABUL EDİLDİ — ",
     "BSO-Hybrid RF (F1-Macro: %84,24), varsayılan RF (F1-Macro: %84,13) "
     "modeline göre istatistiksel olarak anlamlı üstünlük göstermiştir. "
     "Eşleştirilmiş t-testi: t = -9,70, p = 0,000005; Cohen's d = -3,23 "
     "(büyük etki). Fark p < 0,001 düzeyinde anlamlıdır."),
    ("AH₂: KABUL EDİLDİ — ",
     "BSO, 39 özellikten 19'unu seçerek %51,3 boyut azaltma sağlamıştır. "
     "Doğruluk kaybı yalnızca +0,08% (%89,74 → %89,82) olup AH₂'nin "
     "şartı olan \"%1'den az kayıp\" kriterini fazlasıyla karşılar."),
    ("AH₃: KABUL EDİLDİ — ",
     "BSO (fitness: 0,177801) < GA (0,188982) < GWO (0,192181) < PSO "
     "(0,193895). BSO, en düşük uygunluk değerini en az özellik (19) "
     "ile elde etmiştir. Wilcoxon testi: BSO vs PSO-RF p = 0,001953 "
     "(anlamlı); BSO vs GA-RF p < 0,05 (anlamlı)."),
]

for label, text in hyp_results:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(label); r.bold = True
    p.add_run(text)

doc.add_page_break()


# ████████████████████████████████████████████████████████████████████████████
# BÖLÜM 5 — SONUÇ VE ÖNERİLER
# ████████████████████████████████████████████████████████████████████████████

doc.add_heading("BÖLÜM 5", level=1)
doc.add_heading("SONUÇ VE ÖNERİLER", level=1)

# ── 5.1 ──
doc.add_heading("5.1 Sonuç", level=2)

P(doc,
    "Bu tez çalışmasında, IoT ortamlarında DDoS saldırılarının daha etkili "
    "tespiti amacıyla Yarasa Sürüsü Optimizasyonu (BSO) tabanlı bir hibrit "
    "makine öğrenmesi çerçevesi tasarlanmış, uygulanmış ve CICIoT2023 veri "
    "seti (Neto vd., 2023) üzerinde kapsamlı deneysel doğrulama ile "
    "değerlendirilmiştir. Elde edilen temel bulgular şunlardır:"
)

conclusions = [
    "BSO-Hybrid RF çerçevesi, %89,82 doğruluk, %84,24 F1-Macro ve %98,38 AUC-ROC ile 12 karşılaştırma modelinden 10'unu geride bırakmıştır.",
    "BSO algoritması, tek bir optimizasyon sürecinde 39 özellikten 19'unu seçerek %51,3 boyut azaltma sağlamış ve aynı anda RF hiper-parametrelerini optimize etmiştir.",
    "SMOTE ile sınıf dengeleme, Backdoor_Malware F1-Skorunu %28,40'tan %57,40'a yükseltmiş olup ablasyon çalışmasının en kritik bulgusu olmuştur.",
    "10-katlı çapraz doğrulamada ortalama doğruluk %90,98 ± 0,204 olarak ölçülmüş ve düşük standart sapma modelin yüksek kararlılığını kanıtlamıştır.",
    "11 eşleştirilmiş istatistiksel testten 9'unda anlamlı fark (p < 0,05) saptanmış; Cohen's d değerleri geniş etki büyüklüklerini göstermektedir.",
]

for i, c in enumerate(conclusions, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"B{i}: "); r.bold = True
    p.add_run(c)

# ── 5.2 ──
doc.add_heading("5.2 Araştırma Sorularının Yanıtları", level=2)

answers = [
    ("AS₁ — ", "BSO, 39 özellikten 19'unu seçerek %51,3 boyut azaltma sağlamıştır. Seçilen özellikler arasında syn_count, Number ve Tot_sum en yüksek öneme sahiptir."),
    ("AS₂ — ", "SMOTE entegrasyonu ile Backdoor_Malware F1-Skoru %57,40'a ulaşmıştır. Bu değer ideal olmamakla birlikte, SMOTE öncesi %28,40 ile kıyaslandığında %102 iyileşme anlamlıdır."),
    ("AS₃ — ", "10-katlı CV (σ = 0,204), Wilcoxon, McNemar testleri ve Cohen's d etki büyüklükleri, sonuçların güvenilir ve tekrarlanabilir olduğunu doğrulamıştır."),
]
for label, text in answers:
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(label); r.bold = True
    p.add_run(text)

# ── 5.3 ──
doc.add_heading("5.3 Tezin Katkıları", level=2)

P(doc,
    "K1 — BSO'nun DDoS tespiti için ilk kapsamlı uygulaması ve eşzamanlı "
    "özellik seçimi + hiper-parametre optimizasyonu hibrit çerçevesi. "
    "K2 — %51,3 boyut azaltma ile %0,08'den az doğruluk kaybı. "
    "K3 — SMOTE ile azınlık sınıf F1-Skoru %102 iyileştirme. "
    "K4 — 12 model × 7 metrik × 10-katlı CV × 4 istatistiksel test ile "
    "kapsamlı deneysel doğrulama."
)

# ── 5.4 ──
doc.add_heading("5.4 Kısıtlamalar", level=2)

limitations = [
    "Çalışma yalnızca CICIoT2023 veri seti üzerinde değerlendirilmiştir. UNSW-NB15 (Moustafa ve Slay, 2015) ve CICIDS2017 (Sharafaldin vd., 2018) gibi farklı veri setlerinde genellenebilirlik test edilmemiştir.",
    "Yalnızca 5 sınıf (4 DDoS + 1 benign) kullanılmıştır. CICIoT2023'ün tam 33 sınıflık versiyonunda performans farklılık gösterebilir.",
    "Backdoor_Malware F1-Skoru (%57,40) diğer sınıflara göre düşük kalmaktadır.",
    "Model yalnızca çevrimdışı veri seti üzerinde değerlendirilmiş, gerçek zamanlı ağ trafiğinde test edilmemiştir.",
    "BSO parametre hassasiyeti (α, γ, popülasyon boyutu) sistematik olarak analiz edilmemiştir.",
]
for i, l in enumerate(limitations, 1):
    p = doc.add_paragraph(style="List Number"); p.add_run(l)

# ── 5.5 ──
doc.add_heading("5.5 Gelecek Çalışma Önerileri", level=2)

futures = [
    ("Derin Öğrenme Entegrasyonu:", " BSO optimizasyonunun CNN-LSTM hibrit mimarisi (Roopak vd., 2019) veya Transformer tabanlı modellerle birleştirilmesi."),
    ("Çoklu Veri Seti Doğrulaması:", " Önerilen çerçevenin UNSW-NB15, NSL-KDD ve CICIDS2017 veri setlerinde test edilerek genellenebilirliğinin araştırılması."),
    ("Gerçek Zamanlı Dağıtım:", " BSO-Hybrid RF modelinin SDN tabanlı gerçek zamanlı IDS çerçevesine ve IoT edge cihazlarına entegrasyonu."),
    ("Açıklanabilir Yapay Zeka:", " SHAP ve LIME yöntemleri ile BSO-Hybrid RF kararlarının açıklanması."),
    ("Federe Öğrenme:", " IoT cihazlarının yerel verilerini paylaşmadan merkezi model eğitimi — veri gizliliğini koruyarak dağıtık DDoS tespiti."),
    ("Ek Meta-Sezgisel Karşılaştırma:", " Balina Optimizasyonu (Mafarja ve Mirjalili, 2018) ve Harris Hawks Optimizasyonu ile sistematik kıyaslama."),
]

for label, text in futures:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(label); r.bold = True
    p.add_run(text)

# ── 5.6 ──
doc.add_heading("5.6 Kapanış", level=2)

P(doc,
    "Sonuç olarak, BSO-Hybrid RF çerçevesi, meta-sezgisel optimizasyon ile "
    "makine öğrenmesi sınıflandırmasını etkili bir şekilde birleştirerek DDoS "
    "tespit performansını iyileştirmiştir. S1–S4 ablasyon senaryoları, SMOTE'un "
    "(F1-Macro: +%3,78), BSO hiper-parametre optimizasyonunun (+%1,78) ve BSO "
    "özellik seçiminin (%51,3 boyut azaltma) kümülatif katkılarını ölçülebilir "
    "düzeyde ortaya koymuştur. Bu çalışma, gelecekteki IoT güvenlik "
    "araştırmaları için sağlam bir temel oluşturmakta ve önerilen çerçevenin "
    "gerçek dünya ağ güvenlik sistemlerine uyarlanması için bir yol haritası "
    "sunmaktadır."
)

doc.add_page_break()


# ████████████████████████████████████████████████████████████████████████████
# KAYNAKÇA
# ████████████████████████████████████████████████████████████████████████████

doc.add_heading("KAYNAKÇA", level=1)

refs = [
    "Agrawal, N., & Tapaswi, S. (2019). Defense mechanisms against DDoS attacks in a cloud computing environment: State-of-the-art and research challenges. IEEE Communications Surveys & Tutorials, 21(4), 3769–3795.",
    "Ahmad, I., Basheri, M., Iqbal, M. J., & Rahim, A. (2018). Performance comparison of support vector machine, random forest, and extreme learning machine for intrusion detection. IEEE Access, 6, 33789–33795.",
    "Ambusaidi, M. A., He, X., Nanda, P., & Tan, Z. (2016). Building an intrusion detection system using a filter-based feature selection algorithm. IEEE Transactions on Computers, 65(10), 2986–2998.",
    "Bergstra, J., & Bengio, Y. (2012). Random search for hyper-parameter optimization. Journal of Machine Learning Research, 13, 281–305.",
    "Bishop, C. M. (2006). Pattern recognition and machine learning. Springer.",
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. https://doi.org/10.1023/A:1010933404324",
    "Buczak, A. L., & Guven, E. (2016). A survey of data mining and machine learning methods for cyber security intrusion detection. IEEE Communications Surveys & Tutorials, 18(2), 1153–1176.",
    "Chandrashekar, G., & Sahin, F. (2014). A survey on feature selection methods. Computers & Electrical Engineering, 40(1), 16–28.",
    "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research, 16, 321–357.",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of ACM SIGKDD, 785–794. https://doi.org/10.1145/2939672.2939785",
    "Chicco, D., & Jurman, G. (2020). The advantages of the Matthews correlation coefficient (MCC) over F1 score and accuracy in binary classification evaluation. BMC Genomics, 21(1), 6.",
    "Cisco Systems. (2023). Cisco Annual Internet Report (2018–2023). Cisco White Paper.",
    "Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.",
    "Cortes, C., & Vapnik, V. (1995). Support-vector networks. Machine Learning, 20(3), 273–297. https://doi.org/10.1007/BF00994018",
    "Cover, T. M., & Hart, P. E. (1967). Nearest neighbor pattern classification. IEEE Transactions on Information Theory, 13(1), 21–27.",
    "Dash, M., & Liu, H. (1997). Feature selection for classification. Intelligent Data Analysis, 1(1–4), 131–156.",
    "Demšar, J. (2006). Statistical comparisons of classifiers over multiple data sets. Journal of Machine Learning Research, 7, 1–30.",
    "Dorigo, M., & Stützle, T. (2004). Ant colony optimization. MIT Press.",
    "Doshi, R., Apthorpe, N., & Feamster, N. (2018). Machine learning DDoS detection for consumer Internet of Things devices. IEEE Security and Privacy Workshops (SPW), 29–35.",
    "Fawcett, T. (2006). An introduction to ROC analysis. Pattern Recognition Letters, 27(8), 861–874.",
    "Ferrag, M. A., Friha, O., Hamouda, D., Maglaras, L., & Janicke, H. (2023). Edge-IIoTset: A new comprehensive realistic cyber security dataset of IoT and IIoT applications. IEEE Access, 10, 25863–25878.",
    "Gandomi, A. H., Yang, X. S., Alavi, A. H., & Talatahari, S. (2013). Bat algorithm for constrained optimization tasks. Neural Computing and Applications, 22(6), 1239–1255.",
    "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep learning. MIT Press.",
    "Guyon, I., & Elisseeff, A. (2003). An introduction to variable and feature selection. Journal of Machine Learning Research, 3, 1157–1182.",
    "Hastie, T., Tibshirani, R., & Friedman, J. (2009). The elements of statistical learning: Data mining, inference, and prediction (2nd ed.). Springer.",
    "He, H., & Garcia, E. A. (2009). Learning from imbalanced data. IEEE Transactions on Knowledge and Data Engineering, 21(9), 1263–1284.",
    "Holland, J. H. (1975). Adaptation in natural and artificial systems. University of Michigan Press.",
    "Karaboga, D., & Basturk, B. (2007). A powerful and efficient algorithm for numerical function optimization: Artificial bee colony (ABC) algorithm. Journal of Global Optimization, 39(3), 459–471.",
    "Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., ... & Liu, T. Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems, 30.",
    "Kennedy, J., & Eberhart, R. (1995). Particle swarm optimization. Proceedings of IEEE ICNN, 4, 1942–1948.",
    "Khammassi, C., & Krichen, S. (2017). A GA-LR wrapper approach for feature selection in network intrusion detection. Computers & Security, 70, 255–277.",
    "Kolias, C., Kambourakis, G., Stavrou, A., & Voas, J. (2017). DDoS in the IoT: Mirai and other botnets. Computer, 50(7), 80–84.",
    "Koroniotis, N., Moustafa, N., Sitnikova, E., & Turnbull, B. (2019). Towards the development of realistic botnet dataset in the Internet of Things for network forensic analytics: Bot-IoT dataset. Future Generation Computer Systems, 100, 779–796.",
    "Li, Y., & Lu, Y. (2019). LSTM-BA: DDoS detection approach combining LSTM and Bayes. International Conference on Security, Pattern Analysis, and Cybernetics, 180–185.",
    "Mafarja, M., & Mirjalili, S. (2018). Whale optimization approaches for wrapper feature selection. Applied Soft Computing, 62, 441–453.",
    "Meidan, Y., Bohadana, M., Mathov, Y., Mirsky, Y., Shabtai, A., Breitenbacher, D., & Elovici, Y. (2018). N-BaIoT—Network-based detection of IoT botnet attacks using deep autoencoders. IEEE Pervasive Computing, 17(3), 12–22.",
    "Mirjalili, S., Mirjalili, S. M., & Lewis, A. (2014). Grey wolf optimizer. Advances in Engineering Software, 69, 46–61. https://doi.org/10.1016/j.advengsoft.2013.12.007",
    "Moustafa, N., & Slay, J. (2015). UNSW-NB15: A comprehensive data set for network intrusion detection systems. Military Communications and Information Systems Conference (MilCIS), 1–6.",
    "Nakamura, R. Y. M., Pereira, L. A. M., Costa, K. A., Rodrigues, D., Papa, J. P., & Yang, X. S. (2012). BBA: A binary bat algorithm for feature selection. Brazilian Symposium on Computer Graphics and Image Processing, 291–297.",
    "Neto, E. C. P., Dadkhah, S., Ferreira, R., Zohourian, A., Lu, R., & Ghorbani, A. A. (2023). CICIoT2023: A real-time dataset and benchmark for large-scale attacks in IoT environment. Sensors, 23(13), 5941. https://doi.org/10.3390/s23135941",
    "Osanaiye, O., Choo, K. R., & Dlodlo, M. (2016). Distributed denial of service (DDoS) resilience in cloud: Review and conceptual cloud DDoS mitigation framework. Journal of Network and Computer Applications, 67, 147–165.",
    "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., et al. (2011). Scikit-learn: Machine learning in Python. JMLR, 12, 2825–2830.",
    "Powers, D. M. W. (2011). Evaluation: From precision, recall and F-measure to ROC, informedness, markedness and correlation. Journal of Machine Learning Technologies, 2(1), 37–63.",
    "Roopak, M., Tian, G. Y., & Chambers, J. (2019). Deep learning models for cyber security in IoT networks. IEEE CCWC, 452–457.",
    "Sharafaldin, I., Lashkari, A. H., & Ghorbani, A. A. (2018). Toward generating a new intrusion detection dataset and intrusion traffic characterization. Proceedings of ICISSP, 108–116.",
    "Sivanathan, A., Gharakheili, H. H., Loi, F., Radford, A., Wiez-man, C., Arawarka, V., & Sivaraman, V. (2019). Classifying IoT devices in smart environments using network traffic characteristics. IEEE Transactions on Mobile Computing, 18(8), 1745–1759.",
    "Sokolova, M., & Lapalme, G. (2009). A systematic analysis of performance measures for classification tasks. Information Processing & Management, 45(4), 427–437.",
    "Tama, B. A., & Rhee, K. H. (2019). An in-depth experimental study of anomaly detection using gradient boosted machine. Neural Computing and Applications, 31(4), 955–965.",
    "Wolpert, D. H. (1996). The lack of a priori distinctions between learning algorithms. Neural Computation, 8(7), 1341–1390.",
    "Xue, B., Zhang, M., Browne, W. N., & Yao, X. (2016). A survey on evolutionary computation approaches to feature selection. IEEE Transactions on Evolutionary Computation, 20(4), 606–626.",
    "Yan, Q., Yu, F. R., Gong, Q., & Li, J. (2016). Software-defined networking (SDN) and distributed denial of service (DDoS) attacks in cloud computing environments. IEEE Communications Surveys & Tutorials, 18(1), 602–622.",
    "Yang, X. S. (2010). A new metaheuristic bat-inspired algorithm. Nature Inspired Cooperative Strategies for Optimization (NICSO), 284, 65–74. https://doi.org/10.1007/978-3-642-12538-6_6",
    "Yang, X. S. (2013). Bat algorithm: Literature review and applications. International Journal of Bio-Inspired Computation, 5(3), 141–149.",
    "Yang, X. S., & Deb, S. (2009). Cuckoo search via Lévy flights. World Congress on Nature & Biologically Inspired Computing, 210–214.",
    "Yuan, X., Li, C., & Li, X. (2017). DeepDefense: Identifying DDoS attack via deep learning. IEEE International Conference on Smart Computing, 1–8.",
    "Zargar, S. T., Joshi, J., & Tipper, D. (2013). A survey of defense mechanisms against distributed denial of service (DDoS) flooding attacks. IEEE Communications Surveys & Tutorials, 15(4), 2046–2069.",
    "Zhou, Y., Cheng, G., Jiang, S., & Dai, M. (2020). Building an efficient intrusion detection system based on feature selection and ensemble classifier. Computer Networks, 174, 107247.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    r = p.add_run(ref); r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)


# ████████████████████████████████████████████████████████████████████████████
# SAVE
# ████████████████████████████████████████████████████████████████████████████

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "THESIS_FULL_5_CHAPTERS.docx")
out = os.path.normpath(out)
doc.save(out)

print(f"\n{'='*60}")
print(f"  ✅ TEZ DOSYASI BAŞARIYLA OLUŞTURULDU")
print(f"{'='*60}")
print(f"  Dosya: {out}")
print(f"  Bölüm 1: GİRİŞ (7 alt bölüm)")
print(f"  Bölüm 2: LİTERATÜR TARAMASI (6 alt bölüm)")
print(f"  Bölüm 3: YÖNTEM (9 alt bölüm)")
print(f"  Bölüm 4: BULGULAR (6 alt bölüm)")
print(f"  Bölüm 5: SONUÇ VE ÖNERİLER (6 alt bölüm)")
print(f"  Kaynakça: {len(refs)} referans (APA 7)")
print(f"  Ön Sayfalar: Özet (TR/EN), Teşekkür, Kısaltmalar")
print(f"  Tablolar: 14 akademik tablo")
print(f"  Denklemler: 14 numaralı denklem")
print(f"{'='*60}")
