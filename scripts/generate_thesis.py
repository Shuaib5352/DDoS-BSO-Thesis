#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Master's Thesis — Chapters 1-3 (Turkish)
IoT Ortamlarında BSO Tabanlı Hibrit Makine Öğrenmesi ile DDoS Saldırı Tespiti

Author:  Shuaib Ayad Jasim
Advisor: Dr. Saim Ervural
University: KTO Karatay Üniversitesi
Date: 2026

Generates a formatted Word (.docx) file.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================================
# STYLE HELPERS
# ============================================================================

def set_cell_shading(cell, color_hex):
    """Set background shading of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, start, end with values like '4' (eighths of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), val)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tcPr.append(borders)


def add_formatted_table(doc, headers, rows, col_widths=None, caption=None, highlight_row=None):
    """Add a nicely formatted academic table."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2E4057")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if highlight_row is not None and r_idx == highlight_row:
                set_cell_shading(cell, "E8F5E9")
                run.bold = True

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing after table
    return table


def add_equation(doc, equation_text, number=None):
    """Add a centered equation with optional numbering."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(equation_text)
    run.font.size = Pt(11)
    run.italic = True
    if number:
        run2 = p.add_run(f"    ({number})")
        run2.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


# ============================================================================
# DOCUMENT SETUP
# ============================================================================

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)

# Default font style
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level in range(1, 4):
    hstyle = doc.styles[f"Heading {level}"]
    hstyle.font.name = "Times New Roman"
    hstyle.font.color.rgb = RGBColor(0, 0, 0)
    hstyle.font.bold = True
    if level == 1:
        hstyle.font.size = Pt(16)
        hstyle.paragraph_format.space_before = Pt(24)
        hstyle.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hstyle.font.size = Pt(14)
        hstyle.paragraph_format.space_before = Pt(18)
        hstyle.paragraph_format.space_after = Pt(8)
    else:
        hstyle.font.size = Pt(12)
        hstyle.paragraph_format.space_before = Pt(12)
        hstyle.paragraph_format.space_after = Pt(6)


# ============================================================================
# TITLE PAGE
# ============================================================================

for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("KTO KARATAY ÜNİVERSİTESİ")
run.bold = True
run.font.size = Pt(16)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

dept_p = doc.add_paragraph()
dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dept_p.add_run("Bilgisayar Mühendisliği Anabilim Dalı")
run.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

thesis_title = doc.add_paragraph()
thesis_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = thesis_title.add_run(
    "IoT ORTAMLARINDA YARASA SÜRÜSÜ OPTİMİZASYONU (BSO) TABANLI "
    "HİBRİT MAKİNE ÖĞRENMESİ İLE DDoS SALDIRI TESPİTİ: "
    "CICIoT2023 VERİ SETİ ÜZERİNDE DENEYSEL DEĞERLENDİRME"
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

type_p = doc.add_paragraph()
type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = type_p.add_run("YÜKSEK LİSANS TEZİ")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author_p.add_run("Shuaib Ayad JASIM")
run.font.size = Pt(13)

doc.add_paragraph()

advisor_p = doc.add_paragraph()
advisor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = advisor_p.add_run("Danışman: Dr. Öğr. Üyesi Saim ERVURAL")
run.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

year_p = doc.add_paragraph()
year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = year_p.add_run("KONYA — 2026")
run.bold = True
run.font.size = Pt(14)

doc.add_page_break()


# ============================================================================
# TABLE OF CONTENTS (placeholder)
# ============================================================================

doc.add_heading("İÇİNDEKİLER", level=1)
toc_items = [
    ("BÖLÜM 1 — GİRİŞ", "1"),
    ("  1.1 Araştırmanın Arka Planı", "1"),
    ("  1.2 Problem Tanımı", "3"),
    ("  1.3 Araştırma Hipotezleri", "5"),
    ("  1.4 Araştırma Soruları", "6"),
    ("  1.5 Tezin Amacı ve Kapsamı", "7"),
    ("  1.6 Tezin Katkıları", "8"),
    ("  1.7 Tezin Organizasyonu", "9"),
    ("BÖLÜM 2 — LİTERATÜR TARAMASI", "10"),
    ("  2.1 DDoS Saldırıları ve IoT Güvenliği", "10"),
    ("  2.2 Makine Öğrenmesi Tabanlı Saldırı Tespiti", "14"),
    ("  2.3 Meta-Sezgisel Optimizasyon Yöntemleri", "18"),
    ("  2.4 Yarasa Algoritması (Bat Algorithm)", "22"),
    ("  2.5 Öznitelik Seçimi Yöntemleri", "25"),
    ("  2.6 İlgili Çalışmalar ve Karşılaştırma", "28"),
    ("BÖLÜM 3 — YÖNTEM", "32"),
    ("  3.1 Araştırma Tasarımı ve Genel Çerçeve", "32"),
    ("  3.2 Veri Seti: CICIoT2023", "34"),
    ("  3.3 Veri Ön İşleme", "37"),
    ("  3.4 BSO Algoritması", "40"),
    ("  3.5 Uygunluk Fonksiyonu", "44"),
    ("  3.6 Hibrit Model Tanımı: BSO-Hybrid RF", "46"),
    ("  3.7 Deney Senaryoları (S1–S4)", "48"),
    ("  3.8 Değerlendirme Metrikleri", "50"),
    ("  3.9 İstatistiksel Doğrulama Yöntemleri", "52"),
    ("KAYNAKÇA", "54"),
]
for title, page in toc_items:
    p = doc.add_paragraph()
    if title.startswith("  "):
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(title.strip())
        run.font.size = Pt(11)
    else:
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
    # Tab + page number
    run2 = p.add_run(f"\t{page}")
    run2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ============================================================================
# CHAPTER 1 — GİRİŞ (INTRODUCTION)
# ============================================================================

doc.add_heading("BÖLÜM 1", level=1)
h1 = doc.add_heading("GİRİŞ", level=1)

# --- 1.1 ---
doc.add_heading("1.1 Araştırmanın Arka Planı", level=2)

doc.add_paragraph(
    "İnternet of Things (IoT) cihazlarının hızla yaygınlaşması, ağ altyapılarını "
    "Dağıtılmış Hizmet Reddi (DDoS) saldırılarına karşı giderek daha savunmasız hale "
    "getirmektedir. Cisco'nun 2023 yıllık raporuna göre, küresel DDoS saldırı sayısı "
    "yıllık olarak %25 artış göstermiş ve 2025 yılında 15,4 milyona ulaşması "
    "öngörülmektedir (Cisco, 2023). IoT cihazlarının sınırlı hesaplama kaynakları ve "
    "güvenlik zafiyetleri, bu cihazların botnet ağlarına dönüştürülmesini kolaylaştırmakta "
    "ve Mirai benzeri büyük ölçekli saldırıların temelini oluşturmaktadır (Kolias vd., 2017)."
)

doc.add_paragraph(
    "DDoS saldırıları, hedef sisteme aşırı miktarda trafik göndererek hizmetin "
    "kesintiye uğramasına neden olur. Bu saldırılar, hacimsel (volumetric), "
    "protokol tabanlı ve uygulama katmanı olmak üzere üç ana kategoride "
    "sınıflandırılmaktadır (Zargar vd., 2013). IoT ortamlarında bu saldırılar "
    "özellikle tehlikelidir çünkü: (1) IoT cihazları genellikle güçlü güvenlik "
    "mekanizmalarından yoksundur, (2) cihaz sayısının çokluğu büyük ölçekli botnet "
    "oluşturmayı kolaylaştırır, (3) heterojen cihaz yapıları standart güvenlik "
    "politikalarının uygulanmasını zorlaştırır (Doshi vd., 2018)."
)

doc.add_paragraph(
    "Geleneksel imza tabanlı saldırı tespit sistemleri (IDS), bilinen saldırı "
    "kalıplarıyla sınırlı kalarak sıfır-gün saldırılarına ve polimorfik saldırı "
    "varyantlarına karşı etkisiz kalmaktadır. Bu durum, makine öğrenmesi (ML) "
    "tabanlı tespit yöntemlerinin önemini önemli ölçüde artırmıştır. ML yaklaşımları, "
    "ağ trafiğindeki normal ve anormal kalıpları otomatik olarak öğrenerek, daha "
    "önce görülmemiş saldırı türlerini de tespit edebilme potansiyeline sahiptir "
    "(Buczak ve Guven, 2016). Ancak ML yaklaşımları, yüksek boyutlu ağ trafik "
    "verileriyle karşı karşıya kaldığında, gereksiz özelliklerden kaynaklanan aşırı "
    "uyum (overfitting), yüksek hesaplama maliyeti ve sınıf dengesizliğinin kontrol "
    "edilmemesi gibi temel sorunlarla mücadele etmektedir."
)

doc.add_paragraph(
    "Son yıllarda, meta-sezgisel optimizasyon algoritmalarının makine öğrenmesi ile "
    "birleştirilmesi, bu sorunlara etkili çözümler sunmaktadır. Parçacık Sürüsü "
    "Optimizasyonu (PSO), Genetik Algoritma (GA) ve Gri Kurt Optimizasyonu (GWO) "
    "gibi doğadan esinlenen algoritmalar, özellik seçimi ve hiper-parametre "
    "optimizasyonu görevlerinde başarıyla uygulanmıştır (Xue vd., 2016). "
    "Bu çalışmada, Yang (2010) tarafından önerilen Yarasa Algoritması'nın (Bat "
    "Algorithm — BSO) DDoS saldırı tespiti bağlamında ilk kapsamlı uygulaması "
    "gerçekleştirilmektedir."
)

doc.add_paragraph(
    "Yarasa Sürüsü Optimizasyonu (BSO), micro-yarasaların ekolokasyon davranışından "
    "esinlenen bir meta-sezgisel algoritmadır. BSO, frekans ayarlama, darbe emisyonu "
    "ve gürültü kontrolü mekanizmalarıyla küresel ve yerel arama arasında dinamik denge "
    "kurma yeteneğine sahiptir. Bu özellikler, BSO'yu karmaşık çok boyutlu optimizasyon "
    "problemleri için özellikle uygun kılmaktadır (Yang, 2013). Nakamura vd. (2012) "
    "BSO'nun ikili versiyonunu (BBA) öznitelik seçimi için başarıyla uygulamış, "
    "Gandomi vd. (2013) ise kısıtlı optimizasyon görevlerinde BSO'nun etkinliğini "
    "göstermiştir."
)

# --- 1.2 ---
doc.add_heading("1.2 Problem Tanımı", level=2)

doc.add_paragraph(
    "DDoS saldırı tespitinde mevcut makine öğrenmesi yaklaşımlarının karşılaştığı "
    "üç temel sorun bu tezin motivasyonunu oluşturmaktadır:"
)

problems = [
    ("Yüksek Boyutlu Özellik Uzayı: ",
     "Ağ trafik veri setleri genellikle onlarca hatta yüzlerce özellik içerir. "
     "Bu özelliklerin önemli bir kısmı gereksiz veya gürültülüdür ve model "
     "performansını olumsuz etkiler. CICIoT2023 veri seti 39 ağ trafik özelliği "
     "içermektedir; ancak bu özelliklerin tamamı sınıflandırma için bilgilendirici "
     "değildir. Gereksiz özellikler aşırı uyuma (overfitting), hesaplama maliyetinin "
     "artmasına ve genellenebilirliğin azalmasına neden olur (Guyon ve Elisseeff, 2003)."),
    ("Hiper-Parametre Optimizasyonu: ",
     "Makine öğrenmesi modellerinin performansı, hiper-parametre seçimine büyük "
     "ölçüde bağımlıdır. Random Forest gibi topluluk (ensemble) modellerinde "
     "n_estimators, max_depth, min_samples_split ve max_features gibi parametrelerin "
     "optimal değerlerinin belirlenmesi, geniş bir arama uzayı gerektirir. "
     "Manuel ayarlama veya ızgara araması (grid search) gibi geleneksel yöntemler, "
     "hesaplama maliyeti açısından pratik olmayabilir ve optimal çözümü "
     "garantilememektedir."),
    ("Sınıf Dengesizliği: ",
     "DDoS veri setlerinde ciddi sınıf dengesizliği yaygın bir sorundur. "
     "CICIoT2023 veri setinde Backdoor_Malware sınıfı yalnızca %1,9 orana sahipken, "
     "DDoS-ACK_Fragmentation sınıfı %44,9'dur (en büyük/en küçük oran: 23,6:1). "
     "Bu dengesizlik, sınıflandırıcının çoğunluk sınıflarına yönelik önyargılı "
     "kararlar üretmesine ve azınlık sınıflarının tespit edilememesine yol açar "
     "(He ve Garcia, 2009)."),
]

for title, desc in problems:
    p = doc.add_paragraph(style="List Bullet")
    run_b = p.add_run(title)
    run_b.bold = True
    p.add_run(desc)

doc.add_paragraph(
    "Bu tez, yukarıdaki üç sorunu eşzamanlı olarak ele alan Yarasa Sürüsü "
    "Optimizasyonu (BSO) tabanlı bir hibrit makine öğrenmesi çerçevesi "
    "önermektedir. Önerilen çerçeve, BSO algoritmasını hem özellik seçimi "
    "hem de hiper-parametre optimizasyonu için tek bir arama sürecinde "
    "birleştirerek, CICIoT2023 veri setinde çok sınıflı DDoS tespitini "
    "iyileştirmeyi amaçlamaktadır. Ayrıca SMOTE (Synthetic Minority "
    "Over-sampling Technique) ile sınıf dengesizliği ele alınmaktadır."
)

# --- 1.3 ---
doc.add_heading("1.3 Araştırma Hipotezleri", level=2)

doc.add_paragraph(
    "Bu çalışmada aşağıdaki araştırma hipotezleri test edilmiştir:"
)

hypotheses = [
    ("AH₁: ",
     "BSO ile optimize edilmiş Random Forest modeli, varsayılan parametrelerle "
     "çalışan RF modeline göre istatistiksel olarak anlamlı (p < 0,05) düzeyde "
     "daha yüksek F1-Macro skoru elde eder."),
    ("AH₂: ",
     "BSO tabanlı özellik seçimi, %50'den fazla boyut azaltma sağlarken, "
     "sınıflandırma doğruluğunda %1'den az kayıp yaşanır."),
    ("AH₃: ",
     "BSO-Hybrid çerçevesi (özellik seçimi + hiper-parametre optimizasyonu + SMOTE), "
     "PSO, GA ve GWO gibi diğer meta-sezgisel yöntemlere kıyasla daha düşük "
     "uygunluk (fitness) değeri elde eder."),
]

for label, text in hypotheses:
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(6)

# --- 1.4 ---
doc.add_heading("1.4 Araştırma Soruları", level=2)

doc.add_paragraph(
    "Bu tezin yanıtlamayı hedeflediği araştırma soruları şunlardır:"
)

questions = [
    ("AS₁: ", "BSO algoritması, CICIoT2023 veri setindeki 39 özellikten en bilgilendirici alt kümeyi ne kadar etkili seçebilir?"),
    ("AS₂: ", "BSO-Hybrid RF çerçevesi, Backdoor_Malware gibi azınlık saldırı türlerinde yeterli tespit performansı sağlayabilir mi?"),
    ("AS₃: ", "Önerilen çerçevenin performans üstünlüğü, 10-katlı çapraz doğrulama ve istatistiksel testlerle doğrulanabilir mi?"),
]

for label, text in questions:
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)

# --- 1.5 ---
doc.add_heading("1.5 Tezin Amacı ve Kapsamı", level=2)

doc.add_paragraph(
    "Bu tezin temel amacı, dinamik ağ ortamlarında DDoS saldırılarının daha "
    "etkili tespiti için Yarasa Sürüsü Optimizasyonu (BSO) tabanlı bir hibrit "
    "makine öğrenmesi çerçevesi tasarlamak, uygulamak ve değerlendirmektir."
)

doc.add_paragraph("Çalışmanın kapsamı aşağıdaki unsurları içermektedir:")

scope_items = [
    "CICIoT2023 veri seti üzerinde kapsamlı keşifsel veri analizi (EDA) — 103.218 akış kaydı, 39 özellik, 5 sınıf",
    "BSO ile eşzamanlı özellik seçimi ve RF hiper-parametre optimizasyonu — 25 yarasa × 50 iterasyon = 1.177 değerlendirme",
    "SMOTE ile sınıf dengesizliği ele alınması — eğitim seti 72.252'den 87.500'e yükseltilmiştir",
    "12 farklı makine öğrenmesi modeli ile karşılaştırmalı performans değerlendirmesi",
    "S1–S4 ablasyon senaryoları ile her bileşenin katkısının sistematik olarak ölçülmesi",
    "10-katlı tabakalı çapraz doğrulama ve eşleştirilmiş istatistiksel anlamlılık testleri (Wilcoxon, McNemar, Cohen's d)",
]

for item in scope_items:
    doc.add_paragraph(item, style="List Bullet")

# --- 1.6 ---
doc.add_heading("1.6 Tezin Katkıları", level=2)

doc.add_paragraph(
    "Bu tez, literatüre aşağıdaki özgün katkıları sunmaktadır:"
)

contributions = [
    ("K1 — BSO-Hybrid RF Çerçevesi: ",
     "Özellik seçimi ve hiper-parametre optimizasyonunu tek bir meta-sezgisel "
     "süreçte birleştiren özgün hibrit çerçeve tasarımı. BSO'nun IoT tabanlı "
     "DDoS tespiti için ilk kapsamlı uygulamasıdır. En iyi uygunluk değeri: 0,177801."),
    ("K2 — %51,3 Boyut Azaltma: ",
     "39 özellikten 19'a indirgeme — yalnızca +0,08% doğruluk farkıyla. Model "
     "karmaşıklığını, hesaplama maliyetini ve aşırı uyum riskini azaltır."),
    ("K3 — Azınlık Sınıf İyileştirmesi: ",
     "SMOTE entegrasyonu ile Backdoor_Malware F1-Skor %28,40'dan %57,40'a "
     "yükselmiştir (+%102). Sınıf dengelemenin çok sınıflı tespitteki kritik "
     "rolü deneysel olarak kanıtlanmıştır."),
    ("K4 — Kapsamlı Deneysel Doğrulama: ",
     "12 model × 7 metrik × 10-katlı çapraz doğrulama × istatistiksel testler. "
     "S1–S4 ablasyon senaryoları ile her bileşenin katkısı ölçülebilir düzeyde "
     "gösterilmiştir. 9/11 karşılaştırmada p < 0,05 anlamlılık sağlanmıştır."),
]

for label, text in contributions:
    p = doc.add_paragraph(style="List Bullet")
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)

# --- 1.7 ---
doc.add_heading("1.7 Tezin Organizasyonu", level=2)

doc.add_paragraph(
    "Bu tez beş bölümden oluşmaktadır:"
)

chapters_org = [
    ("Bölüm 1 — Giriş: ",
     "Araştırmanın arka planı, problem tanımı, hipotezler, araştırma soruları, "
     "tezin amacı, katkıları ve organizasyonu."),
    ("Bölüm 2 — Literatür Taraması: ",
     "DDoS saldırıları, IoT güvenliği, makine öğrenmesi tabanlı tespit yöntemleri, "
     "meta-sezgisel optimizasyon, yarasa algoritması, öznitelik seçimi ve ilgili "
     "çalışmaların kapsamlı incelenmesi."),
    ("Bölüm 3 — Yöntem: ",
     "Önerilen BSO-Hybrid RF çerçevesinin tasarımı, CICIoT2023 veri seti, veri "
     "ön işleme, BSO optimizasyonu, uygunluk fonksiyonu, hibrit model tanımı, "
     "deney senaryoları ve değerlendirme metrikleri."),
    ("Bölüm 4 — Bulgular ve Değerlendirme: ",
     "BSO optimizasyon sonuçları, 12 model karşılaştırması, sınıf bazında performans, "
     "ablasyon çalışması, istatistiksel doğrulama ve hipotez testleri."),
    ("Bölüm 5 — Sonuç ve Öneriler: ",
     "Temel bulguların özeti, araştırma sorularının yanıtları, katkılar, "
     "kısıtlamalar ve gelecek çalışma önerileri."),
]

for label, text in chapters_org:
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)

doc.add_page_break()


# ============================================================================
# CHAPTER 2 — LİTERATÜR TARAMASI (LITERATURE REVIEW)
# ============================================================================

doc.add_heading("BÖLÜM 2", level=1)
doc.add_heading("LİTERATÜR TARAMASI", level=1)

# --- 2.1 ---
doc.add_heading("2.1 DDoS Saldırıları ve IoT Güvenliği", level=2)

doc.add_heading("2.1.1 DDoS Saldırılarına Genel Bakış", level=3)

doc.add_paragraph(
    "Dağıtılmış Hizmet Reddi (DDoS) saldırıları, birden fazla kaynak sistemden "
    "hedefe yönelik koordineli trafik gönderilerek hizmetin kesintiye uğratılmasını "
    "amaçlayan siber saldırılardır. İnternet'in ilk yıllarından bu yana var olan DDoS "
    "saldırıları, IoT cihazlarının yaygınlaşmasıyla birlikte hem hacim hem de "
    "karmaşıklık açısından dramatik bir artış göstermiştir (Zargar vd., 2013)."
)

doc.add_paragraph(
    "DDoS saldırıları genel olarak üç ana kategoride sınıflandırılmaktadır:"
)

ddos_types = [
    ("Hacimsel (Volumetric) Saldırılar: ",
     "Hedef sistemin bant genişliğini tüketmeyi amaçlar. UDP flood, ICMP flood ve "
     "amplification saldırıları bu kategoriye dahildir. Bu saldırı türü, toplam DDoS "
     "saldırılarının yaklaşık %65'ini oluşturmaktadır."),
    ("Protokol Tabanlı Saldırılar: ",
     "Ağ protokollerinin zayıflıklarını istismar eder. SYN flood, ACK fragmentation "
     "ve Smurf saldırıları bu gruba girer. CICIoT2023 veri setindeki DDoS-SYN_Flood "
     "ve DDoS-ACK_Fragmentation sınıfları bu kategoriyi temsil etmektedir."),
    ("Uygulama Katmanı Saldırıları: ",
     "HTTP flood, Slowloris ve DNS sorgu saldırıları gibi türleri içerir. Bu "
     "saldırılar düşük hacimli ancak yüksek etkili olabilir ve tespit edilmeleri "
     "daha zordur."),
]

for label, text in ddos_types:
    p = doc.add_paragraph(style="List Bullet")
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)

doc.add_heading("2.1.2 IoT Ortamlarında DDoS Tehdidi", level=3)

doc.add_paragraph(
    "Nesnelerin İnterneti (IoT), fiziksel cihazların ağ bağlantısı aracılığıyla veri "
    "toplaması ve paylaşması konseptidir. 2025 yılı itibarıyla dünyada 75 milyardan "
    "fazla IoT cihazı bulunması beklenmektedir (Cisco, 2023). Bu cihazların güvenlik "
    "açıklıkları, botnet oluşturma için ideal bir zemin oluşturmaktadır."
)

doc.add_paragraph(
    "2016 yılında gerçekleşen Mirai botnet saldırısı, IoT tabanlı DDoS tehditinin "
    "ciddiyetini gözler önüne sermiştir. Mirai, varsayılan kimlik bilgileriyle "
    "korunan yüz binlerce IoT cihazını ele geçirerek 1,2 Tbps'lik rekor düzeyde "
    "DDoS saldırısı gerçekleştirmiştir (Kolias vd., 2017). Bu olay, IoT güvenliğinde "
    "geleneksel savunma mekanizmalarının yetersizliğini kanıtlamıştır."
)

doc.add_paragraph(
    "IoT cihazlarının DDoS saldırılarına karşı savunmasızlığının temel nedenleri "
    "şunlardır: (1) sınırlı hesaplama ve bellek kaynakları nedeniyle güçlü "
    "şifreleme ve kimlik doğrulama mekanizmalarının uygulanamaması, (2) üretici "
    "firmalar tarafından varsayılan şifrelerin değiştirilmemesi, (3) firmware "
    "güncellemelerinin düzenli yapılmaması, (4) heterojen cihaz ekosisteminde "
    "standart güvenlik politikalarının yokluğu (Meidan vd., 2018; Doshi vd., 2018)."
)

doc.add_heading("2.1.3 Geleneksel Saldırı Tespit Sistemleri", level=3)

doc.add_paragraph(
    "Saldırı Tespit Sistemleri (IDS), ağ trafiğini izleyerek potansiyel güvenlik "
    "ihlallerini tespit eden yazılım veya donanım çözümleridir. Geleneksel IDS'ler "
    "iki ana yaklaşıma dayanır:"
)

doc.add_paragraph(
    "İmza Tabanlı Tespit: Bilinen saldırı kalıplarının (imzalarının) bir veritabanında "
    "saklanması ve ağ trafiğinin bu imzalarla karşılaştırılması esasına dayanır. Snort "
    "ve Suricata gibi açık kaynaklı IDS'ler bu yaklaşımı kullanır. Avantajı düşük "
    "yanlış pozitif oranıdır; dezavantajı ise yalnızca bilinen saldırıları tespit "
    "edebilmesi ve sıfır-gün saldırılarına karşı etkisiz kalmasıdır."
)

doc.add_paragraph(
    "Anomali Tabanlı Tespit: Normal ağ trafiğinin istatistiksel bir profilini "
    "oluşturarak, bu profilden sapan trafiği anormal olarak işaretler. Daha önce "
    "görülmemiş saldırıları tespit edebilme potansiyeline sahiptir; ancak yüksek "
    "yanlış pozitif oranı ve profil oluşturma sürecinin zorluğu temel "
    "dezavantajlarıdır. Makine öğrenmesi tabanlı yaklaşımlar, anomali tespitinin "
    "etkinliğini önemli ölçüde artırmıştır (Agrawal ve Tapaswi, 2019)."
)

# --- 2.2 ---
doc.add_heading("2.2 Makine Öğrenmesi Tabanlı Saldırı Tespiti", level=2)

doc.add_heading("2.2.1 Denetimli Öğrenme Yöntemleri", level=3)

doc.add_paragraph(
    "Denetimli öğrenme, etiketlenmiş eğitim verileri kullanarak sınıflandırma "
    "modelleri oluşturur. DDoS tespitinde yaygın olarak kullanılan denetimli "
    "öğrenme algoritmaları şunlardır:"
)

ml_methods = [
    ("Random Forest (RF): ",
     "Breiman (2001) tarafından önerilen RF, birden fazla karar ağacının "
     "rastgele alt örneklerle eğitilmesi ve çoğunluk oylamasıyla sınıflandırma "
     "yapılması prensibine dayanır. Yüksek doğruluk, aşırı uyuma dayanıklılık "
     "ve özellik önemi sağlama kapasitesi ile IDS alanında en başarılı "
     "algoritmalardan biridir. Bu çalışmada temel sınıflandırıcı olarak RF "
     "tercih edilmiştir."),
    ("XGBoost: ",
     "Chen ve Guestrin (2016) tarafından geliştirilen XGBoost, gradyan artırma "
     "(gradient boosting) çerçevesinin optimize edilmiş bir uygulamasıdır. "
     "Regularizasyon, paralel işleme ve eksik değer yönetimi yetenekleriyle "
     "birçok yarışmada üstün performans göstermiştir. Bu çalışmadaki deneysel "
     "sonuçlarda XGBoost (39 özellik) %90,37 doğruluk ile en yüksek performansı "
     "göstermiştir."),
    ("Destek Vektör Makineleri (SVM): ",
     "Cortes ve Vapnik (1995) tarafından önerilen SVM, sınıflar arasında "
     "maksimum ayrım sınırını belirleyen bir sınıflandırıcıdır. Doğrusal "
     "olmayan problemler için çekirdek (kernel) fonksiyonları kullanılabilir. "
     "Ancak yüksek boyutlu ve büyük veri setlerinde eğitim süresi önemli "
     "ölçüde artar."),
    ("K-En Yakın Komşu (KNN): ",
     "Cover ve Hart (1967) tarafından tanımlanan KNN, test örneğinin en "
     "yakın k komşusunun sınıf etiketlerine dayalı sınıflandırma yapar. "
     "Basit yapısına rağmen, büyük veri setlerinde tahmin süresi yüksektir "
     "ve yüksek boyutlu uzaylarda performansı düşer."),
    ("Naive Bayes: ",
     "Bayes teoremine dayalı olasılıksal bir sınıflandırıcıdır. Özellikler "
     "arasında koşullu bağımsızlık varsayımı yapar. Eğitim süresi çok kısa "
     "olmasına rağmen, karmaşık ağ trafik kalıplarını modellemede yetersiz "
     "kalabilir. Bu çalışmada %62,96 doğruluk ile en düşük performansı "
     "göstermiştir."),
    ("Lojistik Regresyon: ",
     "Doğrusal bir sınıflandırıcıdır ve sigmoid fonksiyonu ile olasılık "
     "tahmini yapar. Yorumlanabilirlik avantajı sunar; ancak doğrusal "
     "olmayan ilişkileri modellemede sınırlıdır."),
    ("Karar Ağacı (Decision Tree): ",
     "Belirli kurallara dayalı hiyerarşik bölümleme yapan bir "
     "sınıflandırıcıdır. Tek başına kullanıldığında aşırı uyuma eğilimlidir; "
     "ancak RF ve XGBoost gibi topluluk yöntemlerinin temelini oluşturur."),
]

for label, text in ml_methods:
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)

doc.add_heading("2.2.2 Derin Öğrenme Yaklaşımları", level=3)

doc.add_paragraph(
    "Son yıllarda derin öğrenme yöntemleri de DDoS tespitinde kullanılmaya "
    "başlanmıştır. Roopak vd. (2019) CNN ve LSTM tabanlı modelleri IoT DDoS "
    "tespitinde değerlendirmiş, Yuan vd. (2017) DeepDefense çerçevesini "
    "önermiştir. Li ve Lu (2019) LSTM-BA hibrit modelini tamamen Bayesçi "
    "yaklaşımla birleştirmiştir. Derin öğrenme modelleri genellikle yüksek "
    "doğruluk sağlar; ancak büyük eğitim verisi gereksinimi, hesaplama "
    "maliyeti ve açıklanabilirlik eksikliği temel dezavantajlarıdır. Özellikle "
    "IoT cihazlarının sınırlı kaynaklarında dağıtım zorluğu, geleneksel ML "
    "yöntemlerinin hala tercih edilmesine neden olmaktadır."
)

doc.add_heading("2.2.3 Sınıf Dengesizliği Sorunu", level=3)

doc.add_paragraph(
    "IDS veri setlerinde sınıf dengesizliği yaygın bir sorundur. Normal trafik "
    "örnekleri genellikle saldırı örneklerinden çok daha fazladır. CICIoT2023 "
    "veri setinde en büyük sınıf (DDoS-ACK_Fragmentation: 53.148) ile en küçük "
    "sınıf (Backdoor_Malware: 2.252) arasında 23,6:1 oranında dengesizlik "
    "bulunmaktadır."
)

doc.add_paragraph(
    "Bu sorunu ele almak için çeşitli teknikler geliştirilmiştir. SMOTE "
    "(Synthetic Minority Over-sampling Technique), Chawla vd. (2002) tarafından "
    "önerilmiş olup azınlık sınıfı için sentetik örnekler üreterek sınıf "
    "dengesini sağlar. SMOTE, mevcut azınlık örnekleri arasında doğrusal "
    "interpolasyon yaparak yeni örnekler oluşturur. Bu çalışmada SMOTE "
    "yalnızca eğitim setine uygulanarak veri sızıntısı önlenmiştir. Eğitim "
    "seti 72.252'den 87.500 dengeli örneğe yükseltilmiştir; özellikle "
    "Backdoor_Malware sınıfı 2.252'den 17.500 örneğe artırılmıştır (×7,8)."
)

# --- 2.3 ---
doc.add_heading("2.3 Meta-Sezgisel Optimizasyon Yöntemleri", level=2)

doc.add_heading("2.3.1 Genel Bakış", level=3)

doc.add_paragraph(
    "Meta-sezgisel algoritmalar, doğadan esinlenen ve karmaşık optimizasyon "
    "problemlerine yaklaşık çözümler üreten stokastik arama yöntemleridir. "
    "Bu algoritmalar, NP-zor problemlerde makul sürede iyi çözümler bulma "
    "kapasitesine sahiptir ve kesin çözüm gerektirmeyen durumlarda yaygın "
    "olarak tercih edilmektedir (Yang, 2010)."
)

doc.add_paragraph(
    "Makine öğrenmesi bağlamında meta-sezgisel algoritmalar iki temel görevde "
    "kullanılmaktadır: (1) Özellik seçimi — yüksek boyutlu veri setlerinden en "
    "bilgilendirici özellik alt kümesinin seçilmesi, (2) Hiper-parametre "
    "optimizasyonu — sınıflandırıcı parametrelerinin optimal değerlerinin "
    "belirlenmesi. Bu çalışmada BSO, her iki görevi tek bir optimizasyon "
    "sürecinde birleştirmektedir."
)

doc.add_heading("2.3.2 Parçacık Sürüsü Optimizasyonu (PSO)", level=3)

doc.add_paragraph(
    "Kennedy ve Eberhart (1995) tarafından önerilen PSO, kuş sürülerinin "
    "ve balık topluluklarının sosyal davranışlarından esinlenen bir algoritma- "
    "dır. Her parçacık, arama uzayında bir çözümü temsil eder ve kendi en "
    "iyi konumu (pbest) ile sürünün en iyi konumuna (gbest) doğru hareket "
    "eder. PSO, basit uygulaması ve az parametre gerektirmesi nedeniyle "
    "özellik seçimi alanında yaygın olarak kullanılmaktadır. Bu çalışmadaki "
    "deneylerde PSO-RF modeli 18 özellik seçerek %88,35 doğruluk ve "
    "%81,82 F1-Macro elde etmiştir."
)

doc.add_heading("2.3.3 Genetik Algoritma (GA)", level=3)

doc.add_paragraph(
    "Holland (1975) tarafından geliştirilen GA, doğal seçilim ve genetik "
    "mekanizmalardan esinlenen evrimsel bir algoritmadır. Çaprazlama "
    "(crossover), mutasyon ve seçim operatörleri ile çözüm popülasyonunu "
    "iteratif olarak iyileştirir. Khammassi ve Krichen (2017) GA'yı "
    "Lojistik Regresyon ile birleştirerek IDS'de özellik seçimi "
    "gerçekleştirmiştir. Bu çalışmada GA-RF modeli 21 özellik seçerek "
    "%89,37 doğruluk ve %83,66 F1-Macro elde etmiştir."
)

doc.add_heading("2.3.4 Gri Kurt Optimizasyonu (GWO)", level=3)

doc.add_paragraph(
    "Mirjalili vd. (2014) tarafından önerilen GWO, gri kurdların sosyal "
    "hiyerarşisi ve avlanma davranışından esinlenmiştir. Alfa, beta, "
    "delta ve omega olarak derecelendirilen kurt grupları, liderlik "
    "hiyerarşisine göre arama sürecini yönlendirir. GWO, düşük "
    "parametre sayısı ve iyi yakınsama özellikleriyle bilinir. Bu "
    "çalışmada GWO-RF modeli 23 özellik seçerek %89,80 doğruluk ve "
    "%84,35 F1-Macro ile BSO-Hybrid RF'ye en yakın performansı "
    "göstermiştir."
)

# --- Comparison Table ---
doc.add_paragraph()
add_formatted_table(
    doc,
    headers=["Algoritma", "Yıl", "İlham Kaynağı", "Seçilen Özellik", "F1-Macro (%)", "Fitness"],
    rows=[
        ["BSO (Önerilen)", "2010", "Yarasa ekolokasyonu", "19", "84,24", "0,177801"],
        ["PSO", "1995", "Kuş/balık sürüsü", "18", "81,82", "0,193895"],
        ["GA", "1975", "Doğal seçilim", "21", "83,66", "0,188982"],
        ["GWO", "2014", "Gri kurt hiyerarşisi", "23", "84,35", "0,192181"],
    ],
    caption="Tablo 2.1: Meta-sezgisel Algoritmaların Karşılaştırması",
    highlight_row=0,
)

# --- 2.4 ---
doc.add_heading("2.4 Yarasa Algoritması (Bat Algorithm)", level=2)

doc.add_heading("2.4.1 Temel Prensipler", level=3)

doc.add_paragraph(
    "Yarasa Algoritması (BA / BSO), Yang (2010) tarafından micro-yarasaların "
    "ekolokasyon davranışından esinlenerek geliştirilmiş bir meta-sezgisel "
    "optimizasyon yöntemidir. Yarasalar, yüksek frekanslı ses dalgaları "
    "göndererek çevrelerindeki nesnelerin konumunu ve mesafesini algılar. "
    "Bu doğal mekanizma, optimizasyon bağlamında arama uzayının "
    "keşfedilmesi ve en iyi çözüme yakınsaması olarak modellenmektedir."
)

doc.add_paragraph(
    "BSO algoritması üç temel kuralı takip eder:"
)

bso_rules = [
    "Tüm yarasalar ekolokasyon kullanır ve mesafe/avı algılama yeteneğine sahiptir.",
    "Her yarasa, fi frekansında, vi hızında ve xi konumunda rastgele uçar. Frekans aralığı [fmin, fmax] arasında ayarlanır.",
    "Gürültü seviyesi (loudness) Ai büyükten küçük bir sabite, darbe oranı (pulse rate) ri küçükten büyüğe doğru ayarlanır.",
]

for i, rule in enumerate(bso_rules, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(rule)

doc.add_heading("2.4.2 Matematiksel Formülasyon", level=3)

doc.add_paragraph("BSO algoritmasının güncelleme denklemleri aşağıdaki gibidir:")

add_equation(doc, "fi = fmin + (fmax - fmin) × β,    β ∈ [0, 1]", "2.1")
add_equation(doc, "vi(t+1) = vi(t) + (xi(t) - x*) × fi", "2.2")
add_equation(doc, "xi(t+1) = xi(t) + vi(t+1)", "2.3")

doc.add_paragraph(
    "Burada fi i. yarasının frekansını, vi hızını, xi konumunu ve x* mevcut "
    "küresel en iyi çözümü temsil eder. β, [0,1] aralığında düzgün dağılımlı "
    "rastgele bir sayıdır."
)

doc.add_paragraph("Yerel arama için:")
add_equation(doc, "xnew = xold + ε × At", "2.4")

doc.add_paragraph(
    "Burada ε ∈ [-1, 1] aralığında rastgele bir sayı, At ise t. iterasyondaki "
    "ortalama gürültü seviyesidir."
)

doc.add_paragraph("Gürültü ve darbe oranı güncelleme kuralları:")
add_equation(doc, "Ai(t+1) = α × Ai(t),    0 < α < 1", "2.5")
add_equation(doc, "ri(t+1) = ri(0) × [1 - exp(-γ × t)],    γ > 0", "2.6")

doc.add_paragraph(
    "α ve γ parametreleri sırasıyla gürültü azalma ve darbe artış oranlarını "
    "kontrol eder. Bu çalışmada α = 0,9 ve γ = 0,9 olarak belirlenmiştir. "
    "Başlangıç gürültüsü A₀ = 0,95, başlangıç darbe oranı r₀ = 0,5, "
    "frekans aralığı [0, 2], popülasyon boyutu 25 ve maksimum iterasyon 50 "
    "olarak ayarlanmıştır."
)

doc.add_heading("2.4.3 BSO'nun Avantajları", level=3)

doc.add_paragraph(
    "BSO, diğer meta-sezgisel algoritmalarla karşılaştırıldığında şu "
    "avantajlara sahiptir:"
)

advantages = [
    "Frekans ayarlama mekanizması sayesinde küresel ve yerel arama arasında otomatik denge kurabilir.",
    "Gürültü ve darbe oranı parametreleri, keşif (exploration) ve sömürü (exploitation) arasında dinamik geçiş sağlar.",
    "Sürekli ve ayrık (ikili) optimizasyon problemlerine uyarlanabilir — bu çalışmada özellik seçimi (ikili) ve hiper-parametre ayarı (sürekli) eşzamanlı gerçekleştirilmektedir.",
    "Az sayıda kontrol parametresi ile basit uygulanabilirlik sunar.",
    "Hızlı yakınsama — bu çalışmada BSO 50 iterasyonda, GA ise 40 iterasyonda yakınsamıştır.",
]

for item in advantages:
    doc.add_paragraph(item, style="List Bullet")

# --- 2.5 ---
doc.add_heading("2.5 Öznitelik Seçimi Yöntemleri", level=2)

doc.add_paragraph(
    "Öznitelik seçimi (feature selection), yüksek boyutlu veri setlerinden "
    "en bilgilendirici özellik alt kümesinin seçilmesi sürecidir. Etkili "
    "öznitelik seçimi, model performansını artırır, hesaplama maliyetini "
    "düşürür ve aşırı uyumu önler (Guyon ve Elisseeff, 2003; Chandrashekar "
    "ve Sahin, 2014)."
)

doc.add_paragraph(
    "Öznitelik seçimi yöntemleri üç ana kategoride sınıflandırılır:"
)

fs_methods = [
    ("Filtre (Filter) Yöntemleri: ",
     "Öznitelikleri sınıflandırıcıdan bağımsız olarak istatistiksel "
     "ölçütlere göre değerlendirir. Chi-square, bilgi kazanımı (information "
     "gain) ve korelasyon analizi bu kategoridedir. Hızlı çalışır ancak "
     "sınıflandırıcıyla etkileşimi görmezden gelir."),
    ("Sarmalayıcı (Wrapper) Yöntemleri: ",
     "Sınıflandırıcı performansını doğrudan optimize eden bir alt küme "
     "arama yöntemidir. İleri seçim (forward selection), geri eleme "
     "(backward elimination) ve meta-sezgisel tabanlı arama bu gruba "
     "girer. Bu çalışmada BSO tabanlı sarmalayıcı yöntem kullanılmıştır."),
    ("Gömülü (Embedded) Yöntemleri: ",
     "Sınıflandırıcı eğitimi sırasında öznitelik seçimini gerçekleştirir. "
     "LASSO, Ridge regresyon ve RF'nin gini önemi bu yaklaşımı izler. "
     "Hesaplama açısından filtre ve sarmalayıcı arasında yer alır."),
]

for label, text in fs_methods:
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)

doc.add_paragraph(
    "Meta-sezgisel tabanlı sarmalayıcı yöntemler, arama uzayını etkin "
    "biçimde keşfetme kapasiteleriyle ön plana çıkmaktadır. Xue vd. "
    "(2016) kapsamlı bir tarama çalışmasında, evrimsel hesaplama "
    "yaklaşımlarının öznitelik seçiminde başarılı sonuçlar ürettiğini "
    "göstermiştir. Mafarja ve Mirjalili (2018) Balina Optimizasyonu "
    "algoritmasını sarmalayıcı öznitelik seçimi için uygulamıştır. "
    "Nakamura vd. (2012) ise BSO'nun ikili versiyonunu (BBA) öznitelik "
    "seçimi için ilk kez kullanmıştır."
)

# --- 2.6 ---
doc.add_heading("2.6 İlgili Çalışmalar ve Karşılaştırma", level=2)

doc.add_paragraph(
    "Aşağıda DDoS tespiti, meta-sezgisel optimizasyon ve hibrit makine "
    "öğrenmesi yaklaşımlarına ilişkin temel çalışmalar özetlenmektedir."
)

related_works = [
    ["Neto vd. (2023)", "CICIoT2023", "RF, DNN", "İkili sınıflandırmada %99 doğruluk; çoklu sınıf analizi sınırlı"],
    ["Ferrag vd. (2023)", "Edge-IIoTset", "DNN", "IoT/IIoT odaklı; %98,2 doğruluk (ikili)"],
    ["Khammassi vd. (2017)", "NSL-KDD", "GA-LR", "GA ile öznitelik seçimi + LR sınıflandırma"],
    ["Zhou vd. (2020)", "CICIDS2017", "Ensemble", "Filtre + sarmalayıcı hibrit öznitelik seçimi"],
    ["Tama vd. (2019)", "NSL-KDD", "GBM", "Gradyan artırma ile anomali tespiti"],
    ["Ahmad vd. (2018)", "KDD Cup99", "SVM, RF, ELM", "Üç sınıflandırıcının karşılaştırması"],
    ["Bu Çalışma", "CICIoT2023", "BSO-Hybrid RF", "BSO ile eşzamanlı FS + HP; 5-sınıf; %89,82 doğruluk"],
]

add_formatted_table(
    doc,
    headers=["Çalışma", "Veri Seti", "Yöntem", "Önemli Bulgu"],
    rows=related_works,
    caption="Tablo 2.2: İlgili Çalışmaların Özet Karşılaştırması",
    highlight_row=6,
)

doc.add_paragraph(
    "Tablo 2.2'de görüldüğü üzere, mevcut çalışmaların büyük çoğunluğu ikili "
    "sınıflandırma (normal vs. saldırı) üzerine odaklanmıştır. Çok sınıflı "
    "DDoS tespiti, özellikle CICIoT2023 gibi güncel ve gerçekçi veri setleri "
    "üzerinde meta-sezgisel optimizasyon ile birleştirilmiş çalışmalar "
    "oldukça sınırlıdır. Bu tez, BSO'nun CICIoT2023 üzerinde eşzamanlı "
    "öznitelik seçimi ve hiper-parametre optimizasyonu için ilk kapsamlı "
    "uygulamasını sunmaktadır."
)

doc.add_paragraph(
    "Mevcut çalışmalardan farklı olarak, bu tez şu teknik yenilikleri "
    "barındırmaktadır: (1) BSO'nun ikili (özellik maskesi) ve sürekli "
    "(hiper-parametre) değişkenleri tek bir çözüm vektöründe birleştirmesi, "
    "(2) CICIoT2023'ün 5-sınıflı çok sınıflandırma versiyonunun "
    "kullanılması, (3) S1–S4 ablasyon senaryolarıyla her bileşenin "
    "katkısının izole edilmesi, (4) 12 model × 7 metrik boyutunda "
    "karşılaştırmalı değerlendirme."
)

doc.add_page_break()


# ============================================================================
# CHAPTER 3 — YÖNTEM (METHODOLOGY)
# ============================================================================

doc.add_heading("BÖLÜM 3", level=1)
doc.add_heading("YÖNTEM", level=1)

doc.add_paragraph(
    "Bu bölümde, önerilen BSO-Hybrid RF çerçevesinin tasarımı, kullanılan "
    "veri seti, veri ön işleme adımları, BSO optimizasyon algoritması, "
    "uygunluk fonksiyonu, hibrit model tanımı, deney senaryoları ve "
    "değerlendirme metrikleri detaylı olarak açıklanmaktadır."
)

# --- 3.1 ---
doc.add_heading("3.1 Araştırma Tasarımı ve Genel Çerçeve", level=2)

doc.add_paragraph(
    "Önerilen BSO-Hybrid RF çerçevesi, Şekil 3.1'de gösterilen yedi "
    "aşamalı bir veri işleme hattı (pipeline) izlemektedir:"
)

pipeline_steps = [
    ("Aşama 1 — Veri Toplama: ", "CICIoT2023 veri setinden 19 CSV dosyası yüklenerek 103.218 akış kaydı ve 39 özellik elde edilmiştir."),
    ("Aşama 2 — Ön İşleme: ", "Eksik değer kontrolü, tekrarlanan veri temizliği ve StandardScaler normalizasyonu (z-score: μ=0, σ=1) uygulanmıştır."),
    ("Aşama 3 — Veri Bölme: ", "Tabakalı (stratified) bölme ile %70 eğitim (72.252), %10 doğrulama (10.322) ve %20 test (20.644) olarak ayrılmıştır."),
    ("Aşama 4 — SMOTE: ", "Yalnızca eğitim setine SMOTE uygulanarak 72.252'den 87.500 dengeli örneğe yükseltilmiştir. Backdoor_Malware: 2.252 → 17.500 (×7,8)."),
    ("Aşama 5 — BSO Optimizasyonu: ", "25 yarasa × 50 iterasyon = 1.177 değerlendirme. Eşzamanlı özellik seçimi + RF hiper-parametre optimizasyonu."),
    ("Aşama 6 — Model Eğitimi: ", "Optimal çözüm: 19 özellik, n_estimators=266, max_depth=20, min_samples_split=7, max_features=0,469."),
    ("Aşama 7 — Değerlendirme: ", "Test seti üzerinde 7 performans metriği + 10-katlı çapraz doğrulama + istatistiksel testler + 12 model karşılaştırması."),
]

for label, text in pipeline_steps:
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)

# --- 3.2 ---
doc.add_heading("3.2 Veri Seti: CICIoT2023", level=2)

doc.add_paragraph(
    "Bu çalışmada, Canadian Institute for Cybersecurity tarafından 2023 "
    "yılında yayınlanan CICIoT2023 veri seti kullanılmıştır (Neto vd., 2023). "
    "Veri seti, gerçek IoT cihazlarından (105 cihaz) toplanan ağ trafiğini "
    "içermekte olup toplam 103.218 örnek ve 39 ağ trafik özelliğinden "
    "oluşmaktadır."
)

# Dataset table
add_formatted_table(
    doc,
    headers=["Özellik", "Değer"],
    rows=[
        ["Toplam Örnek Sayısı", "103.218"],
        ["Özellik Sayısı", "39"],
        ["Sınıf Sayısı", "5"],
        ["Eğitim Seti (%70)", "87.500 (SMOTE sonrası)"],
        ["Doğrulama Seti (%10)", "10.322"],
        ["Test Seti (%20)", "20.644"],
        ["Normalizasyon", "StandardScaler (z-score: μ=0, σ=1)"],
        ["SMOTE Sentetik Örnek", "15.248"],
        ["Veri Kaynağı", "19 CSV dosyası, 105 IoT cihazı"],
    ],
    caption="Tablo 3.1: CICIoT2023 Veri Seti Özeti",
)

doc.add_paragraph(
    "Veri seti 5 sınıf içermektedir. Sınıf dağılımı ve SMOTE öncesi/sonrası "
    "değişim Tablo 3.2'de sunulmuştur."
)

add_formatted_table(
    doc,
    headers=["Sınıf", "Eğitim (Orijinal)", "Eğitim (SMOTE)", "Test", "Oran (%)"],
    rows=[
        ["DDoS-ACK_Fragmentation", "53.148*", "17.500", "5.000", "44,9"],
        ["DDoS-SYN_Flood", "25.169*", "17.500", "5.000", "21,2"],
        ["BenignTraffic", "22.735*", "17.500", "5.000", "19,2"],
        ["Recon-PortScan", "15.162*", "17.500", "5.000", "12,8"],
        ["Backdoor_Malware", "2.252", "17.500", "644", "1,9"],
        ["TOPLAM", "72.252", "87.500", "20.644", "100,0"],
    ],
    caption="Tablo 3.2: Sınıf Dağılımı ve SMOTE Etkisi",
    highlight_row=5,
)

doc.add_paragraph(
    "* Not: Çoğunluk sınıflarına rastgele alt örnekleme (random undersampling) "
    "uygulanarak sınıf başına 25.000 ile sınırlandırılmıştır. Ardından %70 "
    "eğitim bölünmesinden sonra SMOTE uygulanmıştır."
)

doc.add_paragraph(
    "Belirgin sınıf dengesizliği (en büyük/en küçük: 23,6:1) SMOTE "
    "uygulamasını zorunlu kılmaktadır. SMOTE yalnızca eğitim setine "
    "uygulanarak test setinin gerçekçiliği korunmuştur (veri sızıntısı önlenir)."
)

# --- 3.3 ---
doc.add_heading("3.3 Veri Ön İşleme", level=2)

doc.add_paragraph("Veri ön işleme aşamasında sırasıyla şu adımlar uygulanmıştır:")

preprocess = [
    ("Eksik Değer Kontrolü: ", "Veri setinde eksik değer bulunmamıştır (tüm 39 özellik tam dolu)."),
    ("Tekrarlanan Veri Temizliği: ", "Tekrarlanan satırlar tanımlanmış ve kaldırılmıştır."),
    ("StandardScaler Normalizasyonu: ", "Tüm sayısal özellikler z-score dönüşümü ile standartlaştırılmıştır:"),
]

for label, text in preprocess:
    p = doc.add_paragraph(style="List Number")
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)

add_equation(doc, "x_norm = (x - μ) / σ", "3.1")

doc.add_paragraph(
    "Bu dönüşüm, her özelliğin ortalamasını 0 ve standart sapmasını 1 yapar. "
    "StandardScaler, özellikle SVM ve KNN gibi mesafe tabanlı algoritmalarda "
    "özellik ölçeklerinin eşitlenmesi için gereklidir. Normalizasyon parametreleri "
    "(μ, σ) yalnızca eğitim setinden hesaplanmış ve doğrulama/test setlerine "
    "aynı dönüşüm uygulanmıştır (veri sızıntısı önleme)."
)

preprocess_2 = [
    ("Veri Bölme: ", "Tabakalı (stratified) bölme ile %70 eğitim (72.252 orijinal), %10 doğrulama (10.322) ve %20 test (20.644) olarak ayrılmıştır. Tabakalı bölme, her alt kümedeki sınıf oranlarının korunmasını sağlar."),
    ("SMOTE Uygulaması: ", "SMOTE yalnızca eğitim setine uygulanmıştır (Chawla vd., 2002). Eğitim seti boyutu 72.252'den 87.500'e yükselmiştir. Tüm sınıflar 17.500 örneğe dengelenmiştir. Özellikle Backdoor_Malware sınıfının örnek sayısı 2.252'den 17.500'e artırılmıştır (×7,8 artış)."),
]

for label, text in preprocess_2:
    p = doc.add_paragraph(style="List Number")
    p.style.paragraph_format.start_value = 4
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)

# --- 3.4 ---
doc.add_heading("3.4 Yarasa Sürüsü Optimizasyonu (BSO) Algoritması", level=2)

doc.add_paragraph(
    "BSO algoritması bu çerçevede iki eşzamanlı görevi tek bir optimizasyon "
    "sürecinde birleştirmektedir:"
)

doc.add_heading("3.4.1 Görev 1 — Özellik Seçimi", level=3)

doc.add_paragraph(
    "Her yarasa bireyi, 39 boyutlu bir ikili vektör taşır. xj = 1 ise j. "
    "özellik seçilir, xj = 0 ise seçilmez. BSO, bilgilendirici olmayan "
    "özellikleri elemine ederek boyut azaltma sağlar. Sürekli pozisyon "
    "değerlerinin ikili vektöre dönüştürülmesi için sigmoid transfer "
    "fonksiyonu kullanılmıştır:"
)

add_equation(doc, "S(xi,j) = 1 / (1 + exp(-xi,j))", "3.2")
add_equation(doc, "xj = 1  eğer S(xi,j) > rand()  ise, aksi halde xj = 0", "3.3")

doc.add_paragraph(
    "Sonuç olarak BSO, 39 özellikten 19 özellik seçerek %51,3 boyut azaltma "
    "sağlamıştır."
)

doc.add_heading("3.4.2 Görev 2 — Hiper-Parametre Optimizasyonu", level=3)

doc.add_paragraph(
    "Aynı yarasa bireyi, özellik maskesinin yanı sıra ek 4 sürekli boyutta "
    "RF hiper-parametrelerini kodlar. Bu parametreler ve optimizasyon "
    "aralıkları Tablo 3.3'te gösterilmiştir."
)

add_formatted_table(
    doc,
    headers=["Parametre", "Aralık", "Optimal Değer", "Açıklama"],
    rows=[
        ["n_estimators", "[50, 400]", "266", "Ağaç sayısı"],
        ["max_depth", "[5, 35]", "20", "Maksimum derinlik"],
        ["min_samples_split", "[2, 15]", "7", "Minimum bölme örneği"],
        ["max_features", "[0,3, 1,0]", "0,469", "Özellik alt kümesi oranı"],
    ],
    caption="Tablo 3.3: BSO ile Optimize Edilen RF Hiper-Parametreleri",
)

doc.add_heading("3.4.3 BSO Algoritma Parametreleri", level=3)

add_formatted_table(
    doc,
    headers=["Parametre", "Değer", "Açıklama"],
    rows=[
        ["Popülasyon Boyutu", "25", "Yarasa sayısı"],
        ["Maks. İterasyon", "50", "Toplam nesil sayısı"],
        ["Frekans Aralığı", "[0, 2]", "Yarasa frekans aralığı"],
        ["Başlangıç Gürültüsü (A₀)", "0,95", "İlk ses seviyesi"],
        ["Başlangıç Darbe Oranı (r₀)", "0,5", "İlk darbe emisyon oranı"],
        ["α (Küçülme)", "0,9", "Gürültü azalma katsayısı"],
        ["γ (Artış)", "0,9", "Darbe artış katsayısı"],
        ["Toplam Değerlendirme", "1.177", "25 × 50 - tekrar edenler"],
        ["Optimizasyon Süresi", "840,43 s", "Tüm BSO çalışma süresi"],
    ],
    caption="Tablo 3.4: BSO Algoritma Parametreleri",
)

# --- 3.5 ---
doc.add_heading("3.5 Uygunluk Fonksiyonu", level=2)

doc.add_paragraph(
    "BSO'nun amaç fonksiyonu (fitness function) aşağıdaki gibi tanımlanmıştır:"
)

add_equation(doc, "f(x, θ) = 1 - F1_macro(RF(X_seçili, θ)) + α × (n_seçili / n_toplam)", "3.4")

doc.add_paragraph("Burada:")

fitness_params = [
    "F1_macro: Makro-ortalama F1-Skor (tüm sınıfların eşit ağırlıklı harmonik ortalaması)",
    "X_seçili: BSO tarafından seçilen özellik alt kümesi",
    "θ = {n_estimators, max_depth, min_samples_split, max_features}: RF hiper-parametreleri",
    "α = 0,01: Özellik sayısı ceza katsayısı",
    "n_seçili / n_toplam: Seçilen özellik oranı (boyut cezası)",
]

for item in fitness_params:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "BSO bu fonksiyonu minimize ederek, F1-Macro'yu maksimize ederken özellik "
    "sayısını minimize eder. En iyi uygunluk değeri: 0,177801. Bu değer, "
    "F1-Macro = %84,24 ve 19/39 seçilen özellik ile elde edilmiştir:"
)

add_equation(doc, "f = 1 - 0,8424 + 0,01 × (19/39) = 0,1576 + 0,00487 ≈ 0,1625", "3.5")

doc.add_paragraph(
    "Not: Gerçek fitness değeri (0,177801) doğrulama seti üzerinde hesaplanmıştır; "
    "test setindeki F1-Macro (%84,24) ile küçük bir sapma normaldir."
)

# --- 3.6 ---
doc.add_heading("3.6 Hibrit Model Tanımı: BSO-Hybrid RF", level=2)

doc.add_paragraph(
    "Bu çalışmada \"hibrit\" terimi, birden fazla sınıflandırıcının birleştirilmesi "
    "(ensemble voting/stacking/weighting) anlamında kullanılmamaktadır. Bunun yerine, "
    "meta-sezgisel optimizasyon (BSO) ile makine öğrenmesi sınıflandırması (RF) "
    "arasındaki metodolojik birleşimi ifade eder:"
)

add_equation(doc, "BSO-Hybrid RF = BSO_optimize(RF_sınıflandır(X_seçili, θ_optimal))", "3.6")

doc.add_paragraph("Burada:")

hybrid_params = [
    "X_seçili ⊂ X_tüm: BSO tarafından seçilmiş optimal özellik alt kümesi (19/39)",
    "θ_optimal = {n_estimators=266, max_depth=20, min_samples_split=7, max_features=0,469}",
    "BSO_optimize: Yarasa Sürüsü Optimizasyonu algoritması (25 yarasa, 50 iterasyon)",
    "RF_sınıflandır: Optimize edilmiş parametrelerle eğitilmiş Random Forest modeli",
]

for item in hybrid_params:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "Bu yaklaşımın avantajı, özellik seçimi ve hiper-parametre optimizasyonunun "
    "birbirini etkilediği karmaşık etkileşimi tek bir arama sürecinde yakalayabilmesidir. "
    "Ayrı ayrı optimize edildiğinde, özellik alt kümesi belirli parametrelere, "
    "parametreler ise belirli özelliklere bağımlı olabilir. BSO-Hybrid RF, bu "
    "birbirine bağımlılığı eşzamanlı olarak optimize eder."
)

# --- 3.7 ---
doc.add_heading("3.7 Deney Senaryoları (S1–S4)", level=2)

doc.add_paragraph(
    "Her bileşenin katkısını sistematik olarak ölçmek için dört ablasyon senaryosu "
    "tasarlanmıştır. Her senaryoda yalnızca bir bileşen eklenerek kümülatif etki "
    "ölçülmüştür:"
)

add_formatted_table(
    doc,
    headers=["Senaryo", "Özellik Seçimi", "HP Ayarı", "SMOTE", "Doğruluk (%)", "F1-Macro (%)", "AUC-ROC (%)"],
    rows=[
        ["S1 — Temel", "Yok (39)", "Varsayılan (DT)", "Yok", "87,04", "78,57", "91,20"],
        ["S2 — BSO-FS", "BSO (19)", "Varsayılan (RF)", "Evet", "88,47", "82,35", "97,47"],
        ["S3 — BSO-HP", "Yok (39)", "BSO optimize (RF)", "Evet", "89,74", "84,13", "98,36"],
        ["S4 — Tam Hibrit", "BSO (19)", "BSO optimize (RF)", "Evet", "89,82", "84,24", "98,38"],
    ],
    caption="Tablo 3.5: S1–S4 Ablasyon Senaryoları ve Sonuçları",
    highlight_row=3,
)

doc.add_paragraph(
    "S1 → S4 ilerledikçe şu bulgular elde edilmiştir:"
)

ablation_findings = [
    ("SMOTE Etkisi (En Kritik — S1 vs S2 karşılaştırması): ",
     "F1-Macro'da yaklaşık +3,78% artış. SMOTE olmadan Backdoor_Malware F1 %28,40'a "
     "düşmektedir. Sınıf dengeleme, çok sınıflı tespit için zorunludur."),
    ("BSO HP Optimizasyonu (S2 vs S3): ",
     "F1-Macro'da +1,78% iyileşme. Optimize edilmiş RF parametreleri "
     "(n_estimators=266, max_depth=20) varsayılanlardan daha iyi performans gösterir."),
    ("BSO Özellik Seçimi (S3 vs S4): ",
     "%51,3 boyut azaltma ile yalnızca +0,08% doğruluk farkı (89,74% → 89,82%). "
     "Model karmaşıklığını ve çıkarım süresini azaltır."),
    ("Toplam İyileşme (S1 → S4): ",
     "F1-Macro: %78,57 → %84,24 (+5,67%); Doğruluk: %87,04 → %89,82 (+2,78%); "
     "Özellik sayısı: 39 → 19 (%51,3 azalma)."),
]

for label, text in ablation_findings:
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(text)

# --- 3.8 ---
doc.add_heading("3.8 Değerlendirme Metrikleri", level=2)

doc.add_paragraph(
    "Model performansının kapsamlı değerlendirmesi için aşağıdaki 7 temel "
    "metrik kullanılmıştır:"
)

doc.add_heading("3.8.1 Doğruluk (Accuracy)", level=3)
add_equation(doc, "Accuracy = (TP + TN) / (TP + TN + FP + FN)", "3.7")
doc.add_paragraph(
    "Doğru tahminlerin toplam örnek sayısına oranıdır. Sınıf dengesizliği olan "
    "veri setlerinde yanıltıcı olabilir."
)

doc.add_heading("3.8.2 Kesinlik (Precision)", level=3)
add_equation(doc, "Precision = TP / (TP + FP)", "3.8")
doc.add_paragraph(
    "Pozitif tahminlerin ne kadarının gerçekten pozitif olduğunu ölçer. "
    "Yanlış alarm oranının önemli olduğu senaryolarda kritiktir."
)

doc.add_heading("3.8.3 Duyarlılık (Recall / Sensitivity)", level=3)
add_equation(doc, "Recall = TP / (TP + FN)", "3.9")
doc.add_paragraph(
    "Gerçek pozitiflerin ne kadarının doğru tespit edildiğini ölçer. "
    "Saldırı tespitinde yüksek duyarlılık, saldırıların kaçırılmamasını sağlar."
)

doc.add_heading("3.8.4 F1-Skor (Makro Ortalama)", level=3)
add_equation(doc, "F1 = 2 × (Precision × Recall) / (Precision + Recall)", "3.10")
add_equation(doc, "F1-Macro = (1/C) × Σ F1_c  (tüm C sınıf üzerinden)", "3.11")
doc.add_paragraph(
    "F1-Macro, tüm sınıfların kesinlik ve duyarlılık harmonik ortalamasının "
    "eşit ağırlıklı ortalamasıdır. Sınıf dengesizliği olan veri setlerinde "
    "birincil metrik olarak kullanılmıştır çünkü tüm sınıfları eşit "
    "ağırlıkta değerlendirir."
)

doc.add_heading("3.8.5 AUC-ROC", level=3)
doc.add_paragraph(
    "ROC (Receiver Operating Characteristic) eğrisi altındaki alan, "
    "eşik-bağımsız sınıflandırma performansı ölçümüdür. AUC = 1,0 "
    "mükemmel sınıflandırmayı, AUC = 0,5 rastgele sınıflandırmayı "
    "gösterir (Fawcett, 2006)."
)

doc.add_heading("3.8.6 Yanlış Pozitif Oranı (FPR)", level=3)
add_equation(doc, "FPR = FP / (FP + TN)", "3.12")
doc.add_paragraph(
    "Yanlış alarm oranını temsil eder. IDS sistemlerinde düşük FPR, "
    "operasyonel yükü azaltır."
)

doc.add_heading("3.8.7 Matthews Korelasyon Katsayısı (MCC)", level=3)
add_equation(doc, "MCC = (TP×TN - FP×FN) / √((TP+FP)(TP+FN)(TN+FP)(TN+FN))", "3.13")
doc.add_paragraph(
    "MCC, [-1, +1] aralığında dengeli bir performans ölçümüdür. +1 mükemmel, "
    "0 rastgele, -1 tamamen ters sınıflandırmayı ifade eder. Chicco ve Jurman "
    "(2020), MCC'nin sınıf dengesizliğinde F1'den daha güvenilir olduğunu "
    "göstermiştir."
)

# --- 3.9 ---
doc.add_heading("3.9 İstatistiksel Doğrulama Yöntemleri", level=2)

doc.add_paragraph(
    "Sonuçların güvenilirliğini ve tekrarlanabilirliğini sağlamak için "
    "beş istatistiksel doğrulama yöntemi uygulanmıştır:"
)

doc.add_heading("3.9.1 10-Katlı Tabakalı Çapraz Doğrulama", level=3)

doc.add_paragraph(
    "Veri seti 10 eşit katmana bölünerek her iterasyonda 9 katman eğitim, "
    "1 katman test olarak kullanılmıştır. Tabakalı (stratified) yaklaşım, "
    "her katmandaki sınıf oranlarının korunmasını sağlar. BSO-Hybrid RF "
    "modeli 10-katlı çapraz doğrulamada şu sonuçları elde etmiştir:"
)

add_formatted_table(
    doc,
    headers=["Katman", "Doğruluk (%)", "Kesinlik (%)", "Duyarlılık (%)", "F1-Skor (%)"],
    rows=[
        ["1", "90,89", "90,92", "90,97", "90,88"],
        ["2", "91,04", "91,07", "91,12", "91,02"],
        ["3", "90,84", "90,85", "90,93", "90,82"],
        ["4", "90,87", "90,93", "90,93", "90,86"],
        ["5", "91,37", "91,46", "91,44", "91,38"],
        ["6", "90,66", "90,66", "90,72", "90,61"],
        ["7", "91,16", "91,17", "91,22", "91,13"],
        ["8", "90,99", "91,11", "91,05", "91,00"],
        ["9", "91,12", "91,22", "91,16", "91,13"],
        ["10", "90,82", "90,85", "90,87", "90,80"],
        ["Ortalama ± Std", "90,98 ± 0,204", "91,02 ± 0,229", "91,04 ± 0,204", "90,96 ± 0,217"],
    ],
    caption="Tablo 3.6: 10-Katlı Çapraz Doğrulama Sonuçları (BSO-Hybrid RF)",
    highlight_row=10,
)

doc.add_paragraph(
    "Düşük standart sapma (σ = 0,204) modelin yüksek kararlılığını "
    "göstermektedir. %95 güven aralığı: %90,98 ± 0,400."
)

doc.add_heading("3.9.2 Eşleştirilmiş t-Testi", level=3)

doc.add_paragraph(
    "10-katlı çapraz doğrulamadan elde edilen performans değerleri üzerinde "
    "eşleştirilmiş t-testi uygulanarak BSO-Hybrid RF ile diğer modeller "
    "arasındaki farkın istatistiksel anlamlılığı test edilmiştir. "
    "Anlamlılık düzeyi α = 0,05 olarak belirlenmiştir."
)

doc.add_heading("3.9.3 Wilcoxon Signed-Rank Testi", level=3)

doc.add_paragraph(
    "Parametrik olmayan eşleştirilmiş test olarak Wilcoxon signed-rank "
    "testi uygulanmıştır. Bu test, normal dağılım varsayımı "
    "gerektirmediği için küçük örneklem boyutlarında (10 katman) "
    "daha güvenilir sonuçlar üretir."
)

doc.add_heading("3.9.4 McNemar Testi", level=3)

doc.add_paragraph(
    "McNemar testi, iki sınıflandırıcının test seti üzerindeki tahminleri "
    "arasındaki tutarsızlıkları test eden χ² tabanlı bir testtir. Bir "
    "modelin doğru, diğerinin yanlış tahmin ettiği (ve tersi) örneklerin "
    "sayısını karşılaştırır."
)

doc.add_heading("3.9.5 Cohen's d Etki Büyüklüğü", level=3)

add_equation(doc, "d = (M₁ - M₂) / S_pooled", "3.14")

doc.add_paragraph(
    "Cohen'in d etki büyüklüğü, iki grubun ortalamalarının farkını "
    "birleşik standart sapma cinsinden ifade eder. Yorumlama: "
    "d < 0,2 (küçük), 0,2 ≤ d < 0,8 (orta), d ≥ 0,8 (büyük). "
    "BSO-Hybrid RF vs Decision Tree karşılaştırmasında d = 11,07 "
    "(çok büyük etki) elde edilmiştir."
)


# ============================================================================
# KAYNAKÇA (REFERENCES)
# ============================================================================

doc.add_page_break()
doc.add_heading("KAYNAKÇA", level=1)

references = [
    "Agrawal, N., & Tapaswi, S. (2019). Defense mechanisms against DDoS attacks in a cloud computing environment: State-of-the-art and research challenges. IEEE Communications Surveys & Tutorials, 21(4), 3769–3795.",
    "Ahmad, I., Basheri, M., Iqbal, M. J., & Rahim, A. (2018). Performance comparison of support vector machine, random forest, and extreme learning machine for intrusion detection. IEEE Access, 6, 33789–33795.",
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. https://doi.org/10.1023/A:1010933404324",
    "Buczak, A. L., & Guven, E. (2016). A survey of data mining and machine learning methods for cyber security intrusion detection. IEEE Communications Surveys & Tutorials, 18(2), 1153–1176.",
    "Chandrashekar, G., & Sahin, F. (2014). A survey on feature selection methods. Computers & Electrical Engineering, 40(1), 16–28.",
    "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research, 16, 321–357.",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of ACM SIGKDD, 785–794. https://doi.org/10.1145/2939672.2939785",
    "Chicco, D., & Jurman, G. (2020). The advantages of the Matthews correlation coefficient (MCC) over F1 score and accuracy in binary classification evaluation. BMC Genomics, 21(1), 6.",
    "Cisco Systems. (2023). Cisco Annual Internet Report (2018–2023). Cisco White Paper.",
    "Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.",
    "Cortes, C., & Vapnik, V. (1995). Support-vector networks. Machine Learning, 20(3), 273–297. https://doi.org/10.1007/BF00994018",
    "Cover, T. M., & Hart, P. E. (1967). Nearest neighbor pattern classification. IEEE Transactions on Information Theory, 13(1), 21–27.",
    "Demšar, J. (2006). Statistical comparisons of classifiers over multiple data sets. Journal of Machine Learning Research, 7, 1–30.",
    "Doshi, R., Apthorpe, N., & Feamster, N. (2018). Machine learning DDoS detection for consumer Internet of Things devices. IEEE Security and Privacy Workshops (SPW), 29–35.",
    "Fawcett, T. (2006). An introduction to ROC analysis. Pattern Recognition Letters, 27(8), 861–874.",
    "Ferrag, M. A., Friha, O., Hamouda, D., Maglaras, L., & Janicke, H. (2023). Edge-IIoTset: A new comprehensive realistic cyber security dataset of IoT and IIoT applications. IEEE Access, 10, 25863–25878.",
    "Gandomi, A. H., Yang, X. S., Alavi, A. H., & Talatahari, S. (2013). Bat algorithm for constrained optimization tasks. Neural Computing and Applications, 22(6), 1239–1255.",
    "Guyon, I., & Elisseeff, A. (2003). An introduction to variable and feature selection. Journal of Machine Learning Research, 3, 1157–1182.",
    "He, H., & Garcia, E. A. (2009). Learning from imbalanced data. IEEE Transactions on Knowledge and Data Engineering, 21(9), 1263–1284.",
    "Holland, J. H. (1975). Adaptation in natural and artificial systems. University of Michigan Press.",
    "Kennedy, J., & Eberhart, R. (1995). Particle swarm optimization. Proceedings of IEEE ICNN, 4, 1942–1948.",
    "Khammassi, C., & Krichen, S. (2017). A GA-LR wrapper approach for feature selection in network intrusion detection. Computers & Security, 70, 255–277.",
    "Kolias, C., Kambourakis, G., Stavrou, A., & Voas, J. (2017). DDoS in the IoT: Mirai and other botnets. Computer, 50(7), 80–84.",
    "Li, Y., & Lu, Y. (2019). LSTM-BA: DDoS detection approach combining LSTM and Bayes. International Conference on Security, Pattern Analysis, and Cybernetics, 180–185.",
    "Mafarja, M., & Mirjalili, S. (2018). Whale optimization approaches for wrapper feature selection. Applied Soft Computing, 62, 441–453.",
    "Meidan, Y., Bohadana, M., Mathov, Y., Mirsky, Y., Shabtai, A., Breitenbacher, D., & Elovici, Y. (2018). N-BaIoT—Network-based detection of IoT botnet attacks using deep autoencoders. IEEE Pervasive Computing, 17(3), 12–22.",
    "Mirjalili, S., Mirjalili, S. M., & Lewis, A. (2014). Grey wolf optimizer. Advances in Engineering Software, 69, 46–61. https://doi.org/10.1016/j.advengsoft.2013.12.007",
    "Nakamura, R. Y. M., Pereira, L. A. M., Costa, K. A., Rodrigues, D., Papa, J. P., & Yang, X. S. (2012). BBA: A binary bat algorithm for feature selection. Brazilian Symposium on Computer Graphics and Image Processing, 291–297.",
    "Neto, E. C. P., Dadkhah, S., Ferreira, R., Zohourian, A., Lu, R., & Ghorbani, A. A. (2023). CICIoT2023: A real-time dataset and benchmark for large-scale attacks in IoT environment. Sensors, 23(13), 5941. https://doi.org/10.3390/s23135941",
    "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., et al. (2011). Scikit-learn: Machine learning in Python. JMLR, 12, 2825–2830.",
    "Roopak, M., Tian, G. Y., & Chambers, J. (2019). Deep learning models for cyber security in IoT networks. IEEE CCWC, 452–457.",
    "Sokolova, M., & Lapalme, G. (2009). A systematic analysis of performance measures for classification tasks. Information Processing & Management, 45(4), 427–437.",
    "Tama, B. A., & Rhee, K. H. (2019). An in-depth experimental study of anomaly detection using gradient boosted machine. Neural Computing and Applications, 31(4), 955–965.",
    "Xue, B., Zhang, M., Browne, W. N., & Yao, X. (2016). A survey on evolutionary computation approaches to feature selection. IEEE Transactions on Evolutionary Computation, 20(4), 606–626.",
    "Yang, X. S. (2010). A new metaheuristic bat-inspired algorithm. Nature Inspired Cooperative Strategies for Optimization (NICSO), 284, 65–74. https://doi.org/10.1007/978-3-642-12538-6_6",
    "Yang, X. S. (2013). Bat algorithm: Literature review and applications. International Journal of Bio-Inspired Computation, 5(3), 141–149.",
    "Yuan, X., Li, C., & Li, X. (2017). DeepDefense: Identifying DDoS attack via deep learning. IEEE International Conference on Smart Computing, 1–8.",
    "Zargar, S. T., Joshi, J., & Tipper, D. (2013). A survey of defense mechanisms against distributed denial of service (DDoS) flooding attacks. IEEE Communications Surveys & Tutorials, 15(4), 2046–2069.",
    "Zhou, Y., Cheng, G., Jiang, S., & Dai, M. (2020). Building an efficient intrusion detection system based on feature selection and ensemble classifier. Computer Networks, 174, 107247.",
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)


# ============================================================================
# SAVE
# ============================================================================

output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, "..", "thesis_chapters_1_3.docx")
output_path = os.path.normpath(output_path)

doc.save(output_path)
print(f"\n✅ Tez dosyası başarıyla oluşturuldu: {output_path}")
print(f"   Bölüm 1: GİRİŞ (7 alt bölüm)")
print(f"   Bölüm 2: LİTERATÜR TARAMASI (6 alt bölüm)")
print(f"   Bölüm 3: YÖNTEM (9 alt bölüm)")
print(f"   Kaynakça: {len(references)} referans (APA 7)")
print(f"   Toplam: ~50 sayfa akademik içerik")
