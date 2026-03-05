#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
================================================================================
BSO-Hybrid RF Dashboard — Comprehensive Word Export
================================================================================
Exports the ENTIRE dashboard content to a Word document with all tables,
figures, charts, and text — exactly mirroring the web dashboard.

Author: SHUAIB AYAD JASIM
Advisor: Dr. Saim Ervural — KTO Karatay Üniversitesi
"""

import os, sys, math, io, warnings
from pathlib import Path

warnings.filterwarnings("ignore")

import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.gridspec import GridSpec

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── output paths ──────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent.parent
OUT_DIR = BASE_DIR / "output"
FIG_DIR = OUT_DIR / "dashboard_figures"
OUT_DIR.mkdir(exist_ok=True)
FIG_DIR.mkdir(exist_ok=True)
OUT_DOCX = OUT_DIR / "DASHBOARD_TAM_EXPORT.docx"

# ── colour palette ────────────────────────────────────────────────────────────
C_PRIMARY = "#2563eb"
C_SUCCESS = "#22c55e"
C_DANGER  = "#ef4444"
C_WARNING = "#f59e0b"
C_PURPLE  = "#8b5cf6"
C_ORANGE  = "#f97316"
C_TEAL    = "#14b8a6"
C_PINK    = "#ec4899"
C_INDIGO  = "#6366f1"
C_GRAY    = "#64748b"
MODEL_COLORS = [C_PRIMARY, C_DANGER, C_WARNING, C_SUCCESS, C_PURPLE, C_ORANGE,
                C_TEAL, C_PINK, C_INDIGO, C_GRAY, "#0ea5e9", "#84cc16"]

plt.rcParams.update({
    "figure.dpi": 180,
    "savefig.dpi": 180,
    "font.size": 9,
    "axes.titlesize": 11,
    "axes.labelsize": 9,
    "xtick.labelsize": 8,
    "ytick.labelsize": 8,
    "legend.fontsize": 7,
    "figure.facecolor": "white",
})

# =============================================================================
# DATA (mirrors lib/ciciot2023-dataset.ts)
# =============================================================================

FEATURES_39 = [
    "Header_Length","Protocol Type","Time_To_Live","Rate",
    "fin_flag_number","syn_flag_number","rst_flag_number","psh_flag_number",
    "ack_flag_number","ece_flag_number","cwr_flag_number","ack_count",
    "syn_count","fin_count","rst_count","HTTP","HTTPS","DNS","Telnet",
    "SMTP","SSH","IRC","TCP","UDP","DHCP","ARP","ICMP","IGMP","IPv",
    "LLC","Tot sum","Min","Max","AVG","Std","Tot size","IAT","Number","Variance",
]

ATTACK_TYPES = [
    {"name":"Backdoor_Malware","trainingSamples":2252,"testingSamples":644,"color":"#ef4444","smoteSamples":17500},
    {"name":"BenignTraffic","trainingSamples":17500,"testingSamples":5000,"color":"#22c55e","smoteSamples":17500},
    {"name":"DDoS-ACK_Fragmentation","trainingSamples":17500,"testingSamples":5000,"color":"#f59e0b","smoteSamples":17500},
    {"name":"DDoS-SYN_Flood","trainingSamples":17500,"testingSamples":5000,"color":"#3b82f6","smoteSamples":17500},
    {"name":"Recon-PortScan","trainingSamples":17500,"testingSamples":5000,"color":"#8b5cf6","smoteSamples":17500},
]

DATASET_STATS = dict(
    totalFeatures=39, selectedFeatures=19, featureReductionPct=51.3,
    totalSamples=118466, classes=5,
    training=87500, validation=10322, testing=20644,
    smoteSyntheticSamples=15248, originalMinorityCount=2252, balancedClassCount=17500,
)

BSO_SELECTED = [
    (1,"syn_count",0.224480,12),(2,"Number",0.183394,37),(3,"Tot sum",0.154063,30),
    (4,"Rate",0.105115,3),(5,"Max",0.085952,32),(6,"Header_Length",0.066085,0),
    (7,"HTTPS",0.051489,16),(8,"Time_To_Live",0.045447,2),(9,"psh_flag_number",0.020764,7),
    (10,"HTTP",0.019776,15),(11,"fin_flag_number",0.012970,4),(12,"UDP",0.012883,23),
    (13,"DNS",0.008363,17),(14,"ARP",0.003231,25),(15,"LLC",0.002768,29),
    (16,"SSH",0.001919,20),(17,"DHCP",0.000821,24),(18,"IGMP",0.000344,27),
    (19,"cwr_flag_number",0.000135,10),
]

MODELS = [
    {"name":"BSO-Hybrid RF (Proposed)","accuracy":89.82,"precision":90.19,"recall":89.82,"f1Score":89.90,"f1Macro":84.24,"aucRoc":98.38,"trainingTime":3.742,"predictionTime":0.0043,"featuresUsed":19,"fpr":9.81,"fnr":10.18,"specificity":90.19,"mcc":0.8676,"featureSet":"BSO"},
    {"name":"BSO-SVM","accuracy":82.19,"precision":88.99,"recall":82.19,"f1Score":84.57,"f1Macro":74.03,"aucRoc":97.0,"trainingTime":5.485,"predictionTime":0.0009,"featuresUsed":19,"fpr":11.01,"fnr":17.81,"specificity":88.99,"mcc":0.7808,"featureSet":"BSO"},
    {"name":"PSO-RF","accuracy":88.35,"precision":88.54,"recall":88.35,"f1Score":88.39,"f1Macro":81.82,"aucRoc":97.87,"trainingTime":2.601,"predictionTime":0.0071,"featuresUsed":18,"fpr":11.46,"fnr":11.65,"specificity":88.54,"mcc":0.848,"featureSet":"PSO"},
    {"name":"GA-RF","accuracy":89.37,"precision":89.53,"recall":89.37,"f1Score":89.41,"f1Macro":83.66,"aucRoc":98.24,"trainingTime":2.816,"predictionTime":0.006,"featuresUsed":21,"fpr":10.47,"fnr":10.63,"specificity":89.53,"mcc":0.8614,"featureSet":"GA"},
    {"name":"GWO-RF","accuracy":89.80,"precision":90.06,"recall":89.80,"f1Score":89.87,"f1Macro":84.35,"aucRoc":98.37,"trainingTime":3.508,"predictionTime":0.006,"featuresUsed":23,"fpr":9.94,"fnr":10.20,"specificity":90.06,"mcc":0.8671,"featureSet":"GWO"},
    {"name":"Random Forest","accuracy":89.74,"precision":89.92,"recall":89.74,"f1Score":89.77,"f1Macro":84.13,"aucRoc":98.36,"trainingTime":4.52,"predictionTime":0.0068,"featuresUsed":39,"fpr":10.08,"fnr":10.26,"specificity":89.92,"mcc":0.8665,"featureSet":"All"},
    {"name":"SVM (Linear)","accuracy":83.11,"precision":89.20,"recall":83.11,"f1Score":85.21,"f1Macro":75.09,"aucRoc":97.25,"trainingTime":68.244,"predictionTime":0.0011,"featuresUsed":39,"fpr":10.80,"fnr":16.89,"specificity":89.20,"mcc":0.7928,"featureSet":"All"},
    {"name":"Decision Tree","accuracy":86.12,"precision":86.47,"recall":86.12,"f1Score":86.27,"f1Macro":78.57,"aucRoc":91.2,"trainingTime":1.007,"predictionTime":0.0003,"featuresUsed":39,"fpr":13.53,"fnr":13.88,"specificity":86.47,"mcc":0.8212,"featureSet":"All"},
    {"name":"K-Nearest Neighbors","accuracy":85.20,"precision":88.14,"recall":85.20,"f1Score":86.19,"f1Macro":77.16,"aucRoc":95.61,"trainingTime":0.006,"predictionTime":0.0514,"featuresUsed":39,"fpr":11.86,"fnr":14.80,"specificity":88.14,"mcc":0.8098,"featureSet":"All"},
    {"name":"Naive Bayes","accuracy":62.96,"precision":89.44,"recall":62.96,"f1Score":66.18,"f1Macro":57.17,"aucRoc":96.09,"trainingTime":0.048,"predictionTime":0.0007,"featuresUsed":39,"fpr":10.56,"fnr":37.04,"specificity":89.44,"mcc":0.5852,"featureSet":"All"},
    {"name":"Logistic Regression","accuracy":82.73,"precision":89.26,"recall":82.73,"f1Score":84.98,"f1Macro":74.80,"aucRoc":97.3,"trainingTime":3.742,"predictionTime":0.0001,"featuresUsed":39,"fpr":10.74,"fnr":17.27,"specificity":89.26,"mcc":0.7876,"featureSet":"All"},
    {"name":"XGBoost","accuracy":90.37,"precision":90.45,"recall":90.37,"f1Score":90.35,"f1Macro":84.74,"aucRoc":98.51,"trainingTime":4.688,"predictionTime":0.002,"featuresUsed":39,"fpr":9.55,"fnr":9.63,"specificity":90.45,"mcc":0.8742,"featureSet":"All"},
]

CONFUSION_MATRICES = {
    "BSO-RF":[[421,122,0,0,101],[163,4280,0,0,557],[0,0,5000,0,0],[0,0,4,4996,0],[239,916,0,0,3845]],
    "BSO-SVM":[[398,153,0,0,93],[1115,3502,0,0,383],[0,1,4995,0,4],[0,0,7,4992,1],[1151,768,0,0,3081]],
    "PSO-RF":[[349,156,0,0,139],[157,4128,0,0,715],[0,0,4999,0,1],[0,0,4,4996,0],[227,1007,0,0,3766]],
    "GA-RF":[[386,128,0,0,130],[140,4200,0,0,660],[0,0,4998,0,2],[0,0,3,4994,3],[198,929,1,0,3872]],
    "GWO-RF":[[415,116,0,0,113],[139,4253,0,0,608],[2,0,4997,0,1],[0,0,3,4997,0],[226,897,0,0,3877]],
    "RF":[[394,128,0,0,122],[130,4263,0,0,607],[0,0,5000,0,0],[0,0,4,4996,0],[205,923,0,0,3872]],
    "SVM":[[426,135,0,0,83],[1063,3551,0,0,386],[0,0,4999,0,1],[0,0,4,4995,1],[1007,807,0,0,3186]],
    "DT":[[310,161,0,0,173],[230,3770,0,0,1000],[0,0,4995,4,1],[0,0,4,4996,0],[289,1003,0,0,3708]],
    "KNN":[[396,147,0,0,101],[612,3836,0,0,552],[1,0,4997,0,2],[0,0,3,4997,0],[653,985,0,0,3362]],
    "NB":[[624,16,0,0,4],[4448,487,0,0,65],[1,0,4998,1,0],[0,0,6,4994,0],[2931,174,1,0,1894]],
    "LR":[[434,132,0,0,78],[1098,3511,0,0,391],[0,0,4999,0,1],[0,0,4,4995,1],[1089,771,0,0,3140]],
    "XGBoost":[[376,133,0,0,135],[113,4341,0,0,546],[0,0,4999,0,1],[0,0,4,4996,0],[160,897,0,0,3943]],
}

CLASS_LABELS = ["Backdoor_Malware","BenignTraffic","DDoS-ACK_Frag","DDoS-SYN_Flood","Recon-PortScan"]

BSO_RF_PER_CLASS = [
    {"className":"Backdoor_Malware","precision":51.15,"recall":65.37,"f1Score":57.40,"support":644},
    {"className":"BenignTraffic","precision":80.48,"recall":85.60,"f1Score":82.96,"support":5000},
    {"className":"DDoS-ACK_Fragmentation","precision":99.92,"recall":100.0,"f1Score":99.96,"support":5000},
    {"className":"DDoS-SYN_Flood","precision":100.0,"recall":99.92,"f1Score":99.96,"support":5000},
    {"className":"Recon-PortScan","precision":85.39,"recall":76.90,"f1Score":80.92,"support":5000},
]

CROSS_VALIDATION = {
    "folds": [
        {"fold":1,"accuracy":90.89,"precision":90.92,"recall":90.97,"f1Score":90.88},
        {"fold":2,"accuracy":91.04,"precision":91.07,"recall":91.12,"f1Score":91.02},
        {"fold":3,"accuracy":90.84,"precision":90.85,"recall":90.93,"f1Score":90.82},
        {"fold":4,"accuracy":90.87,"precision":90.93,"recall":90.93,"f1Score":90.86},
        {"fold":5,"accuracy":91.37,"precision":91.46,"recall":91.44,"f1Score":91.38},
        {"fold":6,"accuracy":90.66,"precision":90.66,"recall":90.72,"f1Score":90.61},
        {"fold":7,"accuracy":91.16,"precision":91.17,"recall":91.22,"f1Score":91.13},
        {"fold":8,"accuracy":90.99,"precision":91.11,"recall":91.05,"f1Score":91.00},
        {"fold":9,"accuracy":91.12,"precision":91.22,"recall":91.16,"f1Score":91.13},
        {"fold":10,"accuracy":90.82,"precision":90.85,"recall":90.87,"f1Score":90.80},
    ],
    "mean":{"accuracy":90.98,"precision":91.02,"recall":91.04,"f1Score":90.96},
    "std":{"accuracy":0.204,"precision":0.229,"recall":0.204,"f1Score":0.217},
}

STATISTICAL_TESTS = [
    {"comparison":"BSO-Hybrid vs Random Forest","improvement":"-0.45%","tStatistic":-9.7019,"pValue":0.000005,"significant":True,"cohenD":-3.234,"effectSize":"large","wilcoxonP":0.001953},
    {"comparison":"BSO-Hybrid vs Decision Tree","improvement":"+3.62%","tStatistic":33.1978,"pValue":0.0,"significant":True,"cohenD":11.0659,"effectSize":"large","wilcoxonP":0.001953},
    {"comparison":"BSO-Hybrid vs Logistic Regression","improvement":"+10.51%","tStatistic":86.4492,"pValue":0.0,"significant":True,"cohenD":28.8164,"effectSize":"large","wilcoxonP":0.001953},
    {"comparison":"BSO-Hybrid vs XGBoost","improvement":"-0.55%","tStatistic":-4.2814,"pValue":0.002038,"significant":True,"cohenD":-1.427,"effectSize":"large","wilcoxonP":0.003906},
    {"comparison":"BSO-Hybrid vs PSO-RF","improvement":"+1.47%","tStatistic":15.3264,"pValue":0.0,"significant":True,"cohenD":5.1088,"effectSize":"large","wilcoxonP":0.001953},
    {"comparison":"BSO-Hybrid vs GA-RF","improvement":"+0.45%","tStatistic":6.8421,"pValue":0.000071,"significant":True,"cohenD":2.2807,"effectSize":"large","wilcoxonP":0.001953},
    {"comparison":"BSO-Hybrid vs GWO-RF","improvement":"+0.02%","tStatistic":0.3142,"pValue":0.7609,"significant":False,"cohenD":0.1047,"effectSize":"small","wilcoxonP":0.625},
    {"comparison":"BSO-Hybrid vs BSO-SVM","improvement":"+7.63%","tStatistic":52.8736,"pValue":0.0,"significant":True,"cohenD":17.6245,"effectSize":"large","wilcoxonP":0.001953},
    {"comparison":"BSO-Hybrid vs KNN","improvement":"+4.62%","tStatistic":38.4217,"pValue":0.0,"significant":True,"cohenD":12.8072,"effectSize":"large","wilcoxonP":0.001953},
    {"comparison":"BSO-Hybrid vs Naive Bayes","improvement":"+26.86%","tStatistic":142.8531,"pValue":0.0,"significant":True,"cohenD":47.6177,"effectSize":"large","wilcoxonP":0.001953},
    {"comparison":"BSO-Hybrid vs SVM (Linear)","improvement":"+6.71%","tStatistic":48.2163,"pValue":0.0,"significant":True,"cohenD":16.0721,"effectSize":"large","wilcoxonP":0.001953},
]

BSO_CONVERGENCE = [
    (0,0.184825,0.200021,27),(1,0.183088,0.191476,25),(2,0.181799,0.194135,21),
    (3,0.181435,0.190601,21),(4,0.181129,0.200868,22),(5,0.181129,0.195327,22),
    (6,0.181129,0.202679,22),(7,0.181129,0.211452,22),(8,0.181014,0.200518,21),
    (9,0.180983,0.195750,20),(10,0.180983,0.198495,20),(11,0.180673,0.195257,21),
    (12,0.180673,0.194536,21),(13,0.180673,0.200093,21),(14,0.180673,0.201612,21),
    (15,0.179905,0.190819,21),(16,0.179905,0.199296,21),(17,0.179905,0.194277,21),
    (18,0.179757,0.199849,20),(19,0.179757,0.204054,20),(20,0.179427,0.192536,19),
    (21,0.179427,0.193079,19),(22,0.179427,0.191688,19),(23,0.179427,0.191101,19),
    (24,0.179427,0.187818,19),(25,0.179427,0.187170,19),(26,0.179427,0.185475,19),
    (27,0.179427,0.194779,19),(28,0.179427,0.191739,19),(29,0.179427,0.192119,19),
    (30,0.179427,0.188633,19),(31,0.179427,0.186241,19),(32,0.179427,0.190814,19),
    (33,0.179427,0.186654,19),(34,0.179427,0.185787,19),(35,0.179427,0.193499,19),
    (36,0.179301,0.186727,20),(37,0.179301,0.188759,20),(38,0.179301,0.182090,20),
    (39,0.179301,0.184430,20),(40,0.179301,0.183888,20),(41,0.179301,0.187773,20),
    (42,0.179301,0.186481,20),(43,0.179301,0.184072,20),(44,0.177801,0.184230,19),
    (45,0.177801,0.189771,19),(46,0.177801,0.189608,19),(47,0.177801,0.185909,19),
    (48,0.177801,0.193035,19),(49,0.177801,0.187681,19),
]

OPTIMIZER_CONV = {
    "BSO":[0.184825,0.183088,0.181799,0.181435,0.181129,0.181129,0.181129,0.181129,0.181014,0.180983,0.180983,0.180673,0.180673,0.180673,0.180673,0.179905,0.179905,0.179905,0.179757,0.179757,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179427,0.179301,0.179301,0.179301,0.179301,0.179301,0.179301,0.179301,0.179301,0.177801,0.177801,0.177801,0.177801,0.177801,0.177801],
    "PSO":[0.199141,0.199141,0.199141,0.197860,0.197860,0.197693,0.197693,0.197693,0.197693,0.197693,0.197490,0.197490,0.197490,0.197490,0.197490,0.197490,0.197490,0.197490,0.195505,0.195505,0.195505,0.195505,0.195505,0.195505,0.195505,0.195505,0.195505,0.195505,0.195505,0.195290,0.195290,0.195290,0.195290,0.195290,0.195290,0.195290,0.195290,0.195290,0.195290,0.193895],
    "GA":[0.199332,0.193521,0.190793,0.190793,0.190793,0.190793,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982,0.188982],
    "GWO":[0.196093,0.193160,0.193160,0.193160,0.193160,0.193160,0.193160,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.193146,0.192181,0.192181,0.192181,0.192181,0.192181,0.192181,0.192181,0.192181],
}

FEATURE_SELECTION_COMP = {
    "BSO":{"nSelected":19,"reductionPct":51.3,"bestFitness":0.177801,"evaluations":1177,"time":840.43,
           "features":["Header_Length","Time_To_Live","Rate","fin_flag_number","psh_flag_number","cwr_flag_number","syn_count","HTTP","HTTPS","DNS","SSH","UDP","DHCP","ARP","IGMP","LLC","Tot sum","Max","Number"]},
    "PSO":{"nSelected":18,"reductionPct":53.8,"bestFitness":0.193895,"evaluations":820,"time":68.91,
           "features":["Time_To_Live","fin_flag_number","syn_flag_number","rst_flag_number","psh_flag_number","ack_flag_number","ack_count","HTTP","HTTPS","DNS","SMTP","SSH","IRC","TCP","UDP","Min","IAT","Number"]},
    "GA":{"nSelected":21,"reductionPct":46.2,"bestFitness":0.188982,"evaluations":722,"time":60.47,
          "features":["Protocol Type","Time_To_Live","psh_flag_number","ece_flag_number","ack_count","syn_count","fin_count","HTTP","HTTPS","DNS","SSH","IRC","TCP","DHCP","ARP","IPv","LLC","Min","Max","Std","IAT"]},
    "GWO":{"nSelected":23,"reductionPct":41.0,"bestFitness":0.192181,"evaluations":820,"time":70.23,
           "features":["Time_To_Live","Rate","fin_flag_number","syn_flag_number","psh_flag_number","ack_flag_number","ece_flag_number","fin_count","rst_count","HTTP","HTTPS","DNS","Telnet","SMTP","SSH","DHCP","IPv","LLC","Tot sum","Max","AVG","Std","IAT"]},
}

BSO_PARAMS = dict(
    populationSize=25, maxIterations=50, frequencyMin=0.0, frequencyMax=2.0,
    initialLoudness=0.95, initialPulseRate=0.5, alpha=0.9, gamma=0.9,
    dimensions=39, selectedDimensions=19, convergenceThreshold=1e-6,
    fitnessFunction="1 - F1_macro + 0.01 * (n_selected / n_total)",
    n_estimators=266, max_depth=20, min_samples_split=7, max_features=0.469,
    totalEvaluations=1177, optimizationTime=840.43,
)

DYNAMIC_ENV = {
    "noiseRobustness":[
        {"noiseLevel":0,"accuracy":89.82,"f1Macro":84.24,"degradation":0},
        {"noiseLevel":0.05,"accuracy":65.73,"f1Macro":55.12,"degradation":24.09},
        {"noiseLevel":0.10,"accuracy":63.32,"f1Macro":51.17,"degradation":26.50},
        {"noiseLevel":0.15,"accuracy":62.25,"f1Macro":49.12,"degradation":27.57},
        {"noiseLevel":0.20,"accuracy":60.37,"f1Macro":46.73,"degradation":29.45},
        {"noiseLevel":0.25,"accuracy":59.91,"f1Macro":46.12,"degradation":29.91},
        {"noiseLevel":0.30,"accuracy":59.32,"f1Macro":45.26,"degradation":30.50},
    ],
    "unknownAttack":[
        {"excluded":"Backdoor_Malware","detectionRate":65.84,"samples":644},
        {"excluded":"DDoS-ACK_Fragmentation","detectionRate":99.98,"samples":5000},
        {"excluded":"DDoS-SYN_Flood","detectionRate":100.0,"samples":5000},
        {"excluded":"Recon-PortScan","detectionRate":9.86,"samples":5000},
    ],
    "throughput":[
        {"batchSize":100,"avgTimeMs":38.25,"samplesPerSecond":2615},
        {"batchSize":500,"avgTimeMs":48.08,"samplesPerSecond":10400},
        {"batchSize":1000,"avgTimeMs":48.43,"samplesPerSecond":20647},
        {"batchSize":5000,"avgTimeMs":59.09,"samplesPerSecond":84621},
        {"batchSize":20644,"avgTimeMs":89.05,"samplesPerSecond":231833},
    ],
    "learningCurve":[
        {"fraction":0.1,"nSamples":8750,"accuracy":88.13,"f1Macro":81.80,"trainingTime":0.513},
        {"fraction":0.2,"nSamples":17500,"accuracy":88.76,"f1Macro":82.70,"trainingTime":0.759},
        {"fraction":0.3,"nSamples":26250,"accuracy":89.01,"f1Macro":83.16,"trainingTime":1.078},
        {"fraction":0.5,"nSamples":43750,"accuracy":89.41,"f1Macro":83.71,"trainingTime":1.834},
        {"fraction":0.7,"nSamples":61249,"accuracy":89.61,"f1Macro":84.00,"trainingTime":2.663},
        {"fraction":1.0,"nSamples":87500,"accuracy":89.85,"f1Macro":84.34,"trainingTime":4.207},
    ],
}

STATE_OF_THE_ART = [
    {"paper":"Neto et al. (2023)","dataset":"CICIoT2023","method":"Random Forest","accuracy":99.0,"f1Score":98.0,"note":"İkili sınıflandırma (benign vs attack)"},
    {"paper":"Ferrag et al. (2023)","dataset":"CICIoT2023","method":"DNN","accuracy":98.2,"f1Score":97.5,"note":"İkili sınıflandırma, derin öğrenme"},
    {"paper":"Proposed (BSO-Hybrid RF)","dataset":"CICIoT2023","method":"BSO-Hybrid RF","accuracy":89.82,"f1Score":89.90,"note":"5 sınıflı, BSO optimizasyonu, SMOTE, 19 özellik"},
]

ROC_AUC_DATA = {
    "BSO-Hybrid RF (Proposed)": 0.9817,
    "BSO-SVM": 0.9549,
    "PSO-RF": 0.9744,
    "GA-RF": 0.9802,
    "GWO-RF": 0.9818,
    "Random Forest": 0.9816,
    "SVM (Linear)": 0.9606,
    "Decision Tree": 0.8799,
    "KNN": 0.9353,
    "Naive Bayes": 0.9451,
    "Logistic Regression": 0.9614,
    "XGBoost": 0.9829,
}

# =============================================================================
# HELPER: WORD DOCUMENT UTILITIES
# =============================================================================
def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0,2,4))

def set_cell_shading(cell, color_hex):
    r,g,b = hex_to_rgb(color_hex)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.lstrip("#")}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, highlight_row=None):
    """Add a formatted table with header shading."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(255,255,255)
        set_cell_shading(cell, "#1e3a5f")
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)
            if highlight_row is not None and r_idx == highlight_row:
                set_cell_shading(cell, "#dbeafe")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def save_figure(fig, name):
    """Save figure to FIG_DIR and return path."""
    path = FIG_DIR / f"{name}.png"
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return str(path)

def add_figure(doc, path, caption, width=Inches(6.0)):
    """Insert figure with centered caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else None
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80,80,80)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(30,58,95)
    return h

# =============================================================================
# FIGURE GENERATORS
# =============================================================================
fig_counter = [0]

def next_fig():
    fig_counter[0] += 1
    return fig_counter[0]

# 1. Class Distribution (Before & After SMOTE)
def fig_class_distribution():
    names = [a["name"] for a in ATTACK_TYPES]
    before = [a["trainingSamples"] for a in ATTACK_TYPES]
    after  = [a["smoteSamples"] for a in ATTACK_TYPES]
    colors_b = [a["color"] for a in ATTACK_TYPES]
    
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))
    # Before SMOTE
    bars = axes[0].barh(names, before, color=colors_b, edgecolor="white")
    axes[0].set_title("SMOTE Öncesi Sınıf Dağılımı", fontweight="bold")
    axes[0].set_xlabel("Örnek Sayısı")
    for bar, v in zip(bars, before):
        axes[0].text(v + 200, bar.get_y() + bar.get_height()/2, f"{v:,}", va="center", fontsize=8)
    # After SMOTE
    bars2 = axes[1].barh(names, after, color=colors_b, edgecolor="white")
    axes[1].set_title("SMOTE Sonrası Sınıf Dağılımı", fontweight="bold")
    axes[1].set_xlabel("Örnek Sayısı")
    for bar, v in zip(bars2, after):
        axes[1].text(v + 200, bar.get_y() + bar.get_height()/2, f"{v:,}", va="center", fontsize=8)
    fig.tight_layout()
    return save_figure(fig, "01_class_distribution")

# 2. SMOTE Before/After Grouped Bar Chart
def fig_smote_comparison():
    names = [a["name"].replace("DDoS-","").replace("_"," ")[:15] for a in ATTACK_TYPES]
    before = [a["trainingSamples"] for a in ATTACK_TYPES]
    after  = [a["smoteSamples"] for a in ATTACK_TYPES]
    x = np.arange(len(names))
    w = 0.35
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(x - w/2, before, w, label="SMOTE Öncesi", color="#94a3b8", edgecolor="white")
    ax.bar(x + w/2, after, w, label="SMOTE Sonrası", color=C_PRIMARY, edgecolor="white")
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=25, ha="right")
    ax.set_ylabel("Örnek Sayısı")
    ax.set_title("SMOTE Dengeleme Karşılaştırması", fontweight="bold")
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    for i, (b, a) in enumerate(zip(before, after)):
        ax.text(i - w/2, b + 300, f"{b:,}", ha="center", fontsize=7)
        ax.text(i + w/2, a + 300, f"{a:,}", ha="center", fontsize=7)
    fig.tight_layout()
    return save_figure(fig, "02_smote_comparison")

# 3. BSO Feature Importance (Top 19)
def fig_feature_importance():
    names = [f[1] for f in BSO_SELECTED]
    imps  = [f[2] for f in BSO_SELECTED]
    colors = [C_PRIMARY if i < 5 else C_TEAL if i < 10 else C_GRAY for i in range(len(names))]
    fig, ax = plt.subplots(figsize=(10, 7))
    bars = ax.barh(range(len(names)-1, -1, -1), imps, color=colors, edgecolor="white")
    ax.set_yticks(range(len(names)-1, -1, -1))
    ax.set_yticklabels(names)
    ax.set_xlabel("Önem Skoru (BSO-Hybrid RF)")
    ax.set_title("BSO Seçilmiş 19 Özellik — Önem Sıralaması", fontweight="bold")
    ax.grid(axis="x", alpha=0.3)
    for bar, v in zip(bars, imps):
        ax.text(v + 0.002, bar.get_y() + bar.get_height()/2, f"{v:.4f}", va="center", fontsize=7)
    fig.tight_layout()
    return save_figure(fig, "03_feature_importance")

# 4. 39 vs 19 Features Comparison
def fig_39_vs_19():
    metrics = ["Accuracy","Precision","Recall","F1 Score","F1 Macro","AUC-ROC"]
    bso = [89.82, 90.19, 89.82, 89.90, 84.24, 98.38]
    rf_all = [89.74, 89.92, 89.74, 89.77, 84.13, 98.36]
    x = np.arange(len(metrics))
    w = 0.35
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(x - w/2, rf_all, w, label="Random Forest (39 özellik)", color="#94a3b8")
    ax.bar(x + w/2, bso, w, label="BSO-Hybrid RF (19 özellik)", color=C_PRIMARY)
    ax.set_xticks(x)
    ax.set_xticklabels(metrics)
    ax.set_ylabel("Skor (%)")
    ax.set_title("39 Özellik vs 19 Özellik (BSO) Performans Karşılaştırması", fontweight="bold")
    ax.legend()
    ax.set_ylim(80, 100)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return save_figure(fig, "04_39_vs_19_features")

# 5. Model Accuracy Comparison (all 12)
def fig_model_accuracy():
    names = [m["name"].replace(" (Proposed)","") for m in MODELS]
    accs  = [m["accuracy"] for m in MODELS]
    colors = [C_PRIMARY if i==0 else MODEL_COLORS[i % len(MODEL_COLORS)] for i in range(len(MODELS))]
    fig, ax = plt.subplots(figsize=(12, 6))
    bars = ax.bar(range(len(names)), accs, color=colors, edgecolor="white")
    ax.set_xticks(range(len(names)))
    ax.set_xticklabels(names, rotation=45, ha="right", fontsize=7)
    ax.set_ylabel("Accuracy (%)")
    ax.set_title("12 Model — Accuracy Karşılaştırması", fontweight="bold")
    ax.set_ylim(55, 95)
    ax.grid(axis="y", alpha=0.3)
    for bar, v in zip(bars, accs):
        ax.text(bar.get_x() + bar.get_width()/2, v + 0.3, f"{v:.2f}%", ha="center", fontsize=7)
    fig.tight_layout()
    return save_figure(fig, "05_model_accuracy")

# 6. Radar Chart — Top 4 Models
def fig_radar_top4():
    top4 = [MODELS[0], MODELS[4], MODELS[5], MODELS[11]]  # BSO-RF, GWO-RF, RF, XGBoost
    metrics = ["Accuracy","Precision","Recall","F1","AUC-ROC","MCC×100"]
    angles = np.linspace(0, 2*np.pi, len(metrics), endpoint=False).tolist()
    angles += angles[:1]
    
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
    colors_r = [C_PRIMARY, C_PURPLE, C_ORANGE, C_SUCCESS]
    for i, m in enumerate(top4):
        vals = [m["accuracy"], m["precision"], m["recall"], m["f1Score"], m["aucRoc"], m["mcc"]*100]
        vals += vals[:1]
        ax.plot(angles, vals, "o-", linewidth=2, label=m["name"].replace(" (Proposed)",""), color=colors_r[i])
        ax.fill(angles, vals, alpha=0.1, color=colors_r[i])
    ax.set_thetagrids(np.degrees(angles[:-1]), metrics)
    ax.set_ylim(80, 100)
    ax.set_title("En İyi 4 Model — Radar Karşılaştırması", fontweight="bold", y=1.08)
    ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.1), fontsize=8)
    fig.tight_layout()
    return save_figure(fig, "06_radar_top4")

# 7. BSO Convergence Curve
def fig_bso_convergence():
    iters = [d[0] for d in BSO_CONVERGENCE]
    best  = [d[1] for d in BSO_CONVERGENCE]
    avg   = [d[2] for d in BSO_CONVERGENCE]
    nfeat = [d[3] for d in BSO_CONVERGENCE]
    
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 8), sharex=True)
    ax1.plot(iters, best, "b-o", markersize=3, label="En İyi Fitness", linewidth=2)
    ax1.plot(iters, avg, "r--", alpha=0.6, label="Ortalama Fitness", linewidth=1)
    ax1.set_ylabel("Fitness Değeri")
    ax1.set_title("BSO-Hybrid Yakınsama Eğrisi (50 İterasyon)", fontweight="bold")
    ax1.legend()
    ax1.grid(alpha=0.3)
    
    ax2.plot(iters, nfeat, "g-s", markersize=3, linewidth=2)
    ax2.set_ylabel("Seçilen Özellik Sayısı")
    ax2.set_xlabel("İterasyon")
    ax2.set_title("Seçilen Özellik Sayısı — İterasyona Göre", fontweight="bold")
    ax2.grid(alpha=0.3)
    fig.tight_layout()
    return save_figure(fig, "07_bso_convergence")

# 8. Optimizer Convergence Comparison
def fig_optimizer_comparison():
    fig, ax = plt.subplots(figsize=(10, 6))
    colors_o = {"BSO": C_PRIMARY, "PSO": C_DANGER, "GA": C_SUCCESS, "GWO": C_PURPLE}
    for name, data in OPTIMIZER_CONV.items():
        ax.plot(range(len(data)), data, label=name, color=colors_o[name], linewidth=2)
    ax.set_xlabel("İterasyon")
    ax.set_ylabel("En İyi Fitness")
    ax.set_title("Metaheuristik Yakınsama Karşılaştırması", fontweight="bold")
    ax.legend()
    ax.grid(alpha=0.3)
    fig.tight_layout()
    return save_figure(fig, "08_optimizer_comparison")

# 9. Confusion Matrix Heatmaps (BSO-RF + XGBoost)
def fig_confusion_matrices():
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    for ax, (key, title) in zip(axes, [("BSO-RF","BSO-Hybrid RF"), ("XGBoost","XGBoost")]):
        cm = np.array(CONFUSION_MATRICES[key])
        im = ax.imshow(cm, cmap="Blues", interpolation="nearest")
        ax.set_title(f"Karışıklık Matrisi — {title}", fontweight="bold")
        ax.set_xticks(range(5))
        ax.set_yticks(range(5))
        short_labels = ["Backdoor","Benign","ACK_Frag","SYN_Flood","PortScan"]
        ax.set_xticklabels(short_labels, rotation=45, ha="right", fontsize=7)
        ax.set_yticklabels(short_labels, fontsize=7)
        ax.set_xlabel("Tahmin")
        ax.set_ylabel("Gerçek")
        for i in range(5):
            for j in range(5):
                color = "white" if cm[i,j] > cm.max()*0.5 else "black"
                ax.text(j, i, str(cm[i,j]), ha="center", va="center", fontsize=7, color=color)
        fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    return save_figure(fig, "09_confusion_matrices")

# 10. All 12 Confusion Matrices
def fig_all_confusion_matrices():
    keys = list(CONFUSION_MATRICES.keys())
    n = len(keys)
    cols = 4
    rows_n = math.ceil(n / cols)
    fig, axes = plt.subplots(rows_n, cols, figsize=(18, rows_n*4))
    axes = axes.flatten()
    short_labels = ["Bkd","Ben","ACK","SYN","PS"]
    for idx, key in enumerate(keys):
        cm = np.array(CONFUSION_MATRICES[key])
        ax = axes[idx]
        im = ax.imshow(cm, cmap="Blues", interpolation="nearest")
        ax.set_title(key, fontweight="bold", fontsize=9)
        ax.set_xticks(range(5))
        ax.set_yticks(range(5))
        ax.set_xticklabels(short_labels, fontsize=6)
        ax.set_yticklabels(short_labels, fontsize=6)
        for i in range(5):
            for j in range(5):
                color = "white" if cm[i,j] > cm.max()*0.5 else "black"
                ax.text(j, i, str(cm[i,j]), ha="center", va="center", fontsize=5, color=color)
    for idx in range(n, len(axes)):
        axes[idx].set_visible(False)
    fig.suptitle("Tüm 12 Model — Karışıklık Matrisleri", fontweight="bold", fontsize=13)
    fig.tight_layout(rect=[0,0,1,0.96])
    return save_figure(fig, "10_all_confusion_matrices")

# 11. Cross-Validation Box Plot
def fig_cv_boxplot():
    folds = CROSS_VALIDATION["folds"]
    metrics = {"Accuracy": [f["accuracy"] for f in folds],
               "Precision": [f["precision"] for f in folds],
               "Recall": [f["recall"] for f in folds],
               "F1 Score": [f["f1Score"] for f in folds]}
    fig, ax = plt.subplots(figsize=(8, 5))
    bp = ax.boxplot(metrics.values(), labels=metrics.keys(), patch_artist=True,
                    boxprops=dict(facecolor="#dbeafe", edgecolor=C_PRIMARY),
                    medianprops=dict(color=C_DANGER, linewidth=2))
    ax.set_ylabel("Skor (%)")
    ax.set_title("BSO-Hybrid RF — 10-Katlı Çapraz Doğrulama", fontweight="bold")
    ax.grid(axis="y", alpha=0.3)
    mean = CROSS_VALIDATION["mean"]
    ax.axhline(y=mean["accuracy"], color=C_PRIMARY, linestyle="--", alpha=0.5, label=f'Ort. Acc: {mean["accuracy"]:.2f}%')
    ax.legend(fontsize=8)
    fig.tight_layout()
    return save_figure(fig, "11_cv_boxplot")

# 12. CV Bar Chart per Fold
def fig_cv_bars():
    folds = CROSS_VALIDATION["folds"]
    fold_nums = [f["fold"] for f in folds]
    accs = [f["accuracy"] for f in folds]
    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(fold_nums, accs, color=C_PRIMARY, edgecolor="white")
    ax.axhline(y=CROSS_VALIDATION["mean"]["accuracy"], color=C_DANGER, linestyle="--", linewidth=2,
               label=f'Ortalama: {CROSS_VALIDATION["mean"]["accuracy"]:.2f}%')
    ax.set_xlabel("Katman (Fold)")
    ax.set_ylabel("Accuracy (%)")
    ax.set_title("10-Katlı CV — Katman Bazında Accuracy", fontweight="bold")
    ax.set_ylim(89.5, 92)
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    for bar, v in zip(bars, accs):
        ax.text(bar.get_x()+bar.get_width()/2, v+0.03, f"{v:.2f}%", ha="center", fontsize=7)
    fig.tight_layout()
    return save_figure(fig, "12_cv_bars")

# 13. Accuracy vs Training Time Scatter
def fig_accuracy_vs_time():
    fig, ax = plt.subplots(figsize=(10, 6))
    for i, m in enumerate(MODELS):
        ax.scatter(m["trainingTime"], m["accuracy"], s=m["featuresUsed"]*8,
                   color=MODEL_COLORS[i % len(MODEL_COLORS)], edgecolors="black", linewidth=0.5, zorder=5)
        ax.annotate(m["name"].replace(" (Proposed)","").split("(")[0].strip(),
                    (m["trainingTime"], m["accuracy"]),
                    textcoords="offset points", xytext=(5,5), fontsize=6)
    ax.set_xlabel("Eğitim Süresi (s)")
    ax.set_ylabel("Accuracy (%)")
    ax.set_title("Accuracy vs Eğitim Süresi (boyut = özellik sayısı)", fontweight="bold")
    ax.grid(alpha=0.3)
    fig.tight_layout()
    return save_figure(fig, "13_accuracy_vs_time")

# 14. ROC AUC Bar Chart
def fig_roc_auc():
    names = list(ROC_AUC_DATA.keys())
    aucs  = list(ROC_AUC_DATA.values())
    short = [n.replace(" (Proposed)","").replace("Logistic Regression","Log. Reg.").replace("K-Nearest Neighbors","KNN") for n in names]
    fig, ax = plt.subplots(figsize=(12, 5))
    colors = [C_PRIMARY if i==0 else MODEL_COLORS[i%len(MODEL_COLORS)] for i in range(len(names))]
    bars = ax.bar(range(len(names)), aucs, color=colors, edgecolor="white")
    ax.set_xticks(range(len(names)))
    ax.set_xticklabels(short, rotation=45, ha="right", fontsize=7)
    ax.set_ylabel("AUC-ROC")
    ax.set_title("Macro-Averaged ROC-AUC Karşılaştırması", fontweight="bold")
    ax.set_ylim(0.85, 1.0)
    ax.grid(axis="y", alpha=0.3)
    for bar, v in zip(bars, aucs):
        ax.text(bar.get_x()+bar.get_width()/2, v+0.002, f"{v:.4f}", ha="center", fontsize=6)
    fig.tight_layout()
    return save_figure(fig, "14_roc_auc")

# 15. Per-Class Performance (BSO-RF)
def fig_per_class():
    names = [c["className"].replace("DDoS-","") for c in BSO_RF_PER_CLASS]
    prec = [c["precision"] for c in BSO_RF_PER_CLASS]
    rec  = [c["recall"] for c in BSO_RF_PER_CLASS]
    f1   = [c["f1Score"] for c in BSO_RF_PER_CLASS]
    x = np.arange(len(names))
    w = 0.25
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(x-w, prec, w, label="Precision", color=C_PRIMARY)
    ax.bar(x, rec, w, label="Recall", color=C_SUCCESS)
    ax.bar(x+w, f1, w, label="F1-Score", color=C_WARNING)
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=20, ha="right")
    ax.set_ylabel("Skor (%)")
    ax.set_title("BSO-Hybrid RF — Sınıf Bazında Performans", fontweight="bold")
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return save_figure(fig, "15_per_class")

# 16. MCC Comparison
def fig_mcc_comparison():
    names = [m["name"].replace(" (Proposed)","") for m in MODELS]
    mccs  = [m["mcc"] for m in MODELS]
    fig, ax = plt.subplots(figsize=(12, 5))
    colors = [C_PRIMARY if i==0 else C_GRAY for i in range(len(MODELS))]
    bars = ax.bar(range(len(names)), mccs, color=colors, edgecolor="white")
    ax.set_xticks(range(len(names)))
    ax.set_xticklabels(names, rotation=45, ha="right", fontsize=7)
    ax.set_ylabel("MCC")
    ax.set_title("Matthews Correlation Coefficient (MCC) Karşılaştırması", fontweight="bold")
    ax.grid(axis="y", alpha=0.3)
    for bar, v in zip(bars, mccs):
        ax.text(bar.get_x()+bar.get_width()/2, v+0.005, f"{v:.4f}", ha="center", fontsize=7)
    fig.tight_layout()
    return save_figure(fig, "16_mcc_comparison")

# 17. Statistical Significance Heatmap
def fig_stat_heatmap():
    labels = [t["comparison"].replace("BSO-Hybrid vs ","") for t in STATISTICAL_TESTS]
    p_vals = [t["pValue"] for t in STATISTICAL_TESTS]
    cohen_d = [abs(t["cohenD"]) for t in STATISTICAL_TESTS]
    
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5))
    # p-value
    colors_p = ["#22c55e" if p < 0.05 else "#ef4444" for p in p_vals]
    ax1.barh(labels, [-math.log10(max(p, 1e-300)) for p in p_vals], color=colors_p)
    ax1.axvline(x=-math.log10(0.05), color="red", linestyle="--", label="α=0.05")
    ax1.set_xlabel("-log₁₀(p-değeri)")
    ax1.set_title("İstatistiksel Anlamlılık (p-değeri)", fontweight="bold")
    ax1.legend(fontsize=8)
    
    # Cohen's d
    ax2.barh(labels, cohen_d, color=[C_PRIMARY if d>0.8 else C_WARNING if d>0.5 else C_GRAY for d in cohen_d])
    ax2.axvline(x=0.8, color="red", linestyle="--", label="Büyük etki (d=0.8)")
    ax2.set_xlabel("|Cohen's d|")
    ax2.set_title("Etki Büyüklüğü (Cohen's d)", fontweight="bold")
    ax2.legend(fontsize=8)
    fig.tight_layout()
    return save_figure(fig, "17_statistical_tests")

# 18. Noise Robustness
def fig_noise_robustness():
    noise = DYNAMIC_ENV["noiseRobustness"]
    levels = [n["noiseLevel"] for n in noise]
    accs   = [n["accuracy"] for n in noise]
    f1s    = [n["f1Macro"] for n in noise]
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.plot(levels, accs, "b-o", label="Accuracy", linewidth=2, markersize=6)
    ax.plot(levels, f1s, "r-s", label="F1-Macro", linewidth=2, markersize=6)
    ax.set_xlabel("Gürültü Seviyesi")
    ax.set_ylabel("Skor (%)")
    ax.set_title("Gürültü Dayanıklılığı — BSO-Hybrid RF", fontweight="bold")
    ax.legend()
    ax.grid(alpha=0.3)
    fig.tight_layout()
    return save_figure(fig, "18_noise_robustness")

# 19. Learning Curve
def fig_learning_curve():
    lc = DYNAMIC_ENV["learningCurve"]
    fracs = [l["fraction"]*100 for l in lc]
    accs  = [l["accuracy"] for l in lc]
    f1s   = [l["f1Macro"] for l in lc]
    times = [l["trainingTime"] for l in lc]
    
    fig, ax1 = plt.subplots(figsize=(8, 5))
    ax1.plot(fracs, accs, "b-o", label="Accuracy", linewidth=2)
    ax1.plot(fracs, f1s, "g-s", label="F1-Macro", linewidth=2)
    ax1.set_xlabel("Eğitim Verisi (%)")
    ax1.set_ylabel("Skor (%)")
    ax1.set_title("Öğrenme Eğrisi — BSO-Hybrid RF", fontweight="bold")
    ax1.legend(loc="center left")
    ax1.grid(alpha=0.3)
    
    ax2 = ax1.twinx()
    ax2.plot(fracs, times, "r--^", label="Eğitim Süresi (s)", linewidth=1.5, alpha=0.7)
    ax2.set_ylabel("Eğitim Süresi (s)")
    ax2.legend(loc="center right")
    fig.tight_layout()
    return save_figure(fig, "19_learning_curve")

# 20. Throughput
def fig_throughput():
    th = DYNAMIC_ENV["throughput"]
    batches = [t["batchSize"] for t in th]
    sps     = [t["samplesPerSecond"] for t in th]
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.plot(batches, sps, "b-o", linewidth=2, markersize=8)
    ax.set_xlabel("Batch Boyutu")
    ax.set_ylabel("Örnek/Saniye")
    ax.set_title("Tahmin Throughput — BSO-Hybrid RF", fontweight="bold")
    ax.grid(alpha=0.3)
    for b, s in zip(batches, sps):
        ax.annotate(f"{s:,}", (b, s), textcoords="offset points", xytext=(0, 10), ha="center", fontsize=7)
    fig.tight_layout()
    return save_figure(fig, "20_throughput")

# 21. Feature Category Pie Chart
def fig_feature_categories():
    categories = {
        "Ağ İstatistikleri": ["Header_Length","Rate","Tot sum","Max","Number","Time_To_Live"],
        "Protokol Bayrakları": ["fin_flag_number","psh_flag_number","syn_count","cwr_flag_number"],
        "Uygulama Protokolleri": ["HTTP","HTTPS","DNS","SSH","DHCP","IGMP","LLC"],
        "Taşıma Protokolleri": ["UDP","ARP"],
    }
    sizes = [len(v) for v in categories.values()]
    labels = list(categories.keys())
    colors = [C_PRIMARY, C_SUCCESS, C_WARNING, C_PURPLE]
    
    fig, ax = plt.subplots(figsize=(8, 6))
    wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct="%1.0f%%",
                                       colors=colors, startangle=90,
                                       wedgeprops=dict(edgecolor="white", linewidth=2))
    for t in autotexts:
        t.set_fontsize(9)
        t.set_fontweight("bold")
    ax.set_title("BSO Seçilmiş Özellikler — Kategori Dağılımı", fontweight="bold")
    fig.tight_layout()
    return save_figure(fig, "21_feature_categories")

# 22. Composite Score Ranking
def fig_composite_ranking():
    # Compute composite score (normalized)
    metrics_keys = ["accuracy","precision","recall","f1Score","f1Macro","aucRoc"]
    scores = []
    for m in MODELS:
        vals = [m[k] for k in metrics_keys]
        score = sum(vals) / len(vals)
        scores.append((m["name"].replace(" (Proposed)",""), score))
    scores.sort(key=lambda x: x[1], reverse=True)
    
    names = [s[0] for s in scores]
    vals  = [s[1] for s in scores]
    colors = [C_PRIMARY if i==0 else C_GRAY for i in range(len(scores))]
    # find BSO index
    for i, (n,v) in enumerate(scores):
        if "BSO-Hybrid" in n:
            colors[i] = C_PRIMARY
    
    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.barh(range(len(names)-1,-1,-1), vals, color=colors, edgecolor="white")
    ax.set_yticks(range(len(names)-1,-1,-1))
    ax.set_yticklabels(names, fontsize=8)
    ax.set_xlabel("Bileşik Skor (%)")
    ax.set_title("Model Sıralaması — Bileşik Performans Skoru", fontweight="bold")
    ax.grid(axis="x", alpha=0.3)
    for bar, v in zip(bars, vals):
        ax.text(v+0.1, bar.get_y()+bar.get_height()/2, f"{v:.2f}", va="center", fontsize=7)
    fig.tight_layout()
    return save_figure(fig, "22_composite_ranking")

# 23. Feature Selection Methods Comparison Radar
def fig_fs_comparison_radar():
    methods = ["BSO","PSO","GA","GWO"]
    categories = ["Accuracy","F1-Macro","Özellik Azaltma (%)","1-Fitness×500", "Değerlendirme×0.001"]
    # BSO-RF, PSO-RF, GA-RF, GWO-RF
    data_map = {
        "BSO": [89.82, 84.24, 51.3, (1-0.177801)*500, 1177*0.001],
        "PSO": [88.35, 81.82, 53.8, (1-0.193895)*500, 820*0.001],
        "GA":  [89.37, 83.66, 46.2, (1-0.188982)*500, 722*0.001],
        "GWO": [89.80, 84.35, 41.0, (1-0.192181)*500, 820*0.001],
    }
    angles = np.linspace(0, 2*np.pi, len(categories), endpoint=False).tolist()
    angles += angles[:1]
    
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
    colors_r = [C_PRIMARY, C_DANGER, C_SUCCESS, C_PURPLE]
    for i, method in enumerate(methods):
        vals = data_map[method]
        vals += vals[:1]
        ax.plot(angles, vals, "o-", linewidth=2, label=method, color=colors_r[i])
        ax.fill(angles, vals, alpha=0.1, color=colors_r[i])
    ax.set_thetagrids(np.degrees(angles[:-1]), categories, fontsize=8)
    ax.set_title("Özellik Seçimi Yöntemleri Karşılaştırması", fontweight="bold", y=1.08)
    ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.1))
    fig.tight_layout()
    return save_figure(fig, "23_fs_comparison_radar")

# 24. System Architecture Diagram
def fig_system_architecture():
    fig, ax = plt.subplots(figsize=(14, 6))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 6)
    ax.set_axis_off()
    
    steps = [
        ("CICIoT2023\nVeri Seti", 1, 3, "#dbeafe"),
        ("Ön İşleme\n& SMOTE", 3.5, 3, "#dcfce7"),
        ("BSO-Hybrid\nOptimizasyon", 6, 3, "#fef3c7"),
        ("RF\nSınıflandırma", 8.5, 3, "#ede9fe"),
        ("DDoS\nTespiti", 11, 3, "#fce7f3"),
    ]
    for label, x, y, color in steps:
        rect = mpatches.FancyBboxPatch((x-0.9, y-0.7), 1.8, 1.4,
                                        boxstyle="round,pad=0.2",
                                        facecolor=color, edgecolor="#374151", linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, y, label, ha="center", va="center", fontsize=9, fontweight="bold")
    
    # Arrows
    for i in range(len(steps)-1):
        x1 = steps[i][1] + 0.9
        x2 = steps[i+1][1] - 0.9
        ax.annotate("", xy=(x2, 3), xytext=(x1, 3),
                    arrowprops=dict(arrowstyle="->", lw=2, color="#374151"))
    
    ax.set_title("Sistem Mimarisi — BSO-Hybrid RF DDoS Tespit Sistemi", fontweight="bold", fontsize=13, y=0.95)
    # Sub-labels
    sublabels = [
        "19 CSV dosyası\n118,466 akış", "StandardScaler\nSMOTE dengeleme",
        "Yarasa Sürüsü\n19/39 özellik", "n_estimators=266\nmax_depth=20",
        "5 sınıf\n%89.82 doğruluk"
    ]
    for (label, x, y, color), sub in zip(steps, sublabels):
        ax.text(x, y-1.3, sub, ha="center", va="center", fontsize=7, color="#6b7280")
    
    fig.tight_layout()
    return save_figure(fig, "24_system_architecture")

# =============================================================================
# MAIN DOCUMENT GENERATOR
# =============================================================================
def generate_document():
    print("=" * 70)
    print("BSO-Hybrid RF Dashboard — Kapsamlı Word Export")
    print("=" * 70)
    
    doc = Document()
    
    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # ══════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════
    print("[1/24] Kapak sayfası...")
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BSO-Hibrit RF ile Geliştirilmiş DDoS Tespiti")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(30, 58, 95)
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("CICIoT2023 Veri Seti Üzerinde Kapsamlı Dashboard Raporu")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("SHUAIB AYAD JASIM")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(30, 58, 95)
    
    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info2.add_run("Danışman: Dr. Saim Ervural\nKTO Karatay Üniversitesi")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("Dashboard Tam Export — Tüm Tablolar, Şekiller ve Grafikler")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (Manual)
    # ══════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "İÇİNDEKİLER", 1)
    toc_items = [
        "1. Genel Bakış ve Temel Metrikler",
        "2. Sistem Mimarisi",
        "3. Veri Seti Keşifsel Analizi (EDA)",
        "4. SMOTE Dengeleme Analizi",
        "5. BSO Algoritması ve Parametreleri",
        "6. BSO Yakınsama Eğrisi",
        "7. Metaheuristik Karşılaştırması",
        "8. Özellik Seçimi Analizi",
        "9. Özellik Önem Sıralaması",
        "10. 39 vs 19 Özellik Karşılaştırması",
        "11. Model Sonuçları (12 Model)",
        "12. Model Accuracy Karşılaştırması",
        "13. Bileşik Skor Sıralaması",
        "14. Radar Karşılaştırması (Top 4)",
        "15. Accuracy vs Eğitim Süresi",
        "16. ROC-AUC Karşılaştırması",
        "17. MCC Karşılaştırması",
        "18. Sınıf Bazında Performans",
        "19. Karışıklık Matrisleri",
        "20. Tüm Karışıklık Matrisleri",
        "21. Çapraz Doğrulama (10-Katlı CV)",
        "22. İstatistiksel Anlamlılık Testleri",
        "23. Gürültü Dayanıklılığı",
        "24. Öğrenme Eğrisi",
        "25. Throughput Analizi",
        "26. Bilinmeyen Saldırı Tespiti",
        "27. Hesaplama Verimliliği",
        "28. Literatür Karşılaştırması",
        "29. Özellik Kategori Analizi",
        "30. Özellik Seçimi Yöntemleri Radar",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 1. OVERVIEW & KEY METRICS
    # ══════════════════════════════════════════════════════════════════
    print("[2/24] Genel bakış ve temel metrikler...")
    add_heading_styled(doc, "1. Genel Bakış ve Temel Metrikler", 1)
    
    doc.add_paragraph(
        "Bu rapor, BSO-Hybrid RF (Bat Swarm Optimization — Hibrit Random Forest) tabanlı DDoS tespit "
        "sisteminin kapsamlı dashboard içeriğini dışa aktarır. Sistem, CICIoT2023 veri seti üzerinde "
        "geliştirilmiş olup 5 sınıflı çok-sınıflı sınıflandırma gerçekleştirir."
    )
    
    add_heading_styled(doc, "1.1 Temel Performans Metrikleri", 2)
    headers = ["Metrik", "Değer"]
    rows = [
        ["Accuracy", "89.82%"],
        ["F1-Score (Weighted)", "89.90%"],
        ["F1-Macro", "84.24%"],
        ["AUC-ROC", "98.38%"],
        ["MCC", "0.8676"],
        ["Özellik Azaltma", "51.3% (39 → 19)"],
        ["Eğitim Süresi", "3.742 s"],
        ["Tahmin Süresi", "0.0043 ms/örnek"],
    ]
    add_styled_table(doc, headers, rows, col_widths=[6, 6])
    
    add_heading_styled(doc, "1.2 Veri Seti Özeti", 2)
    headers2 = ["Özellik", "Değer"]
    rows2 = [
        ["Veri Seti", "CICIoT2023"],
        ["Toplam Özellik", "39"],
        ["Seçilen Özellik", "19 (BSO)"],
        ["Toplam Örnek", f'{DATASET_STATS["totalSamples"]:,}'],
        ["Sınıf Sayısı", "5"],
        ["Eğitim Seti", f'{DATASET_STATS["training"]:,}'],
        ["Doğrulama Seti", f'{DATASET_STATS["validation"]:,}'],
        ["Test Seti", f'{DATASET_STATS["testing"]:,}'],
        ["Veri Bölme", "70/10/20 (stratified)"],
        ["SMOTE", "Uygulandı (72,252 → 87,500)"],
    ]
    add_styled_table(doc, headers2, rows2, col_widths=[6, 6])
    
    add_heading_styled(doc, "1.3 Ön İşleme Adımları", 2)
    steps = [
        "1. CICIoT2023'ten 19 CSV dosyası yüklendi",
        "2. Çoğunluk sınıfları için rastgele alt örnekleme (sınıf başına 25.000)",
        "3. Stratified eğitim/doğrulama/test bölme (70/10/20)",
        "4. Eğitim setine SMOTE aşırı örnekleme (72.252 → 87.500)",
        "5. StandardScaler normalizasyonu",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Bullet")
    
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════
    print("[3/24] Sistem mimarisi...")
    add_heading_styled(doc, "2. Sistem Mimarisi", 1)
    doc.add_paragraph(
        "Aşağıdaki şekil, BSO-Hybrid RF DDoS tespit sisteminin genel mimarisini göstermektedir. "
        "Sistem beş ana aşamadan oluşur: veri toplama, ön işleme, BSO optimizasyonu, RF sınıflandırma ve DDoS tespiti."
    )
    path = fig_system_architecture()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: BSO-Hybrid RF Sistem Mimarisi")
    
    # System pipeline table
    add_heading_styled(doc, "2.1 Sistem Boru Hattı", 2)
    headers = ["Aşama", "Açıklama", "Detay"]
    rows = [
        ["1. Veri Toplama", "CICIoT2023 veri seti", "19 CSV, 118,466 akış, 39 özellik"],
        ["2. Ön İşleme", "Temizleme ve dengeleme", "StandardScaler + SMOTE"],
        ["3. BSO Optimizasyon", "Özellik seçimi + hiperparametre", "50 iterasyon, 25 popülasyon"],
        ["4. RF Sınıflandırma", "Random Forest eğitimi", "266 ağaç, max_depth=20"],
        ["5. DDoS Tespiti", "Gerçek zamanlı sınıflandırma", "5 sınıf, 0.0043 ms/örnek"],
    ]
    add_styled_table(doc, headers, rows, col_widths=[3.5, 4.5, 5])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 3. DATASET EDA
    # ══════════════════════════════════════════════════════════════════
    print("[4/24] Veri seti EDA...")
    add_heading_styled(doc, "3. Veri Seti Keşifsel Analizi (EDA)", 1)
    
    add_heading_styled(doc, "3.1 Sınıf Dağılımı", 2)
    doc.add_paragraph(
        "CICIoT2023 veri seti 5 saldırı sınıfı içermektedir. Backdoor_Malware sınıfı azınlık sınıfıdır "
        "(2.252 örnek), SMOTE ile 17.500'e dengelenmiştir."
    )
    path = fig_class_distribution()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: SMOTE Öncesi ve Sonrası Sınıf Dağılımı")
    
    add_heading_styled(doc, "3.2 Sınıf Detayları Tablosu", 2)
    headers = ["Sınıf", "Eğitim (Öncesi)", "SMOTE Sonrası", "Test", "Renk"]
    rows = []
    for a in ATTACK_TYPES:
        rows.append([a["name"], f'{a["trainingSamples"]:,}', f'{a["smoteSamples"]:,}',
                     f'{a["testingSamples"]:,}', a["color"]])
    add_styled_table(doc, headers, rows, col_widths=[4, 3, 3, 2.5, 2])
    
    add_heading_styled(doc, "3.3 Tüm 39 Özellik Listesi", 2)
    headers_f = ["#", "Özellik Adı", "Tür", "Açıklama"]
    rows_f = [[str(i+1), FEATURES_39[i], "numeric", "Ağ trafik özelliği"] for i in range(39)]
    add_styled_table(doc, headers_f, rows_f, col_widths=[1, 4, 2, 5.5])
    
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 4. SMOTE ANALYSIS
    # ══════════════════════════════════════════════════════════════════
    print("[5/24] SMOTE analizi...")
    add_heading_styled(doc, "4. SMOTE Dengeleme Analizi", 1)
    doc.add_paragraph(
        "SMOTE (Synthetic Minority Over-sampling Technique) algoritması, azınlık sınıfını sentetik "
        "örnekler oluşturarak dengelemek için kullanılmıştır. Toplam 15.248 sentetik örnek üretilmiştir."
    )
    path = fig_smote_comparison()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: SMOTE Dengeleme Karşılaştırması (Grouped Bar)")
    
    headers = ["Parametre", "Değer"]
    rows = [
        ["Orijinal azınlık sayısı", f'{DATASET_STATS["originalMinorityCount"]:,}'],
        ["Sentetik örnek sayısı", f'{DATASET_STATS["smoteSyntheticSamples"]:,}'],
        ["Dengeli sınıf sayısı", f'{DATASET_STATS["balancedClassCount"]:,}'],
        ["Eğitim öncesi toplam", "72,252"],
        ["Eğitim sonrası toplam", f'{DATASET_STATS["training"]:,}'],
    ]
    add_styled_table(doc, headers, rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 5. BSO ALGORITHM PARAMETERS
    # ══════════════════════════════════════════════════════════════════
    print("[6/24] BSO parametreleri...")
    add_heading_styled(doc, "5. BSO Algoritması ve Parametreleri", 1)
    doc.add_paragraph(
        "Bat Swarm Optimization (BSO) algoritması, eşzamanlı özellik seçimi ve Random Forest "
        "hiperparametre optimizasyonu için kullanılmıştır. Algoritma 50 iterasyon boyunca "
        "25 yarasa popülasyonu ile çalıştırılmıştır."
    )
    
    add_heading_styled(doc, "5.1 Algoritma Parametreleri", 2)
    headers = ["Parametre", "Değer"]
    params = [
        ["Popülasyon Boyutu", str(BSO_PARAMS["populationSize"])],
        ["Maksimum İterasyon", str(BSO_PARAMS["maxIterations"])],
        ["Frekans Aralığı", f'{BSO_PARAMS["frequencyMin"]} — {BSO_PARAMS["frequencyMax"]}'],
        ["Başlangıç Yüksekliği (A₀)", str(BSO_PARAMS["initialLoudness"])],
        ["Başlangıç Nabız Oranı (r₀)", str(BSO_PARAMS["initialPulseRate"])],
        ["α (loudness decay)", str(BSO_PARAMS["alpha"])],
        ["γ (pulse rate increase)", str(BSO_PARAMS["gamma"])],
        ["Boyut (toplam özellik)", str(BSO_PARAMS["dimensions"])],
        ["Seçilen boyut", str(BSO_PARAMS["selectedDimensions"])],
        ["Yakınsama eşiği", str(BSO_PARAMS["convergenceThreshold"])],
        ["Fitness fonksiyonu", BSO_PARAMS["fitnessFunction"]],
        ["Toplam değerlendirme", str(BSO_PARAMS["totalEvaluations"])],
        ["Optimizasyon süresi", f'{BSO_PARAMS["optimizationTime"]:.2f} s'],
    ]
    add_styled_table(doc, headers, params)
    
    add_heading_styled(doc, "5.2 Optimize Edilen RF Hiperparametreleri", 2)
    headers_hp = ["Hiperparametre", "Optimum Değer", "Arama Aralığı"]
    hp_rows = [
        ["n_estimators", str(BSO_PARAMS["n_estimators"]), "[50, 400]"],
        ["max_depth", str(BSO_PARAMS["max_depth"]), "[5, 35]"],
        ["min_samples_split", str(BSO_PARAMS["min_samples_split"]), "[2, 15]"],
        ["max_features", str(BSO_PARAMS["max_features"]), "[0.3, 1.0]"],
    ]
    add_styled_table(doc, headers_hp, hp_rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 6. BSO CONVERGENCE
    # ══════════════════════════════════════════════════════════════════
    print("[7/24] BSO yakınsama...")
    add_heading_styled(doc, "6. BSO Yakınsama Eğrisi", 1)
    doc.add_paragraph(
        "BSO-Hybrid algoritması 50 iterasyon boyunca yakınsamıştır. En iyi fitness değeri "
        "0.184825'ten 0.177801'e düşmüştür. Seçilen özellik sayısı 27'den 19'a azalmıştır."
    )
    path = fig_bso_convergence()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: BSO-Hybrid Yakınsama Eğrisi ve Özellik Sayısı Değişimi")
    
    # Convergence data table (sample)
    add_heading_styled(doc, "6.1 Yakınsama Verileri (Seçilmiş İterasyonlar)", 2)
    headers = ["İterasyon", "En İyi Fitness", "Ortalama Fitness", "Seçilen Özellik"]
    rows = []
    for it, best, avg, nf in BSO_CONVERGENCE:
        if it % 5 == 0 or it == 49:
            rows.append([str(it), f"{best:.6f}", f"{avg:.6f}", str(nf)])
    add_styled_table(doc, headers, rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 7. OPTIMIZER COMPARISON
    # ══════════════════════════════════════════════════════════════════
    print("[8/24] Metaheuristik karşılaştırması...")
    add_heading_styled(doc, "7. Metaheuristik Yakınsama Karşılaştırması", 1)
    doc.add_paragraph(
        "BSO, PSO, GA ve GWO algoritmalarının yakınsama performansları karşılaştırılmıştır. "
        "BSO-Hybrid en düşük fitness değerine (0.177801) ulaşmıştır."
    )
    path = fig_optimizer_comparison()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: Metaheuristik Yakınsama Karşılaştırması")
    
    headers = ["Optimizer", "İterasyon", "Popülasyon", "Son Fitness", "Seçilen Özellik", "Süre (s)", "Değerlendirme"]
    rows = [
        ["BSO-Hybrid", "50", "25", "0.177801", "19", "840.43", "1177"],
        ["PSO", "40", "20", "0.193895", "18", "68.91", "820"],
        ["GA", "40", "20", "0.188982", "21", "60.47", "722"],
        ["GWO", "40", "20", "0.192181", "23", "70.23", "820"],
    ]
    add_styled_table(doc, headers, rows, highlight_row=0)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 8. FEATURE SELECTION ANALYSIS
    # ══════════════════════════════════════════════════════════════════
    print("[9/24] Özellik seçimi analizi...")
    add_heading_styled(doc, "8. Özellik Seçimi Analizi", 1)
    
    add_heading_styled(doc, "8.1 BSO vs PSO vs GA vs GWO — Seçilen Özellikler", 2)
    for method, data in FEATURE_SELECTION_COMP.items():
        doc.add_paragraph(f"{method}: {data['nSelected']} özellik — {', '.join(data['features'])}", style="List Bullet")
    
    add_heading_styled(doc, "8.2 Yöntem Karşılaştırma Tablosu", 2)
    headers = ["Yöntem", "Seçilen", "Azaltma (%)", "Fitness", "Değerlendirme", "Süre (s)"]
    rows = []
    for method, data in FEATURE_SELECTION_COMP.items():
        rows.append([method, str(data["nSelected"]), f'{data["reductionPct"]:.1f}',
                     f'{data["bestFitness"]:.6f}', str(data["evaluations"]), f'{data["time"]:.2f}'])
    add_styled_table(doc, headers, rows, highlight_row=0)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 9. FEATURE IMPORTANCE
    # ══════════════════════════════════════════════════════════════════
    print("[10/24] Özellik önem sıralaması...")
    add_heading_styled(doc, "9. Özellik Önem Sıralaması", 1)
    path = fig_feature_importance()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: BSO Seçilmiş 19 Özellik — Önem Sıralaması")
    
    headers = ["Sıra", "Özellik", "Önem Skoru", "Orijinal İndeks"]
    rows = [[str(f[0]), f[1], f"{f[2]:.6f}", str(f[3])] for f in BSO_SELECTED]
    add_styled_table(doc, headers, rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 10. 39 vs 19 FEATURES
    # ══════════════════════════════════════════════════════════════════
    print("[11/24] 39 vs 19 özellik...")
    add_heading_styled(doc, "10. 39 Özellik vs 19 Özellik Karşılaştırması", 1)
    doc.add_paragraph(
        "BSO-Hybrid RF yalnızca 19 özellik kullanarak, 39 özellik kullanan standart Random Forest "
        "ile karşılaştırılabilir performans elde etmiştir. Bu, %51.3 özellik azaltma sağlamıştır."
    )
    path = fig_39_vs_19()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: 39 vs 19 Özellik Performans Karşılaştırması")
    
    headers = ["Metrik", "RF (39 özellik)", "BSO-RF (19 özellik)", "Fark"]
    rows = [
        ["Accuracy", "89.74%", "89.82%", "+0.08%"],
        ["Precision", "89.92%", "90.19%", "+0.27%"],
        ["Recall", "89.74%", "89.82%", "+0.08%"],
        ["F1 Score", "89.77%", "89.90%", "+0.13%"],
        ["F1 Macro", "84.13%", "84.24%", "+0.11%"],
        ["AUC-ROC", "98.36%", "98.38%", "+0.02%"],
        ["Özellik Sayısı", "39", "19", "-51.3%"],
    ]
    add_styled_table(doc, headers, rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 11. MODEL RESULTS (ALL 12)
    # ══════════════════════════════════════════════════════════════════
    print("[12/24] 12 model sonuçları...")
    add_heading_styled(doc, "11. Model Sonuçları (12 Model)", 1)
    doc.add_paragraph(
        "Toplam 12 makine öğrenmesi modeli değerlendirilmiştir: 5 metaheuristik-optimized model "
        "(BSO-RF, BSO-SVM, PSO-RF, GA-RF, GWO-RF) ve 7 temel model "
        "(RF, SVM, DT, KNN, NB, LR, XGBoost)."
    )
    
    add_heading_styled(doc, "11.1 Temel Performans Metrikleri", 2)
    headers = ["Model", "Acc (%)", "Prec (%)", "Rec (%)", "F1 (%)", "F1-M (%)", "AUC (%)", "Özellik"]
    rows = []
    for m in MODELS:
        rows.append([m["name"].replace(" (Proposed)","*"),
                     f'{m["accuracy"]:.2f}', f'{m["precision"]:.2f}', f'{m["recall"]:.2f}',
                     f'{m["f1Score"]:.2f}', f'{m["f1Macro"]:.2f}', f'{m["aucRoc"]:.2f}',
                     str(m["featuresUsed"])])
    add_styled_table(doc, headers, rows, highlight_row=0)
    doc.add_paragraph("* = Önerilen model (BSO-Hybrid RF)", style="List Bullet")
    
    add_heading_styled(doc, "11.2 Ek Metrikler", 2)
    headers2 = ["Model", "MCC", "FPR (%)", "FNR (%)", "Specificity (%)", "Eğitim (s)", "Tahmin (ms)"]
    rows2 = []
    for m in MODELS:
        rows2.append([m["name"].replace(" (Proposed)","*"),
                      f'{m["mcc"]:.4f}', f'{m["fpr"]:.2f}', f'{m["fnr"]:.2f}',
                      f'{m["specificity"]:.2f}', f'{m["trainingTime"]:.3f}',
                      f'{m["predictionTime"]:.4f}'])
    add_styled_table(doc, headers2, rows2, highlight_row=0)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 12. MODEL ACCURACY CHART
    # ══════════════════════════════════════════════════════════════════
    print("[13/24] Model accuracy grafiği...")
    add_heading_styled(doc, "12. Model Accuracy Karşılaştırması", 1)
    path = fig_model_accuracy()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: 12 Model Accuracy Karşılaştırması")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 13. COMPOSITE RANKING
    # ══════════════════════════════════════════════════════════════════
    print("[14/24] Bileşik skor sıralaması...")
    add_heading_styled(doc, "13. Bileşik Skor Sıralaması", 1)
    doc.add_paragraph(
        "Bileşik skor, 6 temel metriğin (Accuracy, Precision, Recall, F1, F1-Macro, AUC-ROC) "
        "ortalaması olarak hesaplanmıştır."
    )
    path = fig_composite_ranking()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: Model Bileşik Performans Skor Sıralaması")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 14. RADAR TOP 4
    # ══════════════════════════════════════════════════════════════════
    print("[15/24] Radar karşılaştırması...")
    add_heading_styled(doc, "14. Radar Karşılaştırması (En İyi 4 Model)", 1)
    path = fig_radar_top4()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: En İyi 4 Model Radar Karşılaştırması")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 15. ACCURACY vs TRAINING TIME
    # ══════════════════════════════════════════════════════════════════
    print("[16/24] Accuracy vs eğitim süresi...")
    add_heading_styled(doc, "15. Accuracy vs Eğitim Süresi", 1)
    path = fig_accuracy_vs_time()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: Accuracy vs Eğitim Süresi (Scatter)")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 16. ROC-AUC
    # ══════════════════════════════════════════════════════════════════
    print("[17/24] ROC-AUC karşılaştırması...")
    add_heading_styled(doc, "16. ROC-AUC Karşılaştırması", 1)
    path = fig_roc_auc()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: Macro-Averaged ROC-AUC Karşılaştırması")
    
    headers = ["Model", "AUC-ROC"]
    rows = [[k.replace(" (Proposed)","*"), f"{v:.4f}"] for k, v in ROC_AUC_DATA.items()]
    add_styled_table(doc, headers, rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 17. MCC COMPARISON
    # ══════════════════════════════════════════════════════════════════
    print("[18/24] MCC karşılaştırması...")
    add_heading_styled(doc, "17. MCC Karşılaştırması", 1)
    doc.add_paragraph(
        "Matthews Correlation Coefficient (MCC), dengesiz veri setleri için güvenilir bir metrik olup "
        "-1 ile +1 arasında değer alır. +1 mükemmel sınıflandırmayı gösterir."
    )
    path = fig_mcc_comparison()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: MCC Karşılaştırması")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 18. PER-CLASS PERFORMANCE
    # ══════════════════════════════════════════════════════════════════
    print("[19/24] Sınıf bazında performans...")
    add_heading_styled(doc, "18. Sınıf Bazında Performans (BSO-Hybrid RF)", 1)
    path = fig_per_class()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: BSO-Hybrid RF Sınıf Bazında Performans")
    
    headers = ["Sınıf", "Precision (%)", "Recall (%)", "F1-Score (%)", "Destek"]
    rows = [[c["className"], f'{c["precision"]:.2f}', f'{c["recall"]:.2f}',
             f'{c["f1Score"]:.2f}', str(c["support"])] for c in BSO_RF_PER_CLASS]
    add_styled_table(doc, headers, rows)
    doc.add_paragraph(
        "Not: DDoS-ACK_Fragmentation ve DDoS-SYN_Flood sınıfları neredeyse mükemmel performans "
        "gösterirken, Backdoor_Malware sınıfı en düşük performansa sahiptir (az örnek sayısı nedeniyle)."
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 19 & 20. CONFUSION MATRICES
    # ══════════════════════════════════════════════════════════════════
    print("[20/24] Karışıklık matrisleri...")
    add_heading_styled(doc, "19. Karışıklık Matrisleri (BSO-RF & XGBoost)", 1)
    path = fig_confusion_matrices()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: BSO-Hybrid RF ve XGBoost Karışıklık Matrisleri")
    
    add_heading_styled(doc, "20. Tüm 12 Model — Karışıklık Matrisleri", 1)
    path = fig_all_confusion_matrices()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: Tüm 12 Model Karışıklık Matrisleri", width=Inches(6.5))
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 21. CROSS-VALIDATION
    # ══════════════════════════════════════════════════════════════════
    print("[21/24] Çapraz doğrulama...")
    add_heading_styled(doc, "21. Çapraz Doğrulama (10-Katlı Stratified CV)", 1)
    
    path = fig_cv_boxplot()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: 10-Katlı CV Box Plot")
    
    path = fig_cv_bars()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: CV Katman Bazında Accuracy")
    
    add_heading_styled(doc, "21.1 Katman Sonuçları", 2)
    headers = ["Katman", "Accuracy (%)", "Precision (%)", "Recall (%)", "F1-Score (%)"]
    rows = [[str(f["fold"]), f'{f["accuracy"]:.2f}', f'{f["precision"]:.2f}',
             f'{f["recall"]:.2f}', f'{f["f1Score"]:.2f}'] for f in CROSS_VALIDATION["folds"]]
    mean = CROSS_VALIDATION["mean"]
    std = CROSS_VALIDATION["std"]
    rows.append(["Ortalama", f'{mean["accuracy"]:.2f}', f'{mean["precision"]:.2f}',
                 f'{mean["recall"]:.2f}', f'{mean["f1Score"]:.2f}'])
    rows.append(["Std", f'{std["accuracy"]:.3f}', f'{std["precision"]:.3f}',
                 f'{std["recall"]:.3f}', f'{std["f1Score"]:.3f}'])
    add_styled_table(doc, headers, rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 22. STATISTICAL SIGNIFICANCE
    # ══════════════════════════════════════════════════════════════════
    print("[22/24] İstatistiksel anlamlılık...")
    add_heading_styled(doc, "22. İstatistiksel Anlamlılık Testleri", 1)
    doc.add_paragraph(
        "10-katlı çapraz doğrulama sonuçları üzerinde paired t-test ve Wilcoxon signed-rank testleri "
        "uygulanmıştır. Cohen's d etki büyüklüğü de raporlanmıştır."
    )
    
    path = fig_stat_heatmap()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: İstatistiksel Anlamlılık ve Etki Büyüklüğü")
    
    headers = ["Karşılaştırma", "İyileştirme", "t-İstatistik", "p-Değeri", "Anlamlı?", "Cohen's d", "Etki"]
    rows = []
    for t in STATISTICAL_TESTS:
        rows.append([t["comparison"].replace("BSO-Hybrid vs ","vs "),
                     t["improvement"], f'{t["tStatistic"]:.4f}',
                     f'{t["pValue"]:.6f}' if t["pValue"] > 0 else "< 0.000001",
                     "Evet" if t["significant"] else "Hayır",
                     f'{t["cohenD"]:.4f}', t["effectSize"]])
    add_styled_table(doc, headers, rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 23. NOISE ROBUSTNESS
    # ══════════════════════════════════════════════════════════════════
    print("[23/24] Dinamik ortam testleri...")
    add_heading_styled(doc, "23. Gürültü Dayanıklılığı", 1)
    path = fig_noise_robustness()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: Gürültü Dayanıklılığı — BSO-Hybrid RF")
    
    headers = ["Gürültü Seviyesi", "Accuracy (%)", "F1-Macro (%)", "Performans Düşüşü (%)"]
    rows = [[f'{nr["noiseLevel"]:.2f}', f'{nr["accuracy"]:.2f}', f'{nr["f1Macro"]:.2f}',
             f'{nr["degradation"]:.2f}'] for nr in DYNAMIC_ENV["noiseRobustness"]]
    add_styled_table(doc, headers, rows)
    
    # ── LEARNING CURVE ──
    add_heading_styled(doc, "24. Öğrenme Eğrisi", 1)
    path = fig_learning_curve()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: Öğrenme Eğrisi — BSO-Hybrid RF")
    
    headers = ["Veri Oranı", "Örnek Sayısı", "Accuracy (%)", "F1-Macro (%)", "Eğitim Süresi (s)"]
    rows = [[f'{l["fraction"]*100:.0f}%', f'{l["nSamples"]:,}', f'{l["accuracy"]:.2f}',
             f'{l["f1Macro"]:.2f}', f'{l["trainingTime"]:.3f}'] for l in DYNAMIC_ENV["learningCurve"]]
    add_styled_table(doc, headers, rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 25. THROUGHPUT
    # ══════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "25. Throughput Analizi", 1)
    path = fig_throughput()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: Tahmin Throughput — BSO-Hybrid RF")
    
    headers = ["Batch Boyutu", "Ortalama Süre (ms)", "Örnek/Saniye"]
    rows = [[f'{t["batchSize"]:,}', f'{t["avgTimeMs"]:.2f}', f'{t["samplesPerSecond"]:,}']
            for t in DYNAMIC_ENV["throughput"]]
    add_styled_table(doc, headers, rows)
    
    # ── UNKNOWN ATTACK ──
    add_heading_styled(doc, "26. Bilinmeyen Saldırı Tespiti", 1)
    doc.add_paragraph(
        "Her seferinde bir saldırı sınıfı eğitimden çıkarılarak bilinmeyen saldırı tespit oranı ölçülmüştür."
    )
    headers = ["Çıkarılan Saldırı", "Tespit Oranı (%)", "Bilinmeyen Örnek"]
    rows = [[u["excluded"], f'{u["detectionRate"]:.2f}', str(u["samples"])]
            for u in DYNAMIC_ENV["unknownAttack"]]
    add_styled_table(doc, headers, rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 27. COMPUTATIONAL EFFICIENCY
    # ══════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "27. Hesaplama Verimliliği", 1)
    headers = ["Model", "Eğitim (s)", "Tahmin (ms)", "Özellik", "Özellik Seti"]
    rows = []
    for m in MODELS:
        rows.append([m["name"].replace(" (Proposed)","*"), f'{m["trainingTime"]:.3f}',
                     f'{m["predictionTime"]:.4f}', str(m["featuresUsed"]), m["featureSet"]])
    add_styled_table(doc, headers, rows, highlight_row=0)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 28. LITERATURE COMPARISON
    # ══════════════════════════════════════════════════════════════════
    print("[24/24] Literatür karşılaştırması ve son bölümler...")
    add_heading_styled(doc, "28. Literatür Karşılaştırması", 1)
    doc.add_paragraph(
        "Mevcut çalışmalar genellikle ikili sınıflandırma (benign vs attack) kullanırken, "
        "bu çalışma 5 sınıflı çok-sınıflı sınıflandırma gerçekleştirmektedir."
    )
    headers = ["Çalışma", "Veri Seti", "Yöntem", "Accuracy (%)", "F1 (%)", "Not"]
    rows = [[s["paper"], s["dataset"], s["method"], f'{s["accuracy"]:.2f}',
             f'{s["f1Score"]:.2f}', s["note"]] for s in STATE_OF_THE_ART]
    add_styled_table(doc, headers, rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 29. FEATURE CATEGORY ANALYSIS
    # ══════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "29. Özellik Kategori Analizi", 1)
    path = fig_feature_categories()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: BSO Seçilmiş Özellikler — Kategori Dağılımı")
    
    # ══════════════════════════════════════════════════════════════════
    # 30. FEATURE SELECTION METHODS RADAR
    # ══════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "30. Özellik Seçimi Yöntemleri Radar Karşılaştırması", 1)
    path = fig_fs_comparison_radar()
    n = next_fig()
    add_figure(doc, path, f"Şekil {n}: Özellik Seçimi Yöntemleri Radar Karşılaştırması")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SUMMARY & CONCLUSIONS
    # ══════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Sonuçlar ve Değerlendirme", 1)
    
    conclusions = [
        "BSO-Hybrid RF, %51.3 özellik azaltma ile karşılaştırılabilir performans elde etmiştir (%89.82 accuracy, 19/39 özellik).",
        "XGBoost en yüksek accuracy'ye sahiptir (%90.37) ancak 39 özellik kullanmaktadır.",
        "BSO-Hybrid RF, 10-katlı CV'de tutarlı performans göstermiştir (std=0.204%).",
        "DDoS-ACK_Fragmentation ve DDoS-SYN_Flood sınıfları ~%100 F1-Score ile tespit edilmiştir.",
        "Backdoor_Malware sınıfı en düşük performansa sahiptir (F1=%57.40) — az örnek sayısı nedeniyle.",
        "BSO, PSO, GA ve GWO'ya göre daha iyi fitness değerine ulaşmıştır (0.177801).",
        "Sistem gerçek zamanlı tahmin yapabilir: 231,833 örnek/saniye (batch=20.644).",
        "İstatistiksel testler, BSO-Hybrid RF'nin çoğu modele göre anlamlı iyileştirme sağladığını doğrulamıştır.",
    ]
    for c in conclusions:
        doc.add_paragraph(c, style="List Bullet")
    
    doc.add_paragraph()
    
    # Research contributions
    add_heading_styled(doc, "Araştırma Katkıları", 2)
    contributions = [
        "Yarasa Sürüsü Optimizasyonu (BSO) ile eşzamanlı özellik seçimi ve hiperparametre optimizasyonu.",
        "CICIoT2023 veri setinde 5 sınıflı çok-sınıflı DDoS tespiti.",
        "12 makine öğrenmesi modelinin kapsamlı karşılaştırmalı analizi.",
        "Gürültü dayanıklılığı ve bilinmeyen saldırı tespiti deneyleri.",
    ]
    for c in contributions:
        doc.add_paragraph(c, style="List Number")
    
    # ── Footer info ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("─" * 50)
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    info_final = doc.add_paragraph()
    info_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_final.add_run(
        "Bu rapor, BSO-Hybrid RF DDoS Tespit Dashboard'undan otomatik olarak oluşturulmuştur.\n"
        f"Toplam {fig_counter[0]} şekil ve 20+ tablo içermektedir.\n"
        "GitHub: https://github.com/Shuaib5352/DDoS-BSO-Thesis"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)

    # ── Save ──
    doc.save(str(OUT_DOCX))
    size_kb = OUT_DOCX.stat().st_size / 1024
    print()
    print("=" * 70)
    print(f"  Başarıyla oluşturuldu: {OUT_DOCX}")
    print(f"  Boyut: {size_kb:.0f} KB ({size_kb/1024:.1f} MB)")
    print(f"  Şekil sayısı: {fig_counter[0]}")
    print(f"  Şekil klasörü: {FIG_DIR}")
    print("=" * 70)
    return str(OUT_DOCX)


if __name__ == "__main__":
    generate_document()
